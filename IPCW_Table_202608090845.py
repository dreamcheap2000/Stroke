#!/usr/bin/env python3
"""
IPCW_Table_202608090845.py
Unadjusted vs IPCW-adjusted LassoCV predictor coefficients for 6MWT4,
extended with bootstrap selection frequency per Bootstrap_resample.docx.

Bootstrap_resample.docx instructs:
  - Track, for each predictor, the number of bootstrap resamples (out of 1 000)
    in which that predictor's coefficient is non-zero.
  - Report selection_frequency = count_nonzero / n_bootstrap (0 to 1) for
    both unadjusted and IPCW-adjusted models as a new column in Table S3
    alongside β, 95% CI, and p-value.
  - Require BOTH p < 0.05 AND selection_frequency ≥ 0.70 before bolding a
    coefficient as robustly significant (red bold).  A low selection frequency
    (< 0.5) indicates the p-value is likely unstable/unreliable even if
    nominally significant.

Changes from IPCW_Code_202608082227.py:
  1. boot_unadj / boot_ipcw matrices now store every resample's coefficients.
  2. sel_freq_unadj / sel_freq_ipcw computed as (coef != 0).mean(axis=0).
  3. Results table gains two new columns: sf_unadj, sf_ipcw.
  4. Significance flagging: p < 0.05 AND sf_ipcw >= SELECTION_FREQ_THRESHOLD.
  5. Table gains two new sub-columns for "Sel. Freq." in each model block.
  6. Legend updated to explain Sel. Freq. and the dual-criterion significance flag.

Outputs: IPCW_Table_202608090845.docx
"""

import os
import numpy as np
import pandas as pd
from sklearn.linear_model import LassoCV
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Significance threshold for selection frequency ────────────────────────────
SELECTION_FREQ_THRESHOLD = 0.70  # per Bootstrap_resample.docx (0.70–0.80)

# ── 1. Load data ──────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
artifacts_df = pd.read_csv(os.path.join(SCRIPT_DIR,
                            'Tip_Over_Analysis_artifacts_out_202608081955.csv'))

# ── 2. Feature sets (same as the reference code) ──────────────────────────────
demographics = ['cov_Age', 'cov_Sex, F0 M1']
stroke_info  = ['cov_Dissection', 'cov_ACA', 'cov_Undetermined',
                'cov_HemorrhageStroke', 'cov_Side_Right', 'cov_Side_Left',
                'cov_Side_Bilateral', 'cov_Loc_CortSub', 'cov_Loc_Subcortical',
                'cov_Loc_Infratentorial', 'cov_LVS', 'cov_LVO']
comorbidities = ['cov_AF', 'cov_DM', 'cov_HTN', 'cov_Dyslipidemia', 'cov_CAD',
                 'cov_CKD', 'cov_RestrictiveLung', 'cov_GIUlcer',
                 'cov_LiverCirrhosis', 'cov_Hepatitis', 'cov_Parkinsonism',
                 'cov_Malignancy', 'cov_OldStroke', 'cov_Dementia',
                 'cov_Psychiatric', 'cov_Gout']
nihss_out_cols = ['cov_ConsOut', 'cov_AnswerOut', 'cov_OrderOut', 'cov_EOMOut',
                  'cov_VisualOut', 'cov_FacialOut', 'cov_LUOut', 'cov_RUOut',
                  'cov_LLOut', 'cov_RLOut', 'cov_Coordinateout', 'cov_SensoryOut',
                  'cov_LanguageOut', 'cov_ArticulateOut', 'cov_NeglectOut']
acute_complications = ['cov_Pneumonia', 'cov_UTI', 'cov_GIB', 'cov_Cellulitis']
acute_treatment     = ['cov_tPA', 'cov_IA', 'cov_tPAIA']
functional_t1_plus_gs = ['cov_MRS1', 'cov_BI1', 'cov_FOIS1', 'cov_MNA1',
                          'cov_EuroQoL5D1', 'cov_IADL1', 'cov_BBS1',
                          'cov_Gait_Speed_1', 'cov_FuglUE1', 'cov_FuglSEN1',
                          'cov_CCAT1']

features_raw = (demographics + stroke_info + comorbidities +
                functional_t1_plus_gs + nihss_out_cols +
                acute_complications + acute_treatment)
features = [f for f in features_raw if f in artifacts_df.columns]
missing_feats = set(features_raw) - set(features)
if missing_feats:
    print(f'Features not found (excluded): {missing_feats}')

target = 'outcome_6mwt4'

# ── 3. Completer subset ───────────────────────────────────────────────────────
model_df = artifacts_df[artifacts_df['completer'] == 1].copy()
model_df = model_df[features + [target, 'ipcw']].dropna()
print(f'Sample size (completers with full data): N = {len(model_df)}')

X_raw = model_df[features].values
y     = model_df[target].values
w     = model_df['ipcw'].values

scaler = StandardScaler()
X      = scaler.fit_transform(X_raw)

# ── 4. Fit unadjusted and IPCW-adjusted LassoCV ───────────────────────────────
RANDOM_STATE = 42
CV_FOLDS     = 5
MAX_ITER     = 20000

lasso_unadj = LassoCV(cv=CV_FOLDS, random_state=RANDOM_STATE, max_iter=MAX_ITER)
lasso_unadj.fit(X, y)
print(f'Unadjusted  α={lasso_unadj.alpha_:.4f}')

lasso_ipcw = LassoCV(cv=CV_FOLDS, random_state=RANDOM_STATE, max_iter=MAX_ITER)
lasso_ipcw.fit(X, y, sample_weight=w)
print(f'IPCW-adj    α={lasso_ipcw.alpha_:.4f}')

coef_unadj = lasso_unadj.coef_
coef_ipcw  = lasso_ipcw.coef_

# ── 5. Bootstrap CIs, p-values, and selection frequencies (B = 1 000) ─────────
# Per Bootstrap_resample.docx:
#   coef_matrix stores each resample's coefficients; selection_frequency is
#   computed as the proportion of resamples in which the coefficient is non-zero.
B   = 1000
n   = len(y)
n_p = len(features)
rng = np.random.default_rng(RANDOM_STATE)

boot_unadj = np.zeros((B, n_p))  # store each resample's coefficients
boot_ipcw  = np.zeros((B, n_p))

for b in range(B):
    idx = rng.integers(0, n, size=n)
    Xb, yb, wb = X[idx], y[idx], w[idx]

    m_u = LassoCV(cv=CV_FOLDS, random_state=RANDOM_STATE, max_iter=MAX_ITER)
    m_u.fit(Xb, yb)
    boot_unadj[b] = m_u.coef_

    m_i = LassoCV(cv=CV_FOLDS, random_state=RANDOM_STATE, max_iter=MAX_ITER)
    m_i.fit(Xb, yb, sample_weight=wb)
    boot_ipcw[b] = m_i.coef_

    if (b + 1) % 100 == 0:
        print(f'  Bootstrap {b + 1}/{B}')

# Selection frequency: proportion of resamples with non-zero coefficient
sel_freq_unadj = (boot_unadj != 0).mean(axis=0)
sel_freq_ipcw  = (boot_ipcw  != 0).mean(axis=0)


def ci95(boot_coefs_col):
    lo = np.percentile(boot_coefs_col, 2.5)
    hi = np.percentile(boot_coefs_col, 97.5)
    return lo, hi


def pvalue(point_est, boot_coefs_col):
    """Two-sided bootstrap p-value via sign-flip proportion."""
    if point_est >= 0:
        p = np.mean(boot_coefs_col <= 0)
    else:
        p = np.mean(boot_coefs_col >= 0)
    return float(np.clip(2 * p, 0, 1))


# ── 6. Assemble results table ─────────────────────────────────────────────────
rows = []
for i, feat in enumerate(features):
    c_u = coef_unadj[i]
    c_i = coef_ipcw[i]
    if c_u == 0 and c_i == 0:
        continue
    lo_u, hi_u = ci95(boot_unadj[:, i])
    lo_i, hi_i = ci95(boot_ipcw[:, i])
    p_u = pvalue(c_u, boot_unadj[:, i])
    p_i = pvalue(c_i, boot_ipcw[:, i])
    rows.append({
        'Predictor'    : feat,
        'beta_unadj'   : c_u,
        'ci_lo_unadj'  : lo_u,
        'ci_hi_unadj'  : hi_u,
        'p_unadj'      : p_u,
        'sf_unadj'     : sel_freq_unadj[i],   # NEW: selection frequency
        'beta_ipcw'    : c_i,
        'ci_lo_ipcw'   : lo_i,
        'ci_hi_ipcw'   : hi_i,
        'p_ipcw'       : p_i,
        'sf_ipcw'      : sel_freq_ipcw[i],    # NEW: selection frequency
    })

results_df = (pd.DataFrame(rows)
              .sort_values('beta_ipcw', ascending=False)
              .reset_index(drop=True))
print(f'\nNon-zero predictors: {len(results_df)}')

# ── 7. Write DOCX table ───────────────────────────────────────────────────────
def fmt_coef(v):
    return f'{v:.3f}' if v != 0 else '0.000'

def fmt_ci(lo, hi):
    return f'[{lo:.3f}, {hi:.3f}]'

def fmt_p(p):
    if p < 0.001:
        return '<0.001'
    return f'{p:.3f}'

def fmt_sf(sf):
    """Format selection frequency as a proportion to 2 decimal places."""
    return f'{sf:.2f}'


def shade_row(row, hex_color):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    trPr.append(shd)


doc = Document()

# Title
title_para = doc.add_paragraph()
title_run  = title_para.add_run(
    'Table. Unadjusted and IPCW-Adjusted LassoCV Predictor Coefficients '
    'for 6MWT at Discharge (T4)')
title_run.font.bold = True
title_run.font.size = Pt(10)

# Column layout (9 columns):
# 0: Predictor
# 1: β (unadj)  2: 95% CI (unadj)  3: p-value (unadj)  4: Sel. Freq. (unadj)
# 5: β (ipcw)   6: 95% CI (ipcw)   7: p-value (ipcw)   8: Sel. Freq. (ipcw)
NUM_COLS = 9
table = doc.add_table(rows=2, cols=NUM_COLS)
table.style = 'Table Grid'


def write_header_cell(cell, text, bold=True):
    cell.text = text
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.bold  = bold
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ---- header row 0 ----
hdr0 = table.rows[0]
shade_row(hdr0, '2E4057')

write_header_cell(hdr0.cells[0], 'Predictor')
write_header_cell(hdr0.cells[1], 'Unadjusted')
hdr0.cells[1].merge(hdr0.cells[2])
hdr0.cells[2].merge(hdr0.cells[3])
hdr0.cells[3].merge(hdr0.cells[4])   # span 4 cols for unadjusted
write_header_cell(hdr0.cells[5], 'IPCW-Adjusted')
hdr0.cells[5].merge(hdr0.cells[6])
hdr0.cells[6].merge(hdr0.cells[7])
hdr0.cells[7].merge(hdr0.cells[8])   # span 4 cols for ipcw

# ---- header row 1 ----
hdr1 = table.rows[1]
shade_row(hdr1, '4A90D9')
sub_headers = [
    'Predictor',
    'β', '95% CI', 'p-value', 'Sel. Freq.',
    'β', '95% CI', 'p-value', 'Sel. Freq.',
]
for j, txt in enumerate(sub_headers):
    write_header_cell(hdr1.cells[j], txt)

# ---- data rows ----
for ri, row_data in results_df.iterrows():
    row  = table.add_row()
    fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    shade_row(row, fill)

    # Dual-criterion significance (per Bootstrap_resample.docx):
    #   flag ONLY if p < 0.05 AND selection_frequency >= SELECTION_FREQ_THRESHOLD
    robust_sig = (
        row_data['p_ipcw'] < 0.05 and
        row_data['sf_ipcw'] >= SELECTION_FREQ_THRESHOLD
    )

    vals = [
        row_data['Predictor'].replace('cov_', ''),
        fmt_coef(row_data['beta_unadj']),
        fmt_ci(row_data['ci_lo_unadj'], row_data['ci_hi_unadj']),
        fmt_p(row_data['p_unadj']),
        fmt_sf(row_data['sf_unadj']),
        fmt_coef(row_data['beta_ipcw']),
        fmt_ci(row_data['ci_lo_ipcw'], row_data['ci_hi_ipcw']),
        fmt_p(row_data['p_ipcw']),
        fmt_sf(row_data['sf_ipcw']),
    ]
    for j, v in enumerate(vals):
        cell = row.cells[j]
        cell.text = v
        for para in cell.paragraphs:
            para.alignment = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                              else WD_ALIGN_PARAGRAPH.CENTER)
            for run in para.runs:
                run.font.size = Pt(9)
                # Red bold only for robustly significant IPCW p-value
                if j == 7 and robust_sig:
                    run.font.bold      = True
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                # Bold selection frequency when above threshold (both models)
                if j == 4 and row_data['sf_unadj'] >= SELECTION_FREQ_THRESHOLD:
                    run.font.bold = True
                if j == 8 and row_data['sf_ipcw'] >= SELECTION_FREQ_THRESHOLD:
                    run.font.bold = True

# ---- Legend ----
doc.add_paragraph()
legend = doc.add_paragraph()
legend.add_run('Legend. ').font.bold = True
legend.add_run(
    'β = standardized LassoCV regression coefficient (StandardScaler applied to '
    'all predictors). '
    '95% CI = bootstrap 95% confidence interval (1 000 resamples, percentile method). '
    'p-value = two-sided bootstrap p-value. '
    'Sel. Freq. = bootstrap selection frequency — proportion of 1 000 resamples in '
    'which the predictor\'s LassoCV coefficient was non-zero '
    '(selection_frequency = count_nonzero / 1 000; range 0–1). '
    'Unadjusted = LassoCV fitted on completers without sample weights. '
    'IPCW-Adjusted = LassoCV fitted on completers using stabilized '
    'inverse-probability-of-completion weights '
    '(pre-computed in Tip_Over_Analysis_artifacts_out_202608081955.csv). '
    'Only predictors with a non-zero coefficient in at least one model are shown. '
    f'Red bold p-values indicate robust statistical significance: both p < 0.05 '
    f'AND Sel. Freq. ≥ {SELECTION_FREQ_THRESHOLD:.2f} in the IPCW-adjusted model '
    f'(dual-criterion per Bootstrap_resample.docx; a low selection frequency '
    f'[< 0.50] indicates an unstable p-value even if nominally significant). '
    'Bold Sel. Freq. values indicate stability selection threshold met (≥ '
    f'{SELECTION_FREQ_THRESHOLD:.2f}).'
)
for run in legend.runs:
    run.font.size = Pt(8)

# ── 8. Save ───────────────────────────────────────────────────────────────────
out_docx = os.path.join(SCRIPT_DIR, 'IPCW_Table_202608090845.docx')
doc.save(out_docx)
print(f'\nTable saved to: {out_docx}')
print('\nPreview of results:')
print(results_df[['Predictor',
                   'beta_unadj', 'p_unadj', 'sf_unadj',
                   'beta_ipcw',  'p_ipcw',  'sf_ipcw']].to_string(index=False))
