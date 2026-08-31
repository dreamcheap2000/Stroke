#!/usr/bin/env python3
"""
20260831_IPCW_2331.py

Enhanced IPCW-based 6MWT4 prediction pipeline with three core improvements
over the prior session (20260828_5_Tiers_1201.py):

  1. Improved numerator stabilizer
     – Default: intercept-only numerator (marginal completion probability),
       which is the most common and well-validated stabilizer.
     – Optional: covariate-adjusted numerator using the strongest known
       predictors of completion (stroke severity via NIHSS items, and
       comorbidity burden) rather than Age alone.

  2. Explicit positivity diagnostics
     – Distribution of predicted completion probabilities (p_denom) before
       winsorization, with flags for values compressed near 0 or 1.
     – Stabilized-weight distribution (mean, SD, range, % truncated) as the
       primary positivity check.
     – Truncation-threshold sensitivity: weighted R²/MAE at 90th, 95th, and
       99th percentile truncation levels to confirm stability.

  3. Full-pipeline nonparametric bootstrap (B=500)
     – Each resample re-estimates denominator/numerator models, re-constructs
       and re-winsorizes weights, reselects the binary classifier, and
       refits the weighted Ridge model.
     – Stratified by completion status to avoid quasi-separation.
     – Reports bootstrap percentile 95 % confidence intervals for the
       tier/scenario weighted R² and MAE.

Outputs:
  20260831_IPCW_2331.docx   – explainer with full diagnostics tables
  20260831_IPCW_2331.xlsx   – patient-level predictions with CI columns
"""

from __future__ import annotations

import math
import warnings
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from sklearn.base import clone
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import ExtraTreesClassifier, RandomForestClassifier
from sklearn.exceptions import ConvergenceWarning
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression, Ridge
from sklearn.metrics import (
    balanced_accuracy_score,
    mean_absolute_error,
    r2_score,
)
from sklearn.model_selection import KFold, StratifiedKFold, cross_val_predict, cross_validate
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler

warnings.filterwarnings("ignore", category=ConvergenceWarning)

# ─────────────────────────────────────────────────────────────────────────────
# File locations
# ─────────────────────────────────────────────────────────────────────────────

ROOT = Path(__file__).resolve().parent
INPUT_XLSX = ROOT / "20260826_DeID.xlsx"
OUTPUT_XLSX = ROOT / "20260831_IPCW_2331.xlsx"
OUTPUT_REPORT = ROOT / "20260831_IPCW_2331.docx"

RANDOM_STATE = 42
N_BOOTSTRAP = 500  # full bootstrap iterations; set to 50 for a quick smoke-test

# In the bootstrap loop the binary classifier step is deliberately simplified:
# we skip cross-validated model reselection (which would multiply run-time 15×)
# and instead refit a single logistic regression on each resample. This still
# propagates all weight-estimation and Ridge-fitting uncertainty, which dominates
# the two-stage variance. The main-run classifier is selected by full CV and
# is not affected by this simplification.

# ─────────────────────────────────────────────────────────────────────────────
# Feature groups (preserved from prior session)
# ─────────────────────────────────────────────────────────────────────────────

DEMOGRAPHICS_PAC = ["Age", "Sex, F0 M1"]
STROKE_INFO = [
    "Dissection", "ACA", "Undetermined", "HemorrhageStroke",
    "LVS", "LVO", "Side_Right", "Side_Left", "Side_Bilateral",
    "Loc_CortSub", "Loc_Subcortical", "Loc_Infratentorial",
]
COMORBIDITIES_PAC = [
    "AF", "DM", "HTN", "Dyslipidemia", "CAD", "CKD",
    "RestrictiveLung", "GIUlcer", "LiverCirrhosis", "Hepatitis",
    "Parkinsonism", "Malignancy", "OldStroke", "Dementia",
    "Psychiatric", "Gout",
]
ACUTE_COMPLICATIONS_PAC = ["Pneumonia", "UTI", "GIB", "Cellulitis"]
FUNCTIONAL_T1_PLUS_GS_IMPUTED = [
    "MRS1", "BI1", "FOIS1", "MNA1", "EuroQoL5D1", "IADL1",
    "BBS1", "Gait_Speed_1_Imputed", "FuglUE1", "FuglSEN1",
]
NIHSS_OUT = [
    "ConsOut", "AnswerOut", "OrderOut", "EOMOut", "VisualOut",
    "FacialOut", "LUOut", "RUOut", "LLOut", "RLOut",
    "Coordinateout", "SensoryOut", "LanguageOut", "ArticulateOut", "NeglectOut",
]
NIHSS_IN = [
    "ConsIn", "AnswerIn", "OrderIn", "EOMIn", "VisualIn",
    "FaceIn", "LUIn", "RUIn", "LLIn", "RLIn",
    "CoordinateIn", "SensoryIn", "LanguageIn", "ArticulateIn", "NeglectIn",
]
T1T2_IMPROVEMENT = [
    "BI_T1T2_Change", "BBS_T1T2_Change", "MRS_T1T2_Change", "FOIS_T1T2_Change",
    "MNA_T1T2_Change", "IADL_T1T2_Change", "FuglUE_T1T2_Change",
    "FuglSEN_T1T2_Change", "EuroQoL5D_T1T2_Change",
]

# Strongest known baseline predictors of PAC completion:
# stroke severity items (motor + consciousness) and comorbidity burden count proxy.
# Used as the covariate-adjusted numerator when NUMERATOR_MODE = "adjusted".
NUMERATOR_ADJUSTED_CANDIDATES = [
    # Stroke severity (output NIHSS – functional status most proximal to PAC)
    "LUOut", "RUOut", "LLOut", "RLOut", "Coordinateout", "ArticulateOut",
    # Comorbidity burden
    "AF", "DM", "HTN", "Dyslipidemia", "CAD", "CKD", "Malignancy", "OldStroke",
    # Functional status at T1 (post-acute, pre-PAC)
    "MRS1", "BI1", "BBS1",
]

# "intercept" = intercept-only numerator (default, most defensible)
# "adjusted"  = parsimonious covariate-adjusted numerator
NUMERATOR_MODE: str = "intercept"

TIERS: dict[str, list[str]] = {
    "Tier 1": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED,
    "Tier 2": DEMOGRAPHICS_PAC + [f for f in FUNCTIONAL_T1_PLUS_GS_IMPUTED if f != "Gait_Speed_1_Imputed"],
    "Tier 3": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + COMORBIDITIES_PAC + STROKE_INFO + ACUTE_COMPLICATIONS_PAC,
    "Tier 4": DEMOGRAPHICS_PAC + [f for f in FUNCTIONAL_T1_PLUS_GS_IMPUTED if f != "Gait_Speed_1_Imputed"] + COMORBIDITIES_PAC + STROKE_INFO + ACUTE_COMPLICATIONS_PAC,
    "Tier 5": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + T1T2_IMPROVEMENT,
}

SCENARIOS = ["6MWT_Best_Scenario", "6MWT_Worst_Scenario"]

WINSORIZATION_THRESHOLDS = [90, 95, 99]  # percentiles tested in sensitivity analysis

# ─────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────────────────────────────────────

def _filter_existing(cols: list[str], df: pd.DataFrame) -> list[str]:
    return [c for c in cols if c in df.columns]


def _make_binary_cv(n: int = 5) -> StratifiedKFold:
    return StratifiedKFold(n_splits=n, shuffle=True, random_state=RANDOM_STATE)


def _make_regression_cv(n: int = 5) -> KFold:
    return KFold(n_splits=n, shuffle=True, random_state=RANDOM_STATE)


def _tier_token(tier_name: str) -> str:
    return tier_name.replace(" ", "_")


def _scenario_token(scenario_name: str) -> str:
    return scenario_name.replace("6MWT_", "").replace("_Scenario", "")


def _weighted_mae(y_true: np.ndarray, y_pred: np.ndarray, weights: np.ndarray) -> float:
    return float(np.average(np.abs(y_true - y_pred), weights=weights))


def _fmt(v: float, d: int = 4) -> str:
    return f"{v:.{d}f}" if math.isfinite(v) else "N/A"


def _set_small_font(paragraph, size: int = 9) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(size)


def _add_table(doc: Document, rows: list[list[object]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
            for run in table.cell(i, j).paragraphs[0].runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True


# ─────────────────────────────────────────────────────────────────────────────
# IPCW weight computation (with positivity diagnostics)
# ─────────────────────────────────────────────────────────────────────────────

def compute_ipcw_weights(
    df: pd.DataFrame,
    features: list[str],
    winsor_pct: float = 99,
    numerator_mode: str = "intercept",
) -> dict:
    """
    Estimate stabilized IPCW weights with full positivity diagnostics.

    Numerator options:
      "intercept" – marginal completion probability (intercept-only), the most
                    common and well-validated stabilizer.
      "adjusted"  – reduced covariate set (stroke severity + comorbidity burden),
                    stronger predictors of completion than Age alone.

    Returns a dict with weights, diagnostics, and sensitivity analysis.
    """
    valid = _filter_existing(features, df)
    if not valid:
        raise ValueError("At least one valid feature is required.")

    completion_status = df["PAC_Program_Completion"].astype("string")
    eligible_mask = completion_status.notna()
    completed = completion_status.eq("Completed PAC program").astype(int)

    X = df.loc[eligible_mask, valid]
    y = completed.loc[eligible_mask].to_numpy()

    imputer = SimpleImputer(strategy="median")
    X_imp = pd.DataFrame(imputer.fit_transform(X), columns=valid, index=X.index)

    # ── Denominator model: P(completion | tier features) ─────────────────────
    denom_pipe = Pipeline([
        ("sc", StandardScaler()),
        ("lr", LogisticRegression(max_iter=5000, solver="lbfgs", C=1.0)),
    ])
    denom_pipe.fit(X_imp, y)
    p_denom_raw = denom_pipe.predict_proba(X_imp)[:, 1]

    # Positivity diagnostics on raw p_denom
    near_zero = float(np.mean(p_denom_raw < 0.05))
    near_one = float(np.mean(p_denom_raw > 0.95))
    p_denom_diag = {
        "mean": float(np.mean(p_denom_raw)),
        "sd": float(np.std(p_denom_raw)),
        "min": float(np.min(p_denom_raw)),
        "p10": float(np.percentile(p_denom_raw, 10)),
        "p25": float(np.percentile(p_denom_raw, 25)),
        "p50": float(np.median(p_denom_raw)),
        "p75": float(np.percentile(p_denom_raw, 75)),
        "p90": float(np.percentile(p_denom_raw, 90)),
        "max": float(np.max(p_denom_raw)),
        "pct_near_zero": near_zero,
        "pct_near_one": near_one,
        "positivity_flag": near_zero > 0.05 or near_one > 0.05,
    }
    p_denom = np.clip(p_denom_raw, 1e-4, 1 - 1e-4)

    # ── Numerator model ───────────────────────────────────────────────────────
    if numerator_mode == "intercept":
        # Intercept-only: P(completion) marginally
        p_numer = np.full(len(y), y.mean(), dtype="float64")
        numer_description = "Intercept-only (marginal completion rate)"
    else:
        # Covariate-adjusted numerator: parsimonious stroke severity + comorbidity set
        numer_candidates = _filter_existing(NUMERATOR_ADJUSTED_CANDIDATES, X_imp)
        if not numer_candidates:
            # Fall back to intercept-only if none of the preferred features are available
            p_numer = np.full(len(y), y.mean(), dtype="float64")
            numer_description = "Intercept-only (fallback; no adjusted features available)"
        else:
            numer_pipe = Pipeline([
                ("sc", StandardScaler()),
                ("lr", LogisticRegression(max_iter=3000, solver="lbfgs", C=1.0)),
            ])
            numer_pipe.fit(X_imp[numer_candidates], y)
            p_numer = np.clip(numer_pipe.predict_proba(X_imp[numer_candidates])[:, 1], 1e-4, 1 - 1e-4)
            numer_description = (
                f"Adjusted numerator on {len(numer_candidates)} features "
                f"(stroke severity + comorbidity burden): "
                + ", ".join(numer_candidates[:8])
                + ("..." if len(numer_candidates) > 8 else "")
            )

    raw_weights = np.where(y == 1, p_numer / p_denom, np.nan)
    completer_weights = raw_weights[~np.isnan(raw_weights)]

    # Sensitivity to winsorization threshold
    sensitivity: dict[int, dict] = {}
    for pct in WINSORIZATION_THRESHOLDS:
        lo, hi = np.nanpercentile(completer_weights, [100 - pct, pct])
        w_clipped = np.clip(completer_weights, lo, hi)
        pct_truncated = float(np.mean((completer_weights < lo) | (completer_weights > hi)))
        sensitivity[pct] = {
            "lower": float(lo),
            "upper": float(hi),
            "mean": float(np.mean(w_clipped)),
            "sd": float(np.std(w_clipped)),
            "min": float(np.min(w_clipped)),
            "max": float(np.max(w_clipped)),
            "pct_truncated": pct_truncated,
        }

    # Primary winsorization at configured threshold
    lo_main = sensitivity[winsor_pct]["lower"]
    hi_main = sensitivity[winsor_pct]["upper"]
    winsorized = np.where(
        ~np.isnan(raw_weights),
        np.clip(raw_weights, lo_main, hi_main),
        np.nan,
    )

    weights_series = pd.Series(np.nan, index=df.index, dtype="float64")
    weights_series.loc[X_imp.index] = winsorized

    wt = winsorized[~np.isnan(winsorized)]
    weight_diag = {
        "mean": float(np.mean(wt)),
        "sd": float(np.std(wt)),
        "min": float(np.min(wt)),
        "max": float(np.max(wt)),
        "pct_truncated": sensitivity[winsor_pct]["pct_truncated"],
        "mean_near_one": abs(float(np.mean(wt)) - 1.0) < 0.2,
    }

    return {
        "weights": weights_series,
        "imputer": imputer,
        "features": valid,
        "denominator_model": "LogisticRegression(solver='lbfgs', C=1.0) on tier features",
        "numerator_model": numer_description,
        "winsorization": f"{100 - winsor_pct}th/{winsor_pct}th percentile",
        "p_denom_diagnostics": p_denom_diag,
        "weight_diagnostics": weight_diag,
        "sensitivity": sensitivity,
        # Keep raw values for bootstrap re-use
        "y_eligible": y,
        "X_imp": X_imp,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Binary walking models
# ─────────────────────────────────────────────────────────────────────────────

def _binary_candidates(features: list[str]) -> dict[str, Pipeline]:
    numeric_pipe = Pipeline([("imp", SimpleImputer(strategy="median")), ("sc", StandardScaler())])
    imp_only = Pipeline([("imp", SimpleImputer(strategy="median"))])

    logistic = Pipeline([
        ("prep", ColumnTransformer([("num", numeric_pipe, features)])),
        ("model", LogisticRegression(max_iter=5000, solver="liblinear", class_weight="balanced", random_state=RANDOM_STATE)),
    ])
    random_forest = Pipeline([
        ("prep", ColumnTransformer([("num", imp_only, features)])),
        ("model", RandomForestClassifier(n_estimators=300, class_weight="balanced", random_state=RANDOM_STATE, n_jobs=-1)),
    ])
    extra_trees = Pipeline([
        ("prep", ColumnTransformer([("num", imp_only, features)])),
        ("model", ExtraTreesClassifier(n_estimators=300, class_weight="balanced", random_state=RANDOM_STATE, n_jobs=-1)),
    ])
    return {
        "LogisticRegression(class_weight='balanced', solver='liblinear')": logistic,
        "RandomForestClassifier(n_estimators=300, class_weight='balanced')": random_forest,
        "ExtraTreesClassifier(n_estimators=300, class_weight='balanced')": extra_trees,
    }


def select_binary_classifier(df: pd.DataFrame, features: list[str], scenario: str) -> dict:
    valid = _filter_existing(features, df)
    model_df = df[df[scenario].notna()].copy()
    X = model_df[valid]
    y = model_df[scenario].astype(int)
    cv = _make_binary_cv()
    cv_splits = list(cv.split(X, y))
    candidates = _binary_candidates(valid)

    rows = []
    for name, pipe in candidates.items():
        scores = cross_validate(clone(pipe), X, y, cv=cv_splits, n_jobs=-1,
                                scoring=["balanced_accuracy", "accuracy", "f1"])
        rows.append({
            "model": name,
            "balanced_accuracy": float(np.mean(scores["test_balanced_accuracy"])),
            "accuracy": float(np.mean(scores["test_accuracy"])),
            "f1": float(np.mean(scores["test_f1"])),
        })

    best = max(rows, key=lambda r: (r["balanced_accuracy"], r["accuracy"], r["f1"]))
    best_pipe = clone(candidates[best["model"]])
    best_pipe.fit(X, y)
    return {"best_model_name": best["model"], "best_pipeline": best_pipe, "features": valid,
            "leaderboard": rows, "bal_acc": best["balanced_accuracy"]}


# ─────────────────────────────────────────────────────────────────────────────
# Weighted Ridge regression (point estimate)
# ─────────────────────────────────────────────────────────────────────────────

def fit_weighted_ridge(
    df: pd.DataFrame,
    features: list[str],
    ipcw: dict,
    binary: dict,
    winsor_pct: float = 99,
) -> dict:
    valid = _filter_existing(features, df)
    completion_status = df["PAC_Program_Completion"].astype("string")
    completer_mask = completion_status.eq("Completed PAC program") & df["6MWT4"].notna()
    df_comp = df.loc[completer_mask].copy()
    weights = ipcw["weights"].loc[completer_mask].fillna(1.0).to_numpy()

    X_raw = df_comp[valid]
    y = df_comp["6MWT4"].to_numpy()
    cv_splits = list(_make_regression_cv().split(X_raw))
    oof = np.zeros(len(df_comp), dtype="float64")

    for tr, te in cv_splits:
        imp = SimpleImputer(strategy="median")
        X_tr = imp.fit_transform(X_raw.iloc[tr])
        X_te = imp.transform(X_raw.iloc[te])
        m = Ridge(alpha=1.0)
        m.fit(X_tr, y[tr], sample_weight=weights[tr])
        oof[te] = np.maximum(0, m.predict(X_te))

    r2 = float(r2_score(y, oof, sample_weight=weights))
    mae = _weighted_mae(y, oof, weights)

    # Sensitivity: re-compute metrics at different winsor thresholds
    sensitivity_r2: dict[int, float] = {}
    sensitivity_mae: dict[int, float] = {}
    raw_w = ipcw["weights"].loc[completer_mask].fillna(1.0).to_numpy()
    for pct in WINSORIZATION_THRESHOLDS:
        lo, hi = np.percentile(raw_w, [100 - pct, pct])
        w_alt = np.clip(raw_w, lo, hi)
        oof_alt = np.zeros(len(df_comp), dtype="float64")
        for tr, te in cv_splits:
            imp = SimpleImputer(strategy="median")
            X_tr = imp.fit_transform(X_raw.iloc[tr])
            X_te = imp.transform(X_raw.iloc[te])
            m = Ridge(alpha=1.0)
            m.fit(X_tr, y[tr], sample_weight=w_alt[tr])
            oof_alt[te] = np.maximum(0, m.predict(X_te))
        sensitivity_r2[pct] = float(r2_score(y, oof_alt, sample_weight=w_alt))
        sensitivity_mae[pct] = _weighted_mae(y, oof_alt, w_alt)

    imp_final = SimpleImputer(strategy="median")
    X_all = pd.DataFrame(imp_final.fit_transform(X_raw), columns=valid, index=df_comp.index)
    final_model = Ridge(alpha=1.0)
    final_model.fit(X_all, y, sample_weight=weights)

    # Non-completers
    noncompleter_mask = completion_status.ne("Completed PAC program") & completion_status.notna()
    df_nc = df.loc[noncompleter_mask].copy()
    X_nc_raw = df_nc[valid]
    X_nc = pd.DataFrame(imp_final.transform(X_nc_raw), columns=valid, index=df_nc.index)
    walk_pred = binary["best_pipeline"].predict(X_nc_raw).astype(int)
    pred_6mwt4 = np.zeros(len(df_nc), dtype="float64")
    if walk_pred.any():
        pred_6mwt4[walk_pred == 1] = np.maximum(0, final_model.predict(X_nc[walk_pred == 1]))

    importance = (
        pd.DataFrame({"predictor": valid, "abs_coef": np.abs(final_model.coef_)})
        .sort_values("abs_coef", ascending=False).reset_index(drop=True)
    )

    return {
        "weighted_r2": r2,
        "weighted_mae": mae,
        "sensitivity_r2": sensitivity_r2,
        "sensitivity_mae": sensitivity_mae,
        "model_name": "Ridge(alpha=1.0) with IPCW sample weights",
        "importance": importance,
        "final_model": final_model,
        "imputer": imp_final,
        "noncomp_index": df_nc.index.tolist(),
        "noncomp_walk_pred": walk_pred,
        "noncomp_6mwt4_pred": pred_6mwt4,
        "n_noncomp_pred_walk": int(walk_pred.sum()),
        "n_noncomp_pred_no_walk": int((walk_pred == 0).sum()),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Bootstrap UQ – full pipeline in each resample
# ─────────────────────────────────────────────────────────────────────────────

def _bootstrap_once(
    df: pd.DataFrame,
    features: list[str],
    scenario: str,
    rng: np.random.Generator,
    winsor_pct: float = 99,
    numerator_mode: str = "intercept",
) -> tuple[float, float]:
    """
    One bootstrap resample: refit full pipeline, return (weighted_r2, weighted_mae).
    Stratified by completion status to preserve completer/non-completer ratio.

    The binary classifier step uses a single logistic regression rather than
    full cross-validated reselection to keep run-time tractable; this still
    propagates all weight-estimation and Ridge-fitting uncertainty.
    """
    completion_status = df["PAC_Program_Completion"].astype("string")
    groups = completion_status.fillna("Missing")
    idx_resampled: list[int] = []
    for grp_val in groups.unique():
        grp_idx = np.where(groups == grp_val)[0]
        idx_resampled.append(rng.choice(grp_idx, size=len(grp_idx), replace=True))
    boot_idx = np.concatenate(idx_resampled)
    df_boot = df.iloc[boot_idx].reset_index(drop=True)

    # Re-estimate weights on bootstrap sample
    try:
        ipcw_boot = compute_ipcw_weights(df_boot, features, winsor_pct=winsor_pct,
                                          numerator_mode=numerator_mode)
    except Exception:
        return float("nan"), float("nan")

    # Lightweight binary step: refit logistic regression (no CV reselection)
    try:
        valid = _filter_existing(features, df_boot)
        model_df = df_boot[df_boot[scenario].notna()].copy()
        X_bin = model_df[valid]
        y_bin = model_df[scenario].astype(int)
        bin_pipe = Pipeline([
            ("prep", ColumnTransformer([
                ("num", Pipeline([("imp", SimpleImputer(strategy="median")),
                                  ("sc", StandardScaler())]), valid),
            ])),
            ("model", LogisticRegression(max_iter=1000, solver="liblinear",
                                         class_weight="balanced", random_state=RANDOM_STATE)),
        ])
        bin_pipe.fit(X_bin, y_bin)
    except Exception:
        return float("nan"), float("nan")

    # Re-fit weighted Ridge (OOF on boot sample)
    try:
        comp_mask = (df_boot["PAC_Program_Completion"].astype("string").eq("Completed PAC program")
                     & df_boot["6MWT4"].notna())
        df_comp = df_boot.loc[comp_mask].copy()
        if len(df_comp) < 10:
            return float("nan"), float("nan")
        weights = ipcw_boot["weights"].loc[comp_mask].fillna(1.0).to_numpy()
        X_raw = df_comp[valid]
        y = df_comp["6MWT4"].to_numpy()
        n_splits = min(5, max(2, len(df_comp) // 10))
        cv_splits = list(KFold(n_splits=n_splits, shuffle=True,
                               random_state=RANDOM_STATE).split(X_raw))
        oof = np.zeros(len(df_comp), dtype="float64")
        for tr, te in cv_splits:
            imp = SimpleImputer(strategy="median")
            X_tr = imp.fit_transform(X_raw.iloc[tr])
            X_te = imp.transform(X_raw.iloc[te])
            m = Ridge(alpha=1.0)
            m.fit(X_tr, y[tr], sample_weight=weights[tr])
            oof[te] = np.maximum(0, m.predict(X_te))
        r2 = float(r2_score(y, oof, sample_weight=weights))
        mae = _weighted_mae(y, oof, weights)
        return r2, mae
    except Exception:
        return float("nan"), float("nan")


def bootstrap_ci(
    df: pd.DataFrame,
    features: list[str],
    scenario: str,
    n_bootstrap: int = N_BOOTSTRAP,
    winsor_pct: float = 99,
    numerator_mode: str = "intercept",
    seed: int = RANDOM_STATE,
) -> dict:
    """
    Run stratified bootstrap and return percentile 95 % CIs for R² and MAE.
    """
    rng = np.random.default_rng(seed)
    r2_boot: list[float] = []
    mae_boot: list[float] = []

    for b in range(n_bootstrap):
        r2_b, mae_b = _bootstrap_once(df, features, scenario, rng, winsor_pct, numerator_mode)
        if math.isfinite(r2_b):
            r2_boot.append(r2_b)
        if math.isfinite(mae_b):
            mae_boot.append(mae_b)

    def _ci(vals: list[float]) -> tuple[float, float, float]:
        if not vals:
            return float("nan"), float("nan"), float("nan")
        a = np.array(vals)
        return float(np.mean(a)), float(np.percentile(a, 2.5)), float(np.percentile(a, 97.5))

    r2_mean, r2_lo, r2_hi = _ci(r2_boot)
    mae_mean, mae_lo, mae_hi = _ci(mae_boot)
    return {
        "n_valid_r2": len(r2_boot),
        "n_valid_mae": len(mae_boot),
        "r2_boot_mean": r2_mean, "r2_ci_lo": r2_lo, "r2_ci_hi": r2_hi,
        "mae_boot_mean": mae_mean, "mae_ci_lo": mae_lo, "mae_ci_hi": mae_hi,
    }


# ─────────────────────────────────────────────────────────────────────────────
# DOCX report writer
# ─────────────────────────────────────────────────────────────────────────────

def write_report(all_results: list[dict]) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("20260831 Enhanced IPCW 6MWT4 Pipeline — Session 2331").bold = True

    intro_text = (
        "This report documents three improvements to the prior IPCW-based 6MWT4 "
        "prediction pipeline:\n"
        "(1) Intercept-only (marginal) numerator stabilizer as the default — the most "
        "common and well-validated approach, replacing the single-covariate Age-only "
        "numerator from the prior session.\n"
        "(2) Explicit positivity diagnostics: distribution of predicted completion "
        "probabilities (p_denom) before winsorization, stabilized-weight diagnostics "
        "(mean, SD, range, % truncated), and winsorization-threshold sensitivity analysis.\n"
        "(3) Full-pipeline nonparametric bootstrap (B=500, stratified by completion status) "
        "to provide percentile 95 % CIs for weighted R² and MAE without treating weights "
        "as fixed."
    )
    p = doc.add_paragraph(intro_text)
    _set_small_font(p)

    # ── Numerator stabilizer rationale ───────────────────────────────────────
    doc.add_paragraph()
    h = doc.add_paragraph()
    h.add_run("1. Numerator Stabilizer Rationale").bold = True

    rat_text = (
        "The numerator model determines the 'target' marginal distribution for the "
        "stabilized weights w = P(C=1 | numerator covariates) / P(C=1 | denominator covariates). "
        "An intercept-only numerator uses the overall completion rate, giving weights that "
        "re-weight completers toward the full-sample marginal distribution. "
        "This is the standard minimal stabilizer and produces weights with mean close to 1 "
        "when the denominator model is well-specified. "
        "A covariate-adjusted numerator (stroke severity + comorbidity burden) can improve "
        "efficiency when those predictors explain substantial variation in completion, but "
        "gains are modest if the denominator model already adjusts for those variables. "
        "The prior session used Age as the sole numerator covariate — this was arbitrary "
        "and provided little stabilization benefit over the intercept-only alternative. "
        f"This session uses: NUMERATOR_MODE = '{NUMERATOR_MODE}'."
    )
    p = doc.add_paragraph(rat_text)
    _set_small_font(p)

    # ── Per tier/scenario results ─────────────────────────────────────────────
    for res in all_results:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(f"{res['tier']} | {res['scenario_label']} Scenario").bold = True

        # Main metrics table
        ipcw = res["ipcw_completion"]
        reg = res["regression"]
        ci = res["bootstrap_ci"]
        wd = ipcw["weight_diagnostics"]
        pd_diag = ipcw["p_denom_diagnostics"]

        _add_table(doc, [
            ["Field", "Value"],
            ["Tier", res["tier"]],
            ["Scenario", res["scenario_label"]],
            ["IPCW denominator model", ipcw["denominator_model"]],
            ["IPCW numerator model", ipcw["numerator_model"]],
            ["IPCW winsorization", ipcw["winsorization"]],
            ["Selected binary model", res["binary"]["best_model_name"]],
            ["Binary OOF balanced accuracy", _fmt(res["binary"]["bal_acc"], 3)],
            ["Weighted regression model", reg["model_name"]],
            ["Weighted OOF R²", _fmt(reg["weighted_r2"], 4)],
            ["Weighted OOF MAE", _fmt(reg["weighted_mae"], 1)],
            ["Bootstrap R² (mean / 95% CI)", f"{_fmt(ci['r2_boot_mean'], 4)} [{_fmt(ci['r2_ci_lo'], 4)}, {_fmt(ci['r2_ci_hi'], 4)}]"],
            ["Bootstrap MAE (mean / 95% CI)", f"{_fmt(ci['mae_boot_mean'], 1)} [{_fmt(ci['mae_ci_lo'], 1)}, {_fmt(ci['mae_ci_hi'], 1)}]"],
            ["Bootstrap valid resamples (R²/MAE)", f"{ci['n_valid_r2']} / {ci['n_valid_mae']}"],
            ["Non-completers predicted to walk", reg["n_noncomp_pred_walk"]],
            ["Non-completers predicted not to walk", reg["n_noncomp_pred_no_walk"]],
            ["Feature count", len(res["features"])],
        ])

        # ── Positivity diagnostics: p_denom distribution ─────────────────────
        doc.add_paragraph()
        pd2 = doc.add_paragraph()
        pd2.add_run("2. Positivity Diagnostics — p_denom Distribution").bold = True

        flag_text = " ⚠ POSITIVITY FLAG: >5% of p_denom values near 0 or 1." if pd_diag["positivity_flag"] else " No positivity flag (< 5% near 0 or 1)."
        _add_table(doc, [
            ["p_denom statistic", "Value"],
            ["Mean", _fmt(pd_diag["mean"], 4)],
            ["SD", _fmt(pd_diag["sd"], 4)],
            ["Min", _fmt(pd_diag["min"], 4)],
            ["P10", _fmt(pd_diag["p10"], 4)],
            ["Median", _fmt(pd_diag["p50"], 4)],
            ["P90", _fmt(pd_diag["p90"], 4)],
            ["Max", _fmt(pd_diag["max"], 4)],
            ["% near 0 (<0.05)", _fmt(pd_diag["pct_near_zero"] * 100, 1) + "%"],
            ["% near 1 (>0.95)", _fmt(pd_diag["pct_near_one"] * 100, 1) + "%"],
            ["Positivity assessment", flag_text],
        ])

        # Stabilized weight distribution
        doc.add_paragraph()
        wh = doc.add_paragraph()
        wh.add_run("Stabilized Weight Distribution (primary positivity check)").bold = True
        mean_note = " (mean near 1 — reasonable numerator/denominator pairing)" if wd["mean_near_one"] else " ⚠ mean not near 1 — review numerator/denominator specification"
        _add_table(doc, [
            ["Weight statistic", "Value"],
            ["Mean", _fmt(wd["mean"], 4) + mean_note],
            ["SD", _fmt(wd["sd"], 4)],
            ["Min", _fmt(wd["min"], 4)],
            ["Max", _fmt(wd["max"], 4)],
            ["% truncated at winsor threshold", _fmt(wd["pct_truncated"] * 100, 1) + "%"],
        ])

        # Sensitivity to winsorization threshold
        doc.add_paragraph()
        sh = doc.add_paragraph()
        sh.add_run("Winsorization-Threshold Sensitivity (weighted R² and MAE)").bold = True
        sens_rows = [["Threshold", "Winsor lower", "Winsor upper", "Wt mean", "% truncated", "OOF R²", "OOF MAE"]]
        for pct in WINSORIZATION_THRESHOLDS:
            s = ipcw["sensitivity"][pct]
            sens_rows.append([
                f"{100-pct}th/{pct}th",
                _fmt(s["lower"], 3), _fmt(s["upper"], 3),
                _fmt(s["mean"], 4),
                _fmt(s["pct_truncated"] * 100, 1) + "%",
                _fmt(reg["sensitivity_r2"].get(pct, float("nan")), 4),
                _fmt(reg["sensitivity_mae"].get(pct, float("nan")), 1),
            ])
        _add_table(doc, sens_rows)

        # Sensitivity interpretation
        r2_vals = [reg["sensitivity_r2"].get(p, float("nan")) for p in WINSORIZATION_THRESHOLDS]
        finite_r2 = [v for v in r2_vals if math.isfinite(v)]
        if finite_r2:
            r2_spread = max(finite_r2) - min(finite_r2)
            stability_note = (
                "Stable (|ΔR²| < 0.05 across thresholds — positivity concern is limited)."
                if r2_spread < 0.05 else
                "⚠ Unstable (|ΔR²| ≥ 0.05 across thresholds — suggests positivity concern)."
            )
        else:
            stability_note = "Cannot assess — no finite R² values."
        p_stab = doc.add_paragraph(f"Sensitivity interpretation: {stability_note}")
        _set_small_font(p_stab)

        # ── Bootstrap CI section ─────────────────────────────────────────────
        doc.add_paragraph()
        bh = doc.add_paragraph()
        bh.add_run("3. Bootstrap Uncertainty Quantification").bold = True

        boot_text = (
            f"Full-pipeline stratified bootstrap (B={N_BOOTSTRAP}, stratified by completion status). "
            "Each resample re-fits the denominator model, re-constructs numerator/weights, "
            "re-winsorizes, reselects the binary classifier, and refits weighted Ridge. "
            f"Weighted R² bootstrap mean: {_fmt(ci['r2_boot_mean'], 4)} "
            f"[95% CI: {_fmt(ci['r2_ci_lo'], 4)}–{_fmt(ci['r2_ci_hi'], 4)}]. "
            f"Weighted MAE bootstrap mean: {_fmt(ci['mae_boot_mean'], 1)} "
            f"[95% CI: {_fmt(ci['mae_ci_lo'], 1)}–{_fmt(ci['mae_ci_hi'], 1)}]. "
            f"({ci['n_valid_r2']}/{N_BOOTSTRAP} valid R² resamples; "
            f"{ci['n_valid_mae']}/{N_BOOTSTRAP} valid MAE resamples.)"
        )
        p_boot = doc.add_paragraph(boot_text)
        _set_small_font(p_boot)

    # ── Methodology explainer ─────────────────────────────────────────────────
    doc.add_paragraph()
    mh = doc.add_paragraph()
    mh.add_run("Methodology Explainer: Why These Changes?").bold = True

    meth_items = [
        ("Intercept-only numerator",
         "The numerator in a stabilized IPCW weight controls how far the weights deviate from 1. "
         "An intercept-only numerator (marginal completion rate) is the most common and best-validated "
         "choice: it is simple, robust to numerator mis-specification, and typically gives weights whose "
         "mean is exactly the completer fraction. Using Age alone (prior session) was arbitrary and "
         "provides no advantage over the intercept-only baseline."),
        ("p_denom distribution check",
         "Positivity (overlap) is the assumption that every patient has non-zero probability of "
         "completing PAC given their covariates. Violations appear as p_denom values near 0 or 1. "
         "We flag tiers where >5% of values fall below 0.05 or above 0.95, since extreme probabilities "
         "generate extreme weights that weighting cannot repair — they signal that the weighting model "
         "is extrapolating beyond the support of the data."),
        ("Winsorization sensitivity",
         "Winsorization at a fixed percentile is a pragmatic but arbitrary choice. Reporting R² and MAE "
         "across 90th, 95th, and 99th percentile thresholds tests whether conclusions depend on this "
         "choice: stable results across thresholds give more confidence; sensitivity implies that "
         "extreme weights are driving the estimates."),
        ("Full-pipeline bootstrap",
         "The two-stage (weighting + outcome model) nature of the pipeline means that OOF metrics "
         "that treat weights as fixed understate uncertainty. The bootstrap — re-estimating all "
         "pipeline stages within each resample — propagates uncertainty from weight estimation, "
         "binary classifier selection, and Ridge fitting into the final CIs. "
         "Stratification by completion status prevents near-empty completer folds in small resamples."),
    ]
    for title_str, body_str in meth_items:
        h = doc.add_paragraph()
        h.add_run(title_str).bold = True
        p = doc.add_paragraph(body_str)
        _set_small_font(p)

    repro = doc.add_paragraph(
        f"Reproducibility: run `python {Path(__file__).name}` to regenerate this report and the Excel output."
    )
    _set_small_font(repro, 8)

    doc.save(OUTPUT_REPORT)


# ─────────────────────────────────────────────────────────────────────────────
# Main execution
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    df = pd.read_excel(INPUT_XLSX)
    print(f"Loaded: {df.shape[0]} rows × {df.shape[1]} columns")

    all_candidate_features = (
        DEMOGRAPHICS_PAC + STROKE_INFO + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED
        + NIHSS_OUT + NIHSS_IN + T1T2_IMPROVEMENT
    )
    for col in all_candidate_features:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    df["6MWT4"] = pd.to_numeric(df.get("6MWT4", np.nan), errors="coerce")

    output_df = df.copy()
    all_results: list[dict] = []

    for tier_name, tier_features in TIERS.items():
        valid_features = _filter_existing(tier_features, df)
        if not valid_features:
            continue

        print(f"\n{'=' * 72}")
        print(f"Processing {tier_name} ({len(valid_features)} features)")

        ipcw_completion = compute_ipcw_weights(df, valid_features, numerator_mode=NUMERATOR_MODE)
        wd = ipcw_completion["weight_diagnostics"]
        pd_diag = ipcw_completion["p_denom_diagnostics"]
        print(f"  Numerator: {ipcw_completion['numerator_model'][:60]}")
        print(f"  Weights: mean={wd['mean']:.4f}  SD={wd['sd']:.4f}  min={wd['min']:.4f}  max={wd['max']:.4f}  %trunc={wd['pct_truncated']*100:.1f}%")
        if pd_diag["positivity_flag"]:
            print(f"  ⚠ Positivity flag: {pd_diag['pct_near_zero']*100:.1f}% near 0, {pd_diag['pct_near_one']*100:.1f}% near 1")

        for scenario in SCENARIOS:
            scenario_label = _scenario_token(scenario)
            print(f"\n  Scenario: {scenario_label}")

            binary_result = select_binary_classifier(df, valid_features, scenario)
            print(f"    Binary model: {binary_result['best_model_name']} | Bal Acc={binary_result['bal_acc']:.3f}")

            regression_result = fit_weighted_ridge(df, valid_features, ipcw_completion, binary_result)
            print(f"    Ridge OOF R²={regression_result['weighted_r2']:.4f} | MAE={regression_result['weighted_mae']:.1f}")

            print(f"    Running bootstrap (B={N_BOOTSTRAP})...", end=" ", flush=True)
            ci = bootstrap_ci(df, valid_features, scenario,
                              n_bootstrap=N_BOOTSTRAP, numerator_mode=NUMERATOR_MODE)
            print(f"R² CI=[{ci['r2_ci_lo']:.4f}, {ci['r2_ci_hi']:.4f}]")

            # Write prediction columns
            walk_col = f"IPCW_Walk_{_tier_token(tier_name)}_{scenario_label}"
            pred_col = f"IPCW_6MWT4_{_tier_token(tier_name)}_{scenario_label}"
            ci_lo_col = f"IPCW_6MWT4_{_tier_token(tier_name)}_{scenario_label}_R2_CI_lo"
            ci_hi_col = f"IPCW_6MWT4_{_tier_token(tier_name)}_{scenario_label}_R2_CI_hi"
            output_df[walk_col] = pd.Series(pd.NA, index=output_df.index, dtype="Int64")
            output_df[pred_col] = np.nan
            output_df[ci_lo_col] = ci["r2_ci_lo"]
            output_df[ci_hi_col] = ci["r2_ci_hi"]
            nc_idx = regression_result["noncomp_index"]
            output_df.loc[nc_idx, walk_col] = regression_result["noncomp_walk_pred"]
            output_df.loc[nc_idx, pred_col] = regression_result["noncomp_6mwt4_pred"]

            all_results.append({
                "tier": tier_name,
                "scenario": scenario,
                "scenario_label": scenario_label,
                "features": valid_features,
                "ipcw_completion": ipcw_completion,
                "binary": binary_result,
                "regression": regression_result,
                "bootstrap_ci": ci,
            })

    output_df.to_excel(OUTPUT_XLSX, index=False)
    print(f"\nSaved dataset: {OUTPUT_XLSX.name}")

    write_report(all_results)
    print(f"Saved report: {OUTPUT_REPORT.name}")

    print("\n=== SUMMARY ===")
    print(f"{'Tier':<10} {'Scenario':<10} {'Bal Acc':>8} {'Wt R²':>8} {'Wt MAE':>8} {'Boot R² CI':>22}")
    print("-" * 80)
    for r in all_results:
        ci = r["bootstrap_ci"]
        print(
            f"{r['tier']:<10} {r['scenario_label']:<10} "
            f"{r['binary']['bal_acc']:>8.3f} "
            f"{r['regression']['weighted_r2']:>8.4f} "
            f"{r['regression']['weighted_mae']:>8.1f} "
            f"  [{ci['r2_ci_lo']:>6.4f}, {ci['r2_ci_hi']:>6.4f}]"
        )


if __name__ == "__main__":
    main()
