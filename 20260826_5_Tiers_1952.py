#!/usr/bin/env python3
"""
5-Tier IPCW analysis for 6MWT4 with best/worst scenario binary classification.

Steps:
1. Define variable categories: NIHSS_Out, NIHSS_In, T1T2_Improvement
2. Compute IPCW weights using 5 tiers of predictors
3. For each tier:
   a. Predict 6MWT4 (completers only, no missing handling)
   b. IPCW-extrapolated 6MWT4 for non-completers via best/worst binary models
4. Save results as 20260826_5_Tiers_1952.docx, code as 20260826_5_Tiers_Code_1952.docx,
   updated data as 20260826_DeID_1952.xlsx
"""

from __future__ import annotations

import textwrap
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from sklearn.base import clone
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import ExtraTreesClassifier, RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.inspection import permutation_importance
from sklearn.linear_model import LogisticRegression, Ridge
from sklearn.metrics import (
    accuracy_score,
    balanced_accuracy_score,
    f1_score,
    mean_absolute_error,
    precision_score,
    r2_score,
    recall_score,
)
from sklearn.model_selection import StratifiedKFold, cross_val_predict, cross_val_score, cross_validate
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler

# ─────────────────────────────────────────────────────────────────────────────
# Paths
# ─────────────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent
INPUT_XLSX = ROOT / "20260826_DeID.xlsx"
OUTPUT_XLSX = ROOT / "20260826_DeID_1952.xlsx"
OUTPUT_REPORT = ROOT / "20260826_5_Tiers_1952.docx"
OUTPUT_CODE = ROOT / "20260826_5_Tiers_Code_1952.docx"

RANDOM_STATE = 42

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1 – Variable category definitions
# ─────────────────────────────────────────────────────────────────────────────

# Previously defined groups (from Covariates.docx)
DEMOGRAPHICS_PAC = ["Age", "Sex, F0 M1"]
STROKE_INFO = [
    "Dissection", "ACA", "Undetermined", "HemorrhageStroke",
    "LVS", "LVO",
    "Side_Right", "Side_Left", "Side_Bilateral",
    "Loc_CortSub", "Loc_Subcortical", "Loc_Infratentorial",
]
COMORBIDITIES_PAC = [
    "AF", "DM", "HTN", "Dyslipidemia", "CAD", "CKD",
    "RestrictiveLung", "GIUlcer", "LiverCirrhosis", "Hepatitis",
    "Parkinsonism", "Malignancy", "OldStroke", "Dementia", "Psychiatric", "Gout",
]
ACUTE_COMPLICATIONS_PAC = ["Pneumonia", "UTI", "GIB", "Cellulitis"]
FUNCTIONAL_T1_PLUS_GS_IMPUTED = [
    "MRS1", "BI1", "FOIS1", "MNA1", "EuroQoL5D1", "IADL1",
    "BBS1", "Gait_Speed_1_Imputed", "FuglUE1", "FuglSEN1",
]

# New variable categories (Step 1)
NIHSS_OUT = [
    "ConsOut", "AnswerOut", "OrderOut", "EOMOut", "VisualOut",
    "FacialOut", "LUOut", "RUOut", "LLOut", "RLOut",
    "Coordinateout", "SensoryOut", "LanguageOut", "ArticulateOut", "NeglectOut",
]
NIHSS_IN = [
    "ConsIn", "AnswerIn", "OrderIn", "EOMIn", "VisualIn",
    "FaceIn", "LUIn", "RUIn", "LLIn", "RLIn",
    "CoordinateIn", "SensoryIn", "LanguageIn", "ArticulateIn", "NeglectIn",
]
T1T2_IMPROVEMENT = [
    "BI_T1T2_Change", "BBS_T1T2_Change", "MRS_T1T2_Change",
    "FOIS_T1T2_Change", "MNA_T1T2_Change", "IADL_T1T2_Change",
    "FuglUE_T1T2_Change", "FuglSEN_T1T2_Change", "EuroQoL5D_T1T2_Change",
]

# ─────────────────────────────────────────────────────────────────────────────
# STEP 2 – Tier definitions
# ─────────────────────────────────────────────────────────────────────────────

def _remove(lst: list[str], item: str) -> list[str]:
    return [x for x in lst if x != item]


TIERS: dict[str, list[str]] = {
    "Tier 1": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED,
    "Tier 2": DEMOGRAPHICS_PAC + _remove(FUNCTIONAL_T1_PLUS_GS_IMPUTED, "Gait_Speed_1_Imputed"),
    "Tier 3": (
        DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED
        + COMORBIDITIES_PAC + STROKE_INFO + ACUTE_COMPLICATIONS_PAC
    ),
    "Tier 4": (
        DEMOGRAPHICS_PAC + _remove(FUNCTIONAL_T1_PLUS_GS_IMPUTED, "Gait_Speed_1_Imputed")
        + COMORBIDITIES_PAC + STROKE_INFO + ACUTE_COMPLICATIONS_PAC
    ),
    "Tier 5": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + T1T2_IMPROVEMENT,
}

SCENARIOS = ["6MWT_Best_Scenario", "6MWT_Worst_Scenario"]

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _filter_existing(cols: list[str], df: pd.DataFrame) -> list[str]:
    return [c for c in cols if c in df.columns]


def _make_numeric_pipeline(features: list[str]) -> Pipeline:
    return Pipeline([
        ("prep", ColumnTransformer([
            ("num", Pipeline([
                ("imp", SimpleImputer(strategy="median")),
                ("sc", StandardScaler()),
            ]), features),
        ])),
    ])


def _make_cv() -> StratifiedKFold:
    return StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)


def _make_binary_candidates(features: list[str]) -> dict[str, dict]:
    numeric_pipe = Pipeline([
        ("imp", SimpleImputer(strategy="median")),
        ("sc", StandardScaler()),
    ])
    prep = ColumnTransformer([("num", numeric_pipe, features)])

    logistic = Pipeline([
        ("prep", prep),
        ("model", LogisticRegression(
            max_iter=5000, solver="liblinear",
            class_weight="balanced", random_state=RANDOM_STATE,
        )),
    ])
    rf = Pipeline([
        ("prep", ColumnTransformer([("num", Pipeline([
            ("imp", SimpleImputer(strategy="median")),
        ]), features)])),
        ("model", RandomForestClassifier(
            n_estimators=400, class_weight="balanced",
            random_state=RANDOM_STATE, n_jobs=-1,
        )),
    ])
    et = Pipeline([
        ("prep", ColumnTransformer([("num", Pipeline([
            ("imp", SimpleImputer(strategy="median")),
        ]), features)])),
        ("model", ExtraTreesClassifier(
            n_estimators=400, class_weight="balanced",
            random_state=RANDOM_STATE, n_jobs=-1,
        )),
    ])
    return {
        "LogisticRegression": {"pipeline": logistic, "code": _lr_code(features)},
        "RandomForest": {"pipeline": rf, "code": _rf_code(features)},
        "ExtraTrees": {"pipeline": et, "code": _et_code(features)},
    }


def _lr_code(features: list[str]) -> str:
    return (
        f"features = {features!r}\n"
        "Pipeline([\n"
        "    ('prep', ColumnTransformer([('num', Pipeline([('imp', SimpleImputer(strategy='median')),"
        " ('sc', StandardScaler())]), features)])),\n"
        "    ('model', LogisticRegression(max_iter=5000, solver='liblinear',"
        " class_weight='balanced', random_state=42)),\n"
        "])"
    )


def _rf_code(features: list[str]) -> str:
    return (
        f"features = {features!r}\n"
        "Pipeline([\n"
        "    ('prep', ColumnTransformer([('num', Pipeline([('imp', SimpleImputer(strategy='median'))]), features)])),\n"
        "    ('model', RandomForestClassifier(n_estimators=400, class_weight='balanced',"
        " random_state=42, n_jobs=-1)),\n"
        "])"
    )


def _et_code(features: list[str]) -> str:
    return (
        f"features = {features!r}\n"
        "Pipeline([\n"
        "    ('prep', ColumnTransformer([('num', Pipeline([('imp', SimpleImputer(strategy='median'))]), features)])),\n"
        "    ('model', ExtraTreesClassifier(n_estimators=400, class_weight='balanced',"
        " random_state=42, n_jobs=-1)),\n"
        "])"
    )


# ─────────────────────────────────────────────────────────────────────────────
# IPCW weight computation
# ─────────────────────────────────────────────────────────────────────────────

def compute_ipcw_weights(df: pd.DataFrame, features: list[str]) -> pd.Series:
    """Return stabilized IPCW weights (NaN for non-completers)."""
    valid = _filter_existing(features, df)
    comp = (df["PAC_Program_Completion"] == "Completed PAC program").astype(int)
    mask = comp.notna()
    X = df.loc[mask, valid]
    y = comp.loc[mask].to_numpy()

    imp = SimpleImputer(strategy="median")
    X_imp = pd.DataFrame(imp.fit_transform(X), columns=valid, index=X.index)

    denom_pipe = Pipeline([
        ("sc", StandardScaler()),
        ("lr", LogisticRegression(max_iter=5000, solver="lbfgs", C=1.0)),
    ])
    denom_pipe.fit(X_imp, y)
    p_denom = np.clip(denom_pipe.predict_proba(X_imp)[:, 1], 1e-4, 1 - 1e-4)

    numer_pipe = Pipeline([
        ("sc", StandardScaler()),
        ("lr", LogisticRegression(max_iter=1000, solver="lbfgs", C=1e6)),
    ])
    X_age = X_imp[["Age"]] if "Age" in X_imp.columns else X_imp.iloc[:, :1]
    numer_pipe.fit(X_age, y)
    p_numer = np.clip(numer_pipe.predict_proba(X_age)[:, 1], 1e-4, 1 - 1e-4)

    raw_w = np.where(y == 1, p_numer / p_denom, np.nan)

    comp_vals = raw_w[~np.isnan(raw_w)]
    lo, hi = np.nanpercentile(comp_vals, [1, 99])
    winsorized = np.where(~np.isnan(raw_w), np.clip(raw_w, lo, hi), np.nan)

    result = pd.Series(np.nan, index=df.index)
    result.loc[mask] = winsorized
    return result


# ─────────────────────────────────────────────────────────────────────────────
# Continuous 6MWT4 prediction (completers only, no IPCW)
# ─────────────────────────────────────────────────────────────────────────────

def predict_6mwt4_standard(df: pd.DataFrame, features: list[str]) -> dict:
    """Ridge regression on completers with observed 6MWT4."""
    valid = _filter_existing(features, df)
    comp_mask = (df["PAC_Program_Completion"] == "Completed PAC program") & df["6MWT4"].notna()
    df_comp = df.loc[comp_mask].copy()

    imp = SimpleImputer(strategy="median")
    X = pd.DataFrame(imp.fit_transform(df_comp[valid]), columns=valid)
    y = df_comp["6MWT4"].to_numpy()

    ridge = Ridge(alpha=1.0)
    ridge.fit(X, y)
    y_pred = ridge.predict(X)

    cv_r2 = cross_val_score(Ridge(alpha=1.0), X, y, cv=5, scoring="r2").mean()
    mae = mean_absolute_error(y, y_pred)
    r2 = r2_score(y, y_pred)

    importance_df = pd.DataFrame({
        "predictor": valid,
        "coefficient": ridge.coef_,
        "abs_coefficient": np.abs(ridge.coef_),
    }).sort_values("abs_coefficient", ascending=False).reset_index(drop=True)

    return {
        "n_total": int(len(df_comp)),
        "cv_r2": float(cv_r2),
        "r2": float(r2),
        "mae": float(mae),
        "model": ridge,
        "imputer": imp,
        "features": valid,
        "importance": importance_df,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Binary classification (best/worst scenario)
# ─────────────────────────────────────────────────────────────────────────────

def evaluate_binary(
    df: pd.DataFrame, features: list[str], scenario: str
) -> dict:
    """Select best binary model for given scenario and features."""
    valid = _filter_existing(features, df)
    model_df = df[df[scenario].notna()].copy()
    X = model_df[valid]
    y = model_df[scenario].astype(int)

    candidates = _make_binary_candidates(valid)
    cv = _make_cv()
    cv_splits = list(cv.split(X, y))

    leaderboard_rows = []
    for name, info in candidates.items():
        scores = cross_validate(
            clone(info["pipeline"]), X, y, cv=cv_splits, n_jobs=-1,
            scoring=["accuracy", "balanced_accuracy", "precision", "recall", "f1"],
        )
        leaderboard_rows.append({
            "model": name,
            "accuracy": float(np.mean(scores["test_accuracy"])),
            "balanced_accuracy": float(np.mean(scores["test_balanced_accuracy"])),
            "precision": float(np.mean(scores["test_precision"])),
            "recall": float(np.mean(scores["test_recall"])),
            "f1": float(np.mean(scores["test_f1"])),
        })

    leaderboard = (
        pd.DataFrame(leaderboard_rows)
        .sort_values(["balanced_accuracy", "accuracy", "f1"], ascending=False)
        .reset_index(drop=True)
    )
    best_name = str(leaderboard.iloc[0]["model"])
    best_pipeline = clone(candidates[best_name]["pipeline"])

    oof_preds = cross_val_predict(clone(best_pipeline), X, y, cv=cv_splits, n_jobs=-1)
    metrics = {
        "accuracy": float(accuracy_score(y, oof_preds)),
        "balanced_accuracy": float(balanced_accuracy_score(y, oof_preds)),
        "precision": float(precision_score(y, oof_preds, zero_division=0)),
        "recall": float(recall_score(y, oof_preds, zero_division=0)),
        "f1": float(f1_score(y, oof_preds, zero_division=0)),
    }
    n_pos = int(y.sum())
    n_neg = int((y == 0).sum())

    # Fit best pipeline on all data for prediction
    best_pipeline.fit(X, y)

    # Permutation importance (5-fold)
    fold_importances = []
    for fold_idx, (train_idx, test_idx) in enumerate(cv_splits):
        fp = clone(candidates[best_name]["pipeline"])
        fp.fit(X.iloc[train_idx], y.iloc[train_idx])
        fi = permutation_importance(
            fp, X.iloc[test_idx], y.iloc[test_idx],
            scoring="balanced_accuracy", n_repeats=10,
            random_state=RANDOM_STATE + fold_idx, n_jobs=-1,
        )
        fold_importances.append(fi.importances_mean)

    imp_arr = np.vstack(fold_importances)
    importance_df = (
        pd.DataFrame({
            "predictor": valid,
            "importance_mean": imp_arr.mean(axis=0),
            "importance_std": imp_arr.std(axis=0, ddof=1),
        })
        .sort_values("importance_mean", ascending=False)
        .reset_index(drop=True)
    )

    return {
        "scenario": scenario,
        "n_total": int(len(model_df)),
        "n_positive": n_pos,
        "n_negative": n_neg,
        "leaderboard": leaderboard,
        "best_model_name": best_name,
        "best_model_code": candidates[best_name]["code"],
        "best_pipeline": best_pipeline,
        "features": valid,
        "metrics": metrics,
        "importance": importance_df,
    }


# ─────────────────────────────────────────────────────────────────────────────
# IPCW-extrapolated 6MWT4 for non-completers
# ─────────────────────────────────────────────────────────────────────────────

def predict_6mwt4_ipcw(
    df: pd.DataFrame,
    features: list[str],
    ipcw_weights: pd.Series,
    binary_result: dict,
) -> dict:
    """
    Weighted Ridge regression on completers. Predict walking ability for
    non-completers using best binary model; extrapolate continuous 6MWT4
    for predicted walkers.
    """
    valid = _filter_existing(features, df)

    # Completers with observed 6MWT4
    comp_mask = (df["PAC_Program_Completion"] == "Completed PAC program") & df["6MWT4"].notna()
    df_comp = df.loc[comp_mask].copy()
    w = ipcw_weights.loc[comp_mask].fillna(1.0).to_numpy()

    imp = SimpleImputer(strategy="median")
    X_comp = pd.DataFrame(imp.fit_transform(df_comp[valid]), columns=valid)
    y_comp = df_comp["6MWT4"].to_numpy()

    ridge_w = Ridge(alpha=1.0)
    ridge_w.fit(X_comp, y_comp, sample_weight=w)
    y_pred_comp = ridge_w.predict(X_comp)

    cv_r2 = cross_val_score(Ridge(alpha=1.0), X_comp, y_comp, cv=5, scoring="r2").mean()
    mae_comp = float(np.average(np.abs(y_comp - y_pred_comp), weights=w))
    r2_comp = r2_score(y_comp, y_pred_comp, sample_weight=w)

    importance_df = pd.DataFrame({
        "predictor": valid,
        "coefficient": ridge_w.coef_,
        "abs_coefficient": np.abs(ridge_w.coef_),
    }).sort_values("abs_coefficient", ascending=False).reset_index(drop=True)

    # Non-completers: predict walking via binary model
    noncomp_mask = df["PAC_Program_Completion"] == "Did not complete PAC program"
    df_noncomp = df.loc[noncomp_mask].copy()
    X_noncomp = pd.DataFrame(imp.transform(df_noncomp[valid]), columns=valid)

    binary_pipe = binary_result["best_pipeline"]
    walk_pred = binary_pipe.predict(X_noncomp)
    n_pred_walk = int(walk_pred.sum())
    n_pred_no_walk = int((walk_pred == 0).sum())

    # Predict 6MWT4 for those expected to walk
    walk_idx = np.where(walk_pred == 1)[0]
    predicted_6mwt4 = np.zeros(len(df_noncomp))
    if len(walk_idx) > 0:
        X_walkers = X_noncomp.iloc[walk_idx]
        predicted_6mwt4[walk_idx] = np.maximum(0, ridge_w.predict(X_walkers))

    return {
        "n_completers": int(comp_mask.sum()),
        "n_noncompleters": int(noncomp_mask.sum()),
        "n_noncomp_pred_walk": n_pred_walk,
        "n_noncomp_pred_no_walk": n_pred_no_walk,
        "cv_r2": float(cv_r2),
        "r2_weighted": float(r2_comp),
        "mae_weighted": float(mae_comp),
        "model": ridge_w,
        "imputer": imp,
        "features": valid,
        "importance": importance_df,
        "noncomp_walk_pred": walk_pred,
        "noncomp_6mwt4_pred": predicted_6mwt4,
        "noncomp_index": df_noncomp.index.tolist(),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Report writing helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_table(doc: Document, rows: list[list]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(cell)
            if i == 0:
                for run in table.cell(i, j).paragraphs[0].runs:
                    run.bold = True


def _set_code_font(para, size: int = 8) -> None:
    for run in para.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(size)


def write_report(all_results: list[dict]) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("20260826 5-Tier IPCW 6MWT4 Analysis").bold = True

    note_text = (
        "Variable categories — NIHSS_Out, NIHSS_In, T1T2_Improvement — were added in Step 1. "
        "Five predictor tiers were used for IPCW censoring weighting and outcome regression. "
        "Best/worst scenario binary models were applied to predict walking ability for non-completers."
    )
    note = doc.add_paragraph(note_text)
    for run in note.runs:
        run.font.size = Pt(9)

    for result in all_results:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(f"Tier: {result['tier']}  |  Scenario: {result['scenario']}").bold = True

        # Tier predictor list
        tier_para = doc.add_paragraph(f"Predictors ({len(result['features'])}): {', '.join(result['features'][:10])}{'...' if len(result['features']) > 10 else ''}")
        for run in tier_para.runs:
            run.font.size = Pt(8)

        # Patient counts
        _add_table(doc, [
            ["Metric", "Value"],
            ["Total patients (completers, observed 6MWT4)", result["standard"]["n_total"]],
            ["Patients used in IPCW regression (completers)", result["ipcw"]["n_completers"]],
            ["Non-completers (missing 6MWT4)", result["ipcw"]["n_noncompleters"]],
            ["Non-completers predicted to walk (binary model)", result["ipcw"]["n_noncomp_pred_walk"]],
            ["Non-completers predicted not to walk", result["ipcw"]["n_noncomp_pred_no_walk"]],
            ["Binary model training N (positive / negative)",
             f'{result["binary"]["n_positive"]} / {result["binary"]["n_negative"]}'],
            ["Best binary model", result["binary"]["best_model_name"]],
        ])

        # Standard regression
        doc.add_paragraph()
        doc.add_paragraph().add_run("A. Standard prediction (completers, no missing handling)").bold = True
        _add_table(doc, [
            ["Metric", "Value"],
            ["CV R²", f'{result["standard"]["cv_r2"]:.4f}'],
            ["Train R²", f'{result["standard"]["r2"]:.4f}'],
            ["MAE", f'{result["standard"]["mae"]:.1f}'],
        ])
        imp_std = result["standard"]["importance"].head(10)
        _add_table(doc, [["Predictor", "|Coefficient|"]] + [
            [r["predictor"], f'{r["abs_coefficient"]:.4f}']
            for _, r in imp_std.iterrows()
        ])

        # IPCW regression
        doc.add_paragraph()
        doc.add_paragraph().add_run("B. IPCW-weighted prediction (completers, weighted)").bold = True
        _add_table(doc, [
            ["Metric", "Value"],
            ["CV R²", f'{result["ipcw"]["cv_r2"]:.4f}'],
            ["Weighted Train R²", f'{result["ipcw"]["r2_weighted"]:.4f}'],
            ["Weighted MAE", f'{result["ipcw"]["mae_weighted"]:.1f}'],
        ])
        imp_ipcw = result["ipcw"]["importance"].head(10)
        _add_table(doc, [["Predictor", "|Coefficient|"]] + [
            [r["predictor"], f'{r["abs_coefficient"]:.4f}']
            for _, r in imp_ipcw.iterrows()
        ])

        # Binary model performance
        doc.add_paragraph()
        doc.add_paragraph().add_run("C. Binary walking-ability model (OOF cross-validated)").bold = True
        bm = result["binary"]["metrics"]
        _add_table(doc, [
            ["Metric", "Value"],
            ["Accuracy", f'{bm["accuracy"]:.3f}'],
            ["Balanced Accuracy", f'{bm["balanced_accuracy"]:.3f}'],
            ["Precision", f'{bm["precision"]:.3f}'],
            ["Recall", f'{bm["recall"]:.3f}'],
            ["F1", f'{bm["f1"]:.3f}'],
        ])
        lb = result["binary"]["leaderboard"]
        lb_rows = [["Model", "Accuracy", "Bal.Acc", "Precision", "Recall", "F1"]]
        for _, row in lb.iterrows():
            lb_rows.append([
                row["model"],
                f'{row["accuracy"]:.3f}',
                f'{row["balanced_accuracy"]:.3f}',
                f'{row["precision"]:.3f}',
                f'{row["recall"]:.3f}',
                f'{row["f1"]:.3f}',
            ])
        _add_table(doc, lb_rows)

        doc.add_paragraph()
        doc.add_paragraph().add_run("D. Binary model permutation importance (top 10)").bold = True
        imp_bin = result["binary"]["importance"].head(10)
        _add_table(doc, [["Predictor", "Importance Mean", "Importance SD"]] + [
            [r["predictor"], f'{r["importance_mean"]:.4f}', f'{r["importance_std"]:.4f}']
            for _, r in imp_bin.iterrows()
        ])

    doc.save(OUTPUT_REPORT)


def write_code_doc(all_results: list[dict]) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("20260826 5-Tier IPCW and Binary Model Code").bold = True

    # Variable category definitions
    doc.add_paragraph().add_run("Variable Category Definitions").bold = True
    categories_code = textwrap.dedent(f"""\
        NIHSS_OUT = {NIHSS_OUT!r}

        NIHSS_IN = {NIHSS_IN!r}

        T1T2_IMPROVEMENT = {T1T2_IMPROVEMENT!r}

        TIERS = {{
            "Tier 1": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED,
            "Tier 2": DEMOGRAPHICS_PAC + [x for x in FUNCTIONAL_T1_PLUS_GS_IMPUTED if x != "Gait_Speed_1_Imputed"],
            "Tier 3": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + COMORBIDITIES_PAC + STROKE_INFO + ACUTE_COMPLICATIONS_PAC,
            "Tier 4": DEMOGRAPHICS_PAC + [x for x in FUNCTIONAL_T1_PLUS_GS_IMPUTED if x != "Gait_Speed_1_Imputed"] + COMORBIDITIES_PAC + STROKE_INFO + ACUTE_COMPLICATIONS_PAC,
            "Tier 5": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + T1T2_IMPROVEMENT,
        }}
    """)
    p = doc.add_paragraph(categories_code)
    _set_code_font(p)

    # IPCW code
    doc.add_paragraph().add_run("IPCW Weight Computation").bold = True
    ipcw_code = textwrap.dedent("""\
        # Denominator model: P(complete | X)
        denom_pipe = Pipeline([
            ("sc", StandardScaler()),
            ("lr", LogisticRegression(max_iter=5000, solver="lbfgs", C=1.0)),
        ])
        denom_pipe.fit(X_imp, y_comp)
        p_denom = np.clip(denom_pipe.predict_proba(X_imp)[:, 1], 1e-4, 1 - 1e-4)

        # Numerator model: P(complete | Age)
        numer_pipe = Pipeline([
            ("sc", StandardScaler()),
            ("lr", LogisticRegression(max_iter=1000, solver="lbfgs", C=1e6)),
        ])
        numer_pipe.fit(X_imp[["Age"]], y_comp)
        p_numer = np.clip(numer_pipe.predict_proba(X_imp[["Age"]])[:, 1], 1e-4, 1 - 1e-4)

        # Stabilized weights (completers only), winsorized at 1st/99th percentile
        raw_w = np.where(y_comp == 1, p_numer / p_denom, np.nan)
        lo, hi = np.nanpercentile(raw_w[~np.isnan(raw_w)], [1, 99])
        ipcw = np.where(~np.isnan(raw_w), np.clip(raw_w, lo, hi), np.nan)
    """)
    p = doc.add_paragraph(ipcw_code)
    _set_code_font(p)

    # Per-tier/scenario best model code
    seen = set()
    for result in all_results:
        key = (result["tier"], result["scenario"])
        if key in seen:
            continue
        seen.add(key)
        doc.add_paragraph()
        doc.add_paragraph().add_run(
            f"Best binary model for {result['tier']} / {result['scenario']}: "
            f"{result['binary']['best_model_name']}"
        ).bold = True
        p = doc.add_paragraph(result["binary"]["best_model_code"])
        _set_code_font(p)

    doc.save(OUTPUT_CODE)


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    df = pd.read_excel(INPUT_XLSX)
    print(f"Loaded: {df.shape[0]} rows × {df.shape[1]} columns")

    # Coerce numeric columns
    all_potential = (
        DEMOGRAPHICS_PAC + STROKE_INFO + COMORBIDITIES_PAC + ACUTE_COMPLICATIONS_PAC
        + FUNCTIONAL_T1_PLUS_GS_IMPUTED + NIHSS_OUT + NIHSS_IN + T1T2_IMPROVEMENT
    )
    for c in all_potential:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")
    df["6MWT4"] = pd.to_numeric(df.get("6MWT4", np.nan), errors="coerce")

    # ── Build new column ordering for output Excel ────────────────────────────
    category_cols = [
        c for c in (
            DEMOGRAPHICS_PAC + STROKE_INFO + COMORBIDITIES_PAC + ACUTE_COMPLICATIONS_PAC
            + FUNCTIONAL_T1_PLUS_GS_IMPUTED
            + NIHSS_OUT + NIHSS_IN + T1T2_IMPROVEMENT
        )
        if c in df.columns
    ]
    remaining = [c for c in df.columns if c not in category_cols]
    df_out = df[category_cols + remaining].copy()

    # ── Run tiers × scenarios ─────────────────────────────────────────────────
    all_results = []
    for tier_name, tier_features in TIERS.items():
        valid_feats = _filter_existing(tier_features, df)
        print(f"\n{'='*60}")
        print(f"Computing IPCW for {tier_name} ({len(valid_feats)} predictors)...")
        ipcw_weights = compute_ipcw_weights(df, valid_feats)
        w_mean = ipcw_weights.dropna().mean()
        print(f"  IPCW weight mean: {w_mean:.4f}")

        # Standard (no missing) prediction
        std_result = predict_6mwt4_standard(df, valid_feats)
        print(f"  Standard CV R²: {std_result['cv_r2']:.4f}  MAE: {std_result['mae']:.1f}")

        for scenario in SCENARIOS:
            print(f"  Scenario: {scenario}")
            binary_result = evaluate_binary(df, valid_feats, scenario)
            print(f"    Binary best: {binary_result['best_model_name']}  "
                  f"BalAcc: {binary_result['metrics']['balanced_accuracy']:.3f}")

            ipcw_result = predict_6mwt4_ipcw(df, valid_feats, ipcw_weights, binary_result)
            print(f"    IPCW CV R²: {ipcw_result['cv_r2']:.4f}  "
                  f"Pred walkers (noncomp): {ipcw_result['n_noncomp_pred_walk']}")

            all_results.append({
                "tier": tier_name,
                "scenario": scenario,
                "features": valid_feats,
                "standard": std_result,
                "ipcw": ipcw_result,
                "binary": binary_result,
            })

            # Add predictions to output dataframe
            pred_col = f"IPCW_6MWT4_{tier_name.replace(' ', '_')}_{scenario.replace('6MWT_', '').replace('_Scenario', '')}"
            df_out[pred_col] = np.nan
            noncomp_idx = ipcw_result["noncomp_index"]
            df_out.loc[noncomp_idx, pred_col] = ipcw_result["noncomp_6mwt4_pred"]

    # ── Save outputs ──────────────────────────────────────────────────────────
    df_out.to_excel(OUTPUT_XLSX, index=False)
    print(f"\nSaved dataset: {OUTPUT_XLSX.name}")

    write_report(all_results)
    print(f"Saved report: {OUTPUT_REPORT.name}")

    write_code_doc(all_results)
    print(f"Saved code doc: {OUTPUT_CODE.name}")

    # ── Print summary ─────────────────────────────────────────────────────────
    print("\n=== SUMMARY ===")
    print(f"{'Tier':<10} {'Scenario':<25} {'Std CV R²':>10} {'IPCW CV R²':>12} {'Binary BalAcc':>14} {'Pred Walk':>10}")
    print("-" * 80)
    for r in all_results:
        print(
            f"{r['tier']:<10} {r['scenario']:<25} "
            f"{r['standard']['cv_r2']:>10.4f} {r['ipcw']['cv_r2']:>12.4f} "
            f"{r['binary']['metrics']['balanced_accuracy']:>14.3f} "
            f"{r['ipcw']['n_noncomp_pred_walk']:>10}"
        )


if __name__ == "__main__":
    main()
