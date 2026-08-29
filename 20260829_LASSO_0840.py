#!/usr/bin/env python3
"""
20260829_LASSO_0840.py
======================
Bootstrap LASSO regression models for 6MWT4 prediction.
NO IPCW is used anywhere in this workflow.

Four models are built, each using a different predictor combination
drawn from the Step 1 variable categories defined in
20260826_5_Tiers_1952.py:

    Model 1: demographic + functional T1
    Model 2: demographic + stroke info + comorbidities + acute complications + NIHSS out
    Model 3: demographic + BBS1 + gait speed 1 imputed
    Model 4: demographics + functional T1 + T1T2 improvement

Outputs (written to repo root):
    20260829_LASSO_0840.docx  – summary report per model
    20260829_LASSO_0840.xlsx  – full dataset with per-patient LASSO predictions appended

Author: auto-generated 2026-08-29
"""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt

from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LassoCV
from sklearn.metrics import mean_absolute_error, r2_score
from sklearn.model_selection import KFold, cross_val_predict
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler

# ─────────────────────────────────────────────────────────────────────────────
# FILE PATHS
# ─────────────────────────────────────────────────────────────────────────────
# Resolve paths relative to this script file so the script can be run from
# any working directory.
ROOT = Path(__file__).resolve().parent
INPUT_XLSX = ROOT / "20260826_DeID.xlsx"
OUTPUT_DOCX = ROOT / "20260829_LASSO_0840.docx"
OUTPUT_XLSX = ROOT / "20260829_LASSO_0840.xlsx"

# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL SETTINGS
# ─────────────────────────────────────────────────────────────────────────────
RANDOM_STATE = 42          # reproducibility seed for all stochastic steps
N_BOOTSTRAP = 1000         # number of bootstrap resamples per model
CV_FOLDS = 5               # folds for LassoCV inner tuning and outer OOF eval
# Regularisation search grid: 120 alpha values log-spaced from 1e-4 to 100
ALPHAS = np.logspace(-4, 2, 120)
# Stability threshold: proportion of bootstrap samples in which a predictor
# must have a non-zero LASSO coefficient to be considered "stably selected"
STABILITY_THRESHOLD = 0.70

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1 VARIABLE CATEGORIES (consistent with 20260826_5_Tiers_1952.py)
# ─────────────────────────────────────────────────────────────────────────────
# Each list below mirrors the exact same category lists used in the prior
# Tier-based analysis so that cross-script comparisons remain valid.

DEMOGRAPHICS_PAC = ["Age", "Sex, F0 M1"]

STROKE_INFO = [
    "Dissection", "ACA", "Undetermined", "HemorrhageStroke",
    "LVS", "LVO",
    "Side_Right", "Side_Left", "Side_Bilateral",
    "Loc_CortSub", "Loc_Subcortical", "Loc_Infratentorial",
]

COMORBIDITIES_PAC = [
    "AF", "DM", "HTN", "Dyslipidemia", "CAD", "CKD",
    "RestrictiveLung", "GIUlcer", "LiverCirrhosis", "Hepatitis",
    "Parkinsonism", "Malignancy", "OldStroke", "Dementia", "Psychiatric", "Gout",
]

ACUTE_COMPLICATIONS_PAC = ["Pneumonia", "UTI", "GIB", "Cellulitis"]

FUNCTIONAL_T1_PLUS_GS_IMPUTED = [
    "MRS1", "BI1", "FOIS1", "MNA1", "EuroQoL5D1", "IADL1",
    "BBS1", "Gait_Speed_1_Imputed", "FuglUE1", "FuglSEN1",
]

NIHSS_OUT = [
    "ConsOut", "AnswerOut", "OrderOut", "EOMOut", "VisualOut",
    "FacialOut", "LUOut", "RUOut", "LLOut", "RLOut",
    "Coordinateout", "SensoryOut", "LanguageOut", "ArticulateOut", "NeglectOut",
]

T1T2_IMPROVEMENT = [
    "BI_T1T2_Change", "BBS_T1T2_Change", "MRS_T1T2_Change",
    "FOIS_T1T2_Change", "MNA_T1T2_Change", "IADL_T1T2_Change",
    "FuglUE_T1T2_Change", "FuglSEN_T1T2_Change", "EuroQoL5D_T1T2_Change",
]

# ─────────────────────────────────────────────────────────────────────────────
# MODEL DEFINITIONS
# ─────────────────────────────────────────────────────────────────────────────
# Each entry: human-readable model name → flat list of candidate predictor
# column names.  Columns absent from the actual data are silently dropped
# later (see _filter_existing).
MODEL_SPECS: dict[str, list[str]] = {
    "Model 1: demographic + functional T1": (
        DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED
    ),
    "Model 2: demographic + stroke info + comorbidities + acute complications + NIHSS out": (
        DEMOGRAPHICS_PAC
        + STROKE_INFO
        + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC
        + NIHSS_OUT
    ),
    "Model 3: demographic + BBS1 + gait speed 1 imputed": (
        DEMOGRAPHICS_PAC + ["BBS1", "Gait_Speed_1_Imputed"]
    ),
    "Model 4: demographics + functional T1 + T1T2 improvement": (
        DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + T1T2_IMPROVEMENT
    ),
}

# Short column names used in the Excel output for each model
MODEL_PRED_COLS = {
    "Model 1: demographic + functional T1": "LASSO_Pred_Model1",
    "Model 2: demographic + stroke info + comorbidities + acute complications + NIHSS out": "LASSO_Pred_Model2",
    "Model 3: demographic + BBS1 + gait speed 1 imputed": "LASSO_Pred_Model3",
    "Model 4: demographics + functional T1 + T1T2 improvement": "LASSO_Pred_Model4",
}

# ─────────────────────────────────────────────────────────────────────────────
# UTILITY FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def _filter_existing(cols: list[str], df: pd.DataFrame) -> list[str]:
    """Return only those columns that actually exist in df, preserving order."""
    return [c for c in cols if c in df.columns]


def _build_lasso_pipeline(features: list[str]) -> Pipeline:
    """
    Build a fresh sklearn Pipeline for LASSO:

    1. ColumnTransformer  – selects the named feature columns
    2. SimpleImputer      – median imputation for missing values
    3. StandardScaler     – z-score standardisation (mean=0, sd=1)
    4. LassoCV            – regularised linear regression with inner
                            cross-validation to choose the best alpha

    A fresh pipeline is built each call so bootstrap iterations get
    independent estimator objects.
    """
    numeric_transformer = Pipeline([
        ("imputer", SimpleImputer(strategy="median")),
        ("scaler", StandardScaler()),
    ])
    preprocessor = ColumnTransformer([
        ("num", numeric_transformer, features)
    ])
    lasso = LassoCV(
        alphas=ALPHAS,
        cv=CV_FOLDS,
        random_state=RANDOM_STATE,
        max_iter=50_000,   # high iteration limit avoids convergence warnings
        n_jobs=-1,
    )
    return Pipeline([
        ("prep", preprocessor),
        ("model", lasso),
    ])


def _add_table_to_doc(doc: Document, rows: list[list]) -> None:
    """Insert a Word table with a header row (bold) into doc."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def _small_para(paragraph, size: int = 9) -> None:
    """Set all runs in a paragraph to a small font size."""
    for run in paragraph.runs:
        run.font.size = Pt(size)


# ─────────────────────────────────────────────────────────────────────────────
# CORE MODELLING FUNCTION
# ─────────────────────────────────────────────────────────────────────────────

def fit_bootstrap_lasso(
    df: pd.DataFrame,
    features: list[str],
    model_name: str,
    model_seed: int = RANDOM_STATE,
) -> dict:
    """
    Fit a true bootstrap LASSO model and return a results dictionary.

    Steps
    -----
    1. Restrict analysis set to rows where 6MWT4 is observed.
    2. Build a baseline LASSO (full data fit) to get final coefficients and
       apparent (train-set) performance.
    3. Run outer K-fold cross-validation on the full data to get unbiased
       out-of-fold (OOF) R² and MAE.
    4. Run N_BOOTSTRAP bootstrap resamples:
       - Sample N rows with replacement.
       - Fit a new LASSO pipeline on the bootstrap sample.
       - Record the coefficient vector.
    5. From bootstrap coefficients compute:
       - selection frequency (proportion of resamples with non-zero coef)
       - mean and SD of coefficient across resamples
    6. Return all metrics + importance table.

    Parameters
    ----------
    df          : full dataset (all rows, predictor columns + '6MWT4')
    features    : column names to use as predictors (already filtered to
                  columns that exist in df)
    model_name  : human-readable label for logging / report

    Returns
    -------
    dict with keys described in the body below
    """
    # ── 1. Analysis set: only rows with observed 6MWT4 ──────────────────────
    model_df = df[df["6MWT4"].notna()].copy()
    X = model_df[features]
    y = model_df["6MWT4"].to_numpy(dtype=float)

    n_patients = len(model_df)
    n_input = len(features)

    # ── 2. Baseline (full-data) LASSO fit ────────────────────────────────────
    base_pipe = _build_lasso_pipeline(features)
    base_pipe.fit(X, y)
    base_lasso: LassoCV = base_pipe.named_steps["model"]
    base_coef = base_lasso.coef_.copy()    # shape (n_features,)
    best_alpha = float(base_lasso.alpha_)

    # Apparent (train-set) predictions and performance
    y_train_pred = base_pipe.predict(X)
    train_r2 = float(r2_score(y, y_train_pred))
    train_mae = float(mean_absolute_error(y, y_train_pred))

    # ── 3. Outer CV – out-of-fold R² and MAE ─────────────────────────────────
    # Build a fresh pipeline for the OOF evaluation so the inner LassoCV
    # tuning happens independently inside each fold.
    outer_cv = KFold(n_splits=CV_FOLDS, shuffle=True, random_state=RANDOM_STATE)
    y_oof = cross_val_predict(
        _build_lasso_pipeline(features), X, y,
        cv=outer_cv, n_jobs=-1,
    )
    cv_r2 = float(r2_score(y, y_oof))
    cv_mae = float(mean_absolute_error(y, y_oof))

    # ── 4. Bootstrap resampling ───────────────────────────────────────────────
    # We store one coefficient vector per bootstrap sample.
    # model_seed is derived per model (see main()) so that bootstrap samples
    # are statistically independent across models.
    rng = np.random.default_rng(model_seed)
    coef_boot = np.zeros((N_BOOTSTRAP, n_input), dtype=float)
    alpha_boot = np.zeros(N_BOOTSTRAP, dtype=float)

    for b in range(N_BOOTSTRAP):
        # Draw a bootstrap sample (sample with replacement, same N as training)
        idx = rng.integers(0, n_patients, size=n_patients)
        Xb = X.iloc[idx]
        yb = y[idx]

        bp = _build_lasso_pipeline(features)
        bp.fit(Xb, yb)
        lb: LassoCV = bp.named_steps["model"]

        coef_boot[b, :] = lb.coef_
        alpha_boot[b] = lb.alpha_

    # ── 5. Bootstrap summary statistics ──────────────────────────────────────
    # Use a small tolerance rather than exact == 0 to avoid floating-point
    # artefacts in LASSO coefficient comparisons.
    _COEF_TOL = 1e-9
    sel_freq = (np.abs(coef_boot) > _COEF_TOL).mean(axis=0)
    coef_mean = coef_boot.mean(axis=0)
    coef_sd = coef_boot.std(axis=0, ddof=1)

    # ── 6. Importance table (sorted by selection frequency then |mean coef|) ─
    importance = pd.DataFrame({
        "predictor": features,
        "base_coef": base_coef,
        "bootstrap_coef_mean": coef_mean,
        "bootstrap_coef_sd": coef_sd,
        "selection_frequency": sel_freq,
        "abs_bootstrap_coef_mean": np.abs(coef_mean),
    }).sort_values(
        ["selection_frequency", "abs_bootstrap_coef_mean"],
        ascending=[False, False],
    ).reset_index(drop=True)

    # Count how many predictors are non-zero in the full-data fit (with tolerance)
    n_final_nonzero = int((np.abs(base_coef) > 1e-9).sum())
    # Count stably selected predictors (above stability threshold)
    n_stable = int((sel_freq >= STABILITY_THRESHOLD).sum())

    return {
        "model_name": model_name,
        "features": features,
        "n_patients": n_patients,
        "n_input_variables": n_input,
        "n_final_nonzero_base": n_final_nonzero,
        "n_stable_selected": n_stable,
        "best_alpha_full_fit": best_alpha,
        "alpha_boot_median": float(np.median(alpha_boot)),
        "alpha_boot_q25": float(np.quantile(alpha_boot, 0.25)),
        "alpha_boot_q75": float(np.quantile(alpha_boot, 0.75)),
        "cv_r2": cv_r2,
        "cv_mae": cv_mae,
        "train_r2": train_r2,
        "train_mae": train_mae,
        "base_pipe": base_pipe,          # fitted pipeline for prediction
        "importance": importance,
    }


# ─────────────────────────────────────────────────────────────────────────────
# REPORT GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def write_docx_report(results: list[dict]) -> None:
    """Write a Word document summarising all 4 bootstrap LASSO models."""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.add_run(
        "Bootstrap LASSO Models for 6MWT4 Prediction (No IPCW) – 2026-08-29"
    ).bold = True

    # Introduction / methodology note
    intro = doc.add_paragraph(
        f"This report summarises {len(results)} bootstrap LASSO regression models "
        "predicting the 6-Minute Walk Test at time-point 4 (6MWT4) as a continuous outcome. "
        "No IPCW weighting is applied. "
        "Missing predictor values are handled by median imputation; predictors are "
        "z-score standardised prior to LASSO. "
        f"Regularisation parameter alpha is chosen by {CV_FOLDS}-fold cross-validation "
        f"(LassoCV). True bootstrap resampling ({N_BOOTSTRAP:,} resamples) is used to "
        "assess coefficient stability. Performance is reported both on the full training "
        f"set (apparent) and via {CV_FOLDS}-fold out-of-fold (OOF) cross-validation."
    )
    _small_para(intro, 9)

    for r in results:
        doc.add_paragraph()
        # Model heading
        heading = doc.add_paragraph()
        heading.add_run(r["model_name"]).bold = True

        # Predictor list (truncated for brevity)
        feat_preview = ", ".join(r["features"][:15])
        if len(r["features"]) > 15:
            feat_preview += f" … (+{len(r['features']) - 15} more)"
        pred_para = doc.add_paragraph(f"Predictors ({len(r['features'])}): {feat_preview}")
        _small_para(pred_para, 8)

        # Summary metrics table
        _add_table_to_doc(doc, [
            ["Metric", "Value"],
            ["Patients (non-missing 6MWT4)", r["n_patients"]],
            ["Input variable count", r["n_input_variables"]],
            ["Final selected (non-zero in full-fit)", r["n_final_nonzero_base"]],
            [
                f"Stably selected (selection freq ≥ {STABILITY_THRESHOLD:.0%})",
                r["n_stable_selected"],
            ],
            ["Best alpha (full-data fit)", f"{r['best_alpha_full_fit']:.6g}"],
            [
                "Bootstrap alpha: median [IQR]",
                (
                    f"{r['alpha_boot_median']:.6g}"
                    f" [{r['alpha_boot_q25']:.6g}, {r['alpha_boot_q75']:.6g}]"
                ),
            ],
            ["CV R² (OOF)", f"{r['cv_r2']:.4f}"],
            ["CV MAE (OOF)", f"{r['cv_mae']:.2f} m"],
            ["Train R² (apparent)", f"{r['train_r2']:.4f}"],
            ["Train MAE (apparent)", f"{r['train_mae']:.2f} m"],
        ])

        doc.add_paragraph()
        imp_heading = doc.add_paragraph()
        imp_heading.add_run(
            "Predictor importance – top 20 (sorted by bootstrap selection frequency)"
        ).bold = True

        imp = r["importance"].head(20)
        rows = [[
            "Predictor",
            "Full-fit Coef",
            "Boot Mean Coef",
            "Boot SD",
            "Selection Freq",
        ]]
        for _, row in imp.iterrows():
            rows.append([
                row["predictor"],
                f"{row['base_coef']:.5f}",
                f"{row['bootstrap_coef_mean']:.5f}",
                f"{row['bootstrap_coef_sd']:.5f}",
                f"{row['selection_frequency']:.3f}",
            ])
        _add_table_to_doc(doc, rows)

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX.name}")


def write_excel_predictions(df: pd.DataFrame, results: list[dict]) -> None:
    """
    Write an Excel file containing:
    - All original dataset rows and columns.
    - One prediction column per model appended at the right.

    Predictions are generated for every row that has sufficient (imputed)
    predictor data, including rows where the observed 6MWT4 is missing.
    The LASSO pipeline handles missing predictor values internally via
    median imputation fitted on the analysis set.
    """
    out = df.copy()

    for r in results:
        col_name = MODEL_PRED_COLS[r["model_name"]]
        pipe: Pipeline = r["base_pipe"]
        features: list[str] = r["features"]

        # Predict for all rows (pipeline imputes internally using medians fitted
        # on the analysis set where 6MWT4 was observed; this is intentional so
        # that patients with missing 6MWT4 also receive a predicted value).
        preds = pipe.predict(df[features])
        out[col_name] = preds

    out.to_excel(OUTPUT_XLSX, index=False)
    print(f"Saved: {OUTPUT_XLSX.name}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    """
    Orchestration function:
    1. Load data.
    2. For each of the 4 model specs, filter predictors to existing columns,
       run bootstrap LASSO, collect results.
    3. Write DOCX report.
    4. Write Excel prediction file.
    5. Print concise console summary.
    """
    # ── Load data ─────────────────────────────────────────────────────────────
    print(f"Loading: {INPUT_XLSX.name}")
    df = pd.read_excel(INPUT_XLSX)
    print(f"Dataset: {df.shape[0]:,} rows × {df.shape[1]} columns")

    # ── Coerce all candidate columns to numeric ───────────────────────────────
    # Non-numeric values become NaN and are subsequently median-imputed.
    all_candidate_cols = list(
        {c for cols in MODEL_SPECS.values() for c in cols}
    )
    for c in all_candidate_cols + ["6MWT4"]:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    # ── Fit all models ────────────────────────────────────────────────────────
    results: list[dict] = []
    for i, (model_name, candidate_features) in enumerate(MODEL_SPECS.items()):
        # Keep only columns that actually exist in the dataset
        valid_features = _filter_existing(candidate_features, df)

        print("\n" + "=" * 78)
        print(f"  {model_name}")
        print(
            f"  Candidate predictors: {len(candidate_features)} → "
            f"existing in data: {len(valid_features)}"
        )

        if not valid_features:
            print("  !! No valid features found – skipping model.")
            continue

        res = fit_bootstrap_lasso(df, valid_features, model_name, model_seed=RANDOM_STATE + i)
        results.append(res)

        # Brief per-model console summary
        print(
            f"  Patients: {res['n_patients']:,} | "
            f"Input vars: {res['n_input_variables']} | "
            f"Non-zero (full fit): {res['n_final_nonzero_base']} | "
            f"Stable (≥{STABILITY_THRESHOLD:.0%}): {res['n_stable_selected']}"
        )
        print(
            f"  CV  R²: {res['cv_r2']:.4f}  MAE: {res['cv_mae']:.2f} m  |  "
            f"Train R²: {res['train_r2']:.4f}  MAE: {res['train_mae']:.2f} m"
        )

    # ── Write outputs ─────────────────────────────────────────────────────────
    print("\n" + "=" * 78)
    write_docx_report(results)
    write_excel_predictions(df, results)
    print("Done.\n")


if __name__ == "__main__":
    main()
