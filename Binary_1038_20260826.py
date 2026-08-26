#!/usr/bin/env python3
"""Create scenario-based binary 6MWT outputs and reports."""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from sklearn.base import clone
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import ExtraTreesClassifier, RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.inspection import permutation_importance
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    balanced_accuracy_score,
    confusion_matrix,
    f1_score,
    precision_score,
    recall_score,
)
from sklearn.model_selection import StratifiedKFold, cross_val_predict, cross_validate
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler


ROOT = Path(__file__).resolve().parent
INPUT_XLSX = ROOT / "20260806_DeID.xlsx"
OUTPUT_XLSX = ROOT / "20260826_DeID.xlsx"
OUTPUT_REPORT = ROOT / "20260826_Binary_1038.docx"
OUTPUT_CODE = ROOT / "20260826_Binary_1038_Code.docx"

RANDOM_STATE = 42

DEMOGRAPHICS_PAC = ["Age", "Sex, F0 M1"]
STROKE_INFO = [
    "Dissection",
    "ACA",
    "Undetermined",
    "HemorrhageStroke",
    "LVS",
    "LVO",
    "Side_Right",
    "Side_Left",
    "Side_Bilateral",
    "Loc_CortSub",
    "Loc_Subcortical",
    "Loc_Infratentorial",
]
COMORBIDITIES_PAC = [
    "AF",
    "DM",
    "HTN",
    "Dyslipidemia",
    "CAD",
    "CKD",
    "RestrictiveLung",
    "GIUlcer",
    "LiverCirrhosis",
    "Hepatitis",
    "Parkinsonism",
    "Malignancy",
    "OldStroke",
    "Dementia",
    "Psychiatric",
    "Gout",
]
ACUTE_COMPLICATIONS_PAC = ["Pneumonia", "UTI", "GIB", "Cellulitis"]
ACUTE_TREATMENT_PAC = ["tPA", "EVT", "tPA+EVT"]
FUNCTIONAL_T1_PLUS_GS_IMPUTED = [
    "MRS1",
    "BI1",
    "FOIS1",
    "MNA1",
    "EuroQoL5D1",
    "IADL1",
    "BBS1",
    "Gait_Speed_1_Imputed",
    "FuglUE1",
    "FuglSEN1",
]
BEST_WORST_SCENARIO_BINARY = ["6MWT_Best_Scenario", "6MWT_Worst_Scenario"]

GROUPED_COLUMNS = (
    DEMOGRAPHICS_PAC
    + STROKE_INFO
    + COMORBIDITIES_PAC
    + ACUTE_COMPLICATIONS_PAC
    + ACUTE_TREATMENT_PAC
    + FUNCTIONAL_T1_PLUS_GS_IMPUTED
    + BEST_WORST_SCENARIO_BINARY
)
MODEL_FEATURES = DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED


def make_cv() -> StratifiedKFold:
    return StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)


def set_code_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
            "Courier New",
        )
        run.font.size = Pt(8)


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True


def build_modified_dataset(df: pd.DataFrame) -> pd.DataFrame:
    rename_map = {"IA": "EVT", "tPAIA": "tPA+EVT"}
    missing_rename_sources = [old for old in rename_map if old not in df.columns]
    if missing_rename_sources:
        raise KeyError(f"Required source columns are missing before renaming: {missing_rename_sources}")
    conflicting_targets = [new for old, new in rename_map.items() if old in df.columns and new in df.columns]
    if conflicting_targets:
        raise ValueError(f"Cannot rename columns because targets already exist: {conflicting_targets}")
    df = df.rename(columns=rename_map)

    first_6mwt = df["First_6MWT_TP"].astype("string")
    completion = df["PAC_Program_Completion"].astype("string")

    observed_6mwt = first_6mwt.isin(["T1", "T2", "T3", "T4"])
    never_6mwt = first_6mwt.eq("Never")
    optimistic_noncompletion = completion.isin(["Never", "Did not complete PAC program"])

    best = pd.Series(np.nan, index=df.index, dtype="float64")
    best.loc[observed_6mwt] = 1
    best.loc[never_6mwt & optimistic_noncompletion] = 1
    best.loc[never_6mwt & completion.eq("Completed PAC program")] = 0

    worst = pd.Series(np.nan, index=df.index, dtype="float64")
    worst.loc[observed_6mwt] = 1
    worst.loc[never_6mwt] = 0

    df = pd.concat(
        [
            df,
            pd.DataFrame(
                {
                    "6MWT_Best_Scenario": best.astype("Int64"),
                    "6MWT_Worst_Scenario": worst.astype("Int64"),
                },
                index=df.index,
            ),
        ],
        axis=1,
    )

    missing_grouped = [column for column in GROUPED_COLUMNS if column not in df.columns]
    if missing_grouped:
        raise KeyError(f"Required grouped columns are missing after restructuring: {missing_grouped}")

    remaining = [column for column in df.columns if column not in GROUPED_COLUMNS]
    return df[GROUPED_COLUMNS + remaining]


def build_candidates() -> dict[str, dict[str, object]]:
    numeric_transformer = Pipeline(
        [
            ("imputer", SimpleImputer(strategy="median")),
            ("scaler", StandardScaler()),
        ]
    )
    imputer_only_transformer = Pipeline([("imputer", SimpleImputer(strategy="median"))])

    logistic = Pipeline(
        [
            ("prep", ColumnTransformer([("num", numeric_transformer, MODEL_FEATURES)])),
            (
                "model",
                LogisticRegression(
                    max_iter=5000,
                    solver="liblinear",
                    class_weight="balanced",
                    random_state=RANDOM_STATE,
                ),
            ),
        ]
    )
    random_forest = Pipeline(
        [
            ("prep", ColumnTransformer([("num", imputer_only_transformer, MODEL_FEATURES)])),
            (
                "model",
                RandomForestClassifier(
                    n_estimators=500,
                    min_samples_leaf=2,
                    class_weight="balanced",
                    random_state=RANDOM_STATE,
                    n_jobs=1,
                ),
            ),
        ]
    )
    extra_trees = Pipeline(
        [
            ("prep", ColumnTransformer([("num", imputer_only_transformer, MODEL_FEATURES)])),
            (
                "model",
                ExtraTreesClassifier(
                    n_estimators=500,
                    min_samples_leaf=2,
                    class_weight="balanced",
                    random_state=RANDOM_STATE,
                    n_jobs=1,
                ),
            ),
        ]
    )

    return {
        "Logistic Regression": {
            "pipeline": logistic,
            "code": """best_pipeline = Pipeline([\n    ('prep', ColumnTransformer([\n        ('num', Pipeline([\n            ('imputer', SimpleImputer(strategy='median')),\n            ('scaler', StandardScaler()),\n        ]), MODEL_FEATURES)\n    ])),\n    ('model', LogisticRegression(max_iter=5000, solver='liblinear', class_weight='balanced', random_state=42)),\n])""",
        },
        "Random Forest": {
            "pipeline": random_forest,
            "code": """best_pipeline = Pipeline([\n    ('prep', ColumnTransformer([\n        ('num', Pipeline([\n            ('imputer', SimpleImputer(strategy='median')),\n        ]), MODEL_FEATURES)\n    ])),\n    ('model', RandomForestClassifier(n_estimators=500, min_samples_leaf=2, class_weight='balanced', random_state=42, n_jobs=1)),\n])""",
        },
        "Extra Trees": {
            "pipeline": extra_trees,
            "code": """best_pipeline = Pipeline([\n    ('prep', ColumnTransformer([\n        ('num', Pipeline([\n            ('imputer', SimpleImputer(strategy='median')),\n        ]), MODEL_FEATURES)\n    ])),\n    ('model', ExtraTreesClassifier(n_estimators=500, min_samples_leaf=2, class_weight='balanced', random_state=42, n_jobs=1)),\n])""",
        },
    }


def evaluate_target(df: pd.DataFrame, target: str) -> dict[str, object]:
    model_df = df[MODEL_FEATURES + [target]].dropna(subset=[target]).copy()
    X = model_df[MODEL_FEATURES]
    y = model_df[target].astype(int)
    cv_splits = list(make_cv().split(X, y))

    candidates = build_candidates()
    scoring = {
        "accuracy": "accuracy",
        "balanced_accuracy": "balanced_accuracy",
        "precision": "precision",
        "recall": "recall",
        "f1": "f1",
    }

    leaderboard_rows = []
    for name, payload in candidates.items():
        cv_result = cross_validate(
            payload["pipeline"],
            X,
            y,
            cv=cv_splits,
            scoring=scoring,
            n_jobs=-1,
        )
        leaderboard_rows.append(
            {
                "model": name,
                "accuracy": float(cv_result["test_accuracy"].mean()),
                "balanced_accuracy": float(cv_result["test_balanced_accuracy"].mean()),
                "precision": float(cv_result["test_precision"].mean()),
                "recall": float(cv_result["test_recall"].mean()),
                "f1": float(cv_result["test_f1"].mean()),
            }
        )

    leaderboard = (
        pd.DataFrame(leaderboard_rows)
        .sort_values(["balanced_accuracy", "accuracy", "f1"], ascending=False)
        .reset_index(drop=True)
    )
    best_name = str(leaderboard.iloc[0]["model"])
    best_pipeline_template = clone(candidates[best_name]["pipeline"])

    oof_predictions = cross_val_predict(clone(best_pipeline_template), X, y, cv=cv_splits, n_jobs=-1)
    metrics = {
        "accuracy": float(accuracy_score(y, oof_predictions)),
        "balanced_accuracy": float(balanced_accuracy_score(y, oof_predictions)),
        "precision": float(precision_score(y, oof_predictions, zero_division=0)),
        "recall": float(recall_score(y, oof_predictions, zero_division=0)),
        "f1": float(f1_score(y, oof_predictions, zero_division=0)),
    }

    tn, fp, fn, tp = confusion_matrix(y, oof_predictions, labels=[0, 1]).ravel()
    fold_importances = []
    for fold_index, (train_idx, test_idx) in enumerate(cv_splits):
        fold_pipeline = clone(best_pipeline_template)
        X_train, X_test = X.iloc[train_idx], X.iloc[test_idx]
        y_train, y_test = y.iloc[train_idx], y.iloc[test_idx]
        fold_pipeline.fit(X_train, y_train)
        fold_importance = permutation_importance(
            fold_pipeline,
            X_test,
            y_test,
            scoring="balanced_accuracy",
            n_repeats=30,
            random_state=RANDOM_STATE + fold_index,
            n_jobs=-1,
        )
        fold_importances.append(fold_importance.importances_mean)

    importance_array = np.vstack(fold_importances)
    importance_df = (
        pd.DataFrame(
            {
                "predictor": MODEL_FEATURES,
                "importance_mean": importance_array.mean(axis=0),
                "importance_std": importance_array.std(axis=0, ddof=1),
            }
        )
        .sort_values("importance_mean", ascending=False)
        .reset_index(drop=True)
    )

    return {
        "target": target,
        "n_total": int(len(model_df)),
        "n_positive": int(y.sum()),
        "n_negative": int((1 - y).sum()),
        "leaderboard": leaderboard,
        "best_model_name": best_name,
        "best_model_code": candidates[best_name]["code"],
        "metrics": metrics,
        "confusion": {"tn": int(tn), "fp": int(fp), "fn": int(fn), "tp": int(tp)},
        "importance": importance_df,
    }


def write_report(results: list[dict[str, object]]) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("20260826 Binary 6MWT Scenario Models").bold = True

    note = doc.add_paragraph(
        "Best models were selected by the highest 5-fold cross-validated balanced accuracy using demographics and functional_t1_plus_gs_imputed predictors only."
    )
    for run in note.runs:
        run.font.size = Pt(9)

    assumption = doc.add_paragraph(
        "For 6MWT_Best_Scenario, rows with First_6MWT_TP = Never were set to 1 when PAC_Program_Completion was Never or Did not complete PAC program, and set to 0 when PAC_Program_Completion was Completed PAC program."
    )
    for run in assumption.runs:
        run.font.size = Pt(9)

    for result in results:
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run(str(result["target"])).bold = True

        add_table(
            doc,
            [
                ["Patient count", "Value"],
                ["Total modeled patients", result["n_total"]],
                ["Positive class (1)", result["n_positive"]],
                ["Negative class (0)", result["n_negative"]],
                ["Best model", result["best_model_name"]],
            ],
        )

        leaderboard = result["leaderboard"]
        leaderboard_rows = [["Candidate model", "Accuracy", "Balanced Acc.", "Precision", "Recall", "F1"]]
        for _, row in leaderboard.iterrows():
            leaderboard_rows.append(
                [
                    row["model"],
                    f'{row["accuracy"]:.3f}',
                    f'{row["balanced_accuracy"]:.3f}',
                    f'{row["precision"]:.3f}',
                    f'{row["recall"]:.3f}',
                    f'{row["f1"]:.3f}',
                ]
            )
        add_table(doc, leaderboard_rows)

        metrics = result["metrics"]
        confusion = result["confusion"]
        add_table(
            doc,
            [
                ["Out-of-fold metric", "Value"],
                ["Accuracy", f'{metrics["accuracy"]:.3f}'],
                ["Balanced accuracy", f'{metrics["balanced_accuracy"]:.3f}'],
                ["Precision", f'{metrics["precision"]:.3f}'],
                ["Recall", f'{metrics["recall"]:.3f}'],
                ["F1", f'{metrics["f1"]:.3f}'],
                ["True negative", confusion["tn"]],
                ["False positive", confusion["fp"]],
                ["False negative", confusion["fn"]],
                ["True positive", confusion["tp"]],
            ],
        )

        importance_rows = [["Predictor", "Importance mean", "Importance SD"]]
        for _, row in result["importance"].head(10).iterrows():
            importance_rows.append(
                [
                    row["predictor"],
                    f'{row["importance_mean"]:.4f}',
                    f'{row["importance_std"]:.4f}',
                ]
            )
        add_table(doc, importance_rows)

    doc.save(OUTPUT_REPORT)


def write_code_doc(results: list[dict[str, object]]) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("20260826 Binary 6MWT Best Model Code").bold = True

    feature_paragraph = doc.add_paragraph(
        "MODEL_FEATURES = " + repr(MODEL_FEATURES)
    )
    set_code_font(feature_paragraph)

    for result in results:
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run(f'Best model for {result["target"]}: {result["best_model_name"]}').bold = True

        code = (
            result["best_model_code"]
            + "\n\ncv_splits = list(StratifiedKFold(n_splits=5, shuffle=True, random_state=42).split(X, y))\n"
            + "for fold_index, (train_idx, test_idx) in enumerate(cv_splits):\n"
            + "    fold_pipeline = clone(best_pipeline)\n"
            + "    X_train, X_test = X.iloc[train_idx], X.iloc[test_idx]\n"
            + "    y_train, y_test = y.iloc[train_idx], y.iloc[test_idx]\n"
            + "    fold_pipeline.fit(X_train, y_train)\n"
            + "    permutation_importance(fold_pipeline, X_test, y_test, scoring='balanced_accuracy', n_repeats=30, random_state=42 + fold_index, n_jobs=-1)"
        )
        paragraph = doc.add_paragraph(code)
        set_code_font(paragraph)

    doc.save(OUTPUT_CODE)


def main() -> None:
    df = pd.read_excel(INPUT_XLSX)
    modified_df = build_modified_dataset(df)
    modified_df.to_excel(OUTPUT_XLSX, index=False)

    results = [
        evaluate_target(modified_df, "6MWT_Best_Scenario"),
        evaluate_target(modified_df, "6MWT_Worst_Scenario"),
    ]
    write_report(results)
    write_code_doc(results)

    print(f"Saved dataset: {OUTPUT_XLSX.name}")
    print(f"Saved report: {OUTPUT_REPORT.name}")
    print(f"Saved code doc: {OUTPUT_CODE.name}")
    for result in results:
        metrics = result["metrics"]
        print(
            f'{result["target"]}: {result["best_model_name"]} | '
            f'accuracy={metrics["accuracy"]:.3f} | '
            f'balanced_accuracy={metrics["balanced_accuracy"]:.3f} | '
            f'precision={metrics["precision"]:.3f} | '
            f'recall={metrics["recall"]:.3f}'
        )


if __name__ == "__main__":
    main()
