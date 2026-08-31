#!/usr/bin/env python3
"""
20260831_IPCW_2009.py

Builds a DOCX explainer focused on the latest IPCW-based 6MWT prediction
workflow from prior-session artifacts.

Output:
- 20260831_IPCW_2009.docx
"""

from __future__ import annotations

from collections import Counter
import math
from pathlib import Path
from statistics import mean

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
SOURCE_DOC = ROOT / "20260828_5_Tiers_1201.docx"
SOURCE_SCRIPT = ROOT / "20260828_5_Tiers_1201.py"
OUT_DOC = ROOT / "20260831_IPCW_2009.docx"


def _small(paragraph, size: int = 9) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(size)


def _add_table(doc: Document, rows: list[list[object]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True


def _field_table_to_dict(table) -> dict[str, str]:
    out: dict[str, str] = {}
    for row in table.rows[1:]:
        if len(row.cells) < 2:
            continue
        key = row.cells[0].text.strip()
        value = row.cells[1].text.strip()
        if key:
            out[key] = value
    return out


def _to_float(value: str) -> float:
    try:
        return float(value)
    except ValueError:
        return float("nan")


def _to_int(value: str, default: int = 0) -> int:
    try:
        return int(value)
    except ValueError:
        return default


def _fmt_float(value: float, digits: int) -> str:
    return f"{value:.{digits}f}" if math.isfinite(value) else "N/A"


def parse_latest_ipcw_models() -> list[dict]:
    source = Document(SOURCE_DOC)
    parsed: list[dict] = []

    # Source report stores each tier/scenario block as 3 tables:
    # 1) metadata/metrics 2) top ridge predictors 3) binary leaderboard.
    for idx in range(0, len(source.tables), 3):
        if idx + 2 >= len(source.tables):
            break

        metric_dict = _field_table_to_dict(source.tables[idx])
        if "Tier" not in metric_dict or "Scenario" not in metric_dict:
            continue

        predictor_table = source.tables[idx + 1]
        predictors: list[dict] = []
        for row in predictor_table.rows[1:]:
            if len(row.cells) < 2:
                continue
            predictor = row.cells[0].text.strip()
            abs_coef_raw = row.cells[1].text.strip()
            if not predictor:
                continue
            try:
                abs_coef = float(abs_coef_raw)
            except ValueError:
                continue
            predictors.append({"predictor": predictor, "abs_coef": abs_coef})

        leaderboard_table = source.tables[idx + 2]
        leaderboard: list[dict] = []
        for row in leaderboard_table.rows[1:]:
            if len(row.cells) < 6:
                continue
            model_name = row.cells[0].text.strip()
            if not model_name:
                continue
            try:
                leaderboard.append(
                    {
                        "model": model_name,
                        "accuracy": float(row.cells[1].text.strip()),
                        "balanced_accuracy": float(row.cells[2].text.strip()),
                        "precision": float(row.cells[3].text.strip()),
                        "recall": float(row.cells[4].text.strip()),
                        "f1": float(row.cells[5].text.strip()),
                    }
                )
            except ValueError:
                continue

        leaderboard = sorted(
            leaderboard,
            key=lambda x: (x["balanced_accuracy"], x["accuracy"], x["f1"]),
            reverse=True,
        )
        if not leaderboard:
            continue
        balacc_margin = (
            leaderboard[0]["balanced_accuracy"] - leaderboard[1]["balanced_accuracy"]
            if len(leaderboard) >= 2
            else 0.0
        )

        parsed.append(
            {
                "tier": metric_dict["Tier"],
                "scenario": metric_dict["Scenario"],
                "ipcw_denominator": metric_dict.get("IPCW denominator model", ""),
                "ipcw_numerator": metric_dict.get("IPCW numerator model", ""),
                "ipcw_winsorization": metric_dict.get("IPCW winsorization", ""),
                "binary_model": metric_dict.get("Selected binary model", ""),
                "bal_acc": _to_float(metric_dict.get("Binary OOF balanced accuracy", "nan")),
                "f1": _to_float(metric_dict.get("Binary OOF F1", "nan")),
                "weighted_model": metric_dict.get("Weighted regression model", ""),
                "weighted_r2": _to_float(metric_dict.get("Weighted OOF R²", "nan")),
                "weighted_mae": _to_float(metric_dict.get("Weighted OOF MAE", "nan")),
                "walk_yes": _to_int(metric_dict.get("Non-completers predicted to walk", "0")),
                "walk_no": _to_int(metric_dict.get("Non-completers predicted not to walk", "0")),
                "feature_count": _to_int(metric_dict.get("Feature count", "0")),
                "top_predictors": predictors[:10],
                "leaderboard": leaderboard,
                "balacc_margin_top2": balacc_margin,
            }
        )

    return parsed


def write_ipcw_doc(models: list[dict]) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("Latest IPCW Models for 6MWT Prediction with Explainers (Session 2009)").bold = True

    intro = doc.add_paragraph(
        "This document consolidates the latest IPCW model outputs from prior agent-session artifacts and explains "
        "how to apply the workflow to generate patient-level 6MWT predictions for non-completers."
    )
    _small(intro)

    provenance = doc.add_paragraph(
        f"Provenance: extracted from {SOURCE_DOC.name} (generated by {SOURCE_SCRIPT.name}). "
        "The modeled workflow combines stabilized IPCW, scenario-specific binary walking classification, "
        "and IPCW-weighted Ridge regression for continuous 6MWT4 prediction."
    )
    _small(provenance)

    model_rows = [[
        "Tier",
        "Scenario",
        "Selected binary model",
        "Bal Acc",
        "F1",
        "Weighted R²",
        "Weighted MAE",
        "Pred Walk/No-walk",
    ]]
    for row in models:
        model_rows.append(
            [
                row["tier"],
                row["scenario"],
                row["binary_model"],
                _fmt_float(row["bal_acc"], 3),
                _fmt_float(row["f1"], 3),
                _fmt_float(row["weighted_r2"], 4),
                _fmt_float(row["weighted_mae"], 1),
                f"{row['walk_yes']}/{row['walk_no']}",
            ]
        )
    _add_table(doc, model_rows)

    model_counter = Counter(r["binary_model"] for r in models if r["binary_model"])
    finite_bal = [r["bal_acc"] for r in models if math.isfinite(r["bal_acc"])]
    finite_r2 = [r["weighted_r2"] for r in models if math.isfinite(r["weighted_r2"])]
    finite_fit_rows = [r for r in models if math.isfinite(r["weighted_r2"])]
    finite_cls_rows = [r for r in models if math.isfinite(r["bal_acc"])]

    missing = []
    if not finite_bal:
        missing.append("balanced_accuracy")
    if not finite_r2:
        missing.append("weighted_r2")
    if not finite_fit_rows:
        missing.append("best_fit_rows")
    if not finite_cls_rows:
        missing.append("best_classification_rows")
    if missing:
        raise RuntimeError(f"Missing finite IPCW metrics for summary: {', '.join(missing)}")

    avg_bal = mean(finite_bal)
    avg_r2 = mean(finite_r2)
    best_fit = max(finite_fit_rows, key=lambda x: x["weighted_r2"])
    best_cls = max(finite_cls_rows, key=lambda x: x["bal_acc"])
    common_model, common_count = model_counter.most_common(1)[0] if model_counter else ("N/A", 0)

    summary = doc.add_paragraph(
        f"Across {len(models)} tier/scenario models, mean binary balanced accuracy was {avg_bal:.3f} and mean "
        f"IPCW-weighted regression R² was {avg_r2:.4f}. Most frequent binary classifier was "
        f"{common_model} ({common_count}/{len(models)} models). "
        f"Best weighted fit: {best_fit['tier']} {best_fit['scenario']} (R²={best_fit['weighted_r2']:.4f}); "
        f"best classification discrimination: {best_cls['tier']} {best_cls['scenario']} (Bal Acc={best_cls['bal_acc']:.3f})."
    )
    _small(summary)

    doc.add_paragraph()
    explainer_header = doc.add_paragraph()
    explainer_header.add_run("Tier/scenario explainers").bold = True

    for row in models:
        top_features = ", ".join(
            f"{x['predictor']} ({x['abs_coef']:.2f})" for x in row["top_predictors"][:5]
        ) or "No predictor table available"
        p = doc.add_paragraph(
            f"• {row['tier']} {row['scenario']}: {row['binary_model']} reached Bal Acc {_fmt_float(row['bal_acc'], 3)} "
            f"(top-2 margin {_fmt_float(row['balacc_margin_top2'], 3)}); weighted model {row['weighted_model']} achieved R² "
            f"{_fmt_float(row['weighted_r2'], 4)} and MAE {_fmt_float(row['weighted_mae'], 1)}. Predicted non-completers walking: "
            f"{row['walk_yes']} vs not walking: {row['walk_no']}. Top weighted predictors: {top_features}."
        )
        _small(p)

    doc.add_paragraph()
    apply_header = doc.add_paragraph()
    apply_header.add_run("How to apply this IPCW workflow for 6MWT prediction").bold = True

    apply_steps = [
        "1) Prepare baseline predictors used by the target tier and ensure missingness is imputed consistently.",
        "2) Estimate stabilized IPCW from completion status: denominator P(completion|tier features), numerator P(completion|stabilizer), then winsorize completer weights at the 1st/99th percentile.",
        "3) For each scenario (Best/Worst), compare candidate binary classifiers and select the top model by cross-validated balanced accuracy.",
        "4) Fit a weighted Ridge model on observed completer 6MWT4 using IPCW sample weights and evaluate out-of-fold weighted R²/MAE.",
        "5) Apply the selected scenario-specific binary model to non-completers; assign 0 to predicted non-walkers and weighted-Ridge predictions to predicted walkers.",
        "6) Export tier/scenario patient-level predictions together with model diagnostics and feature explainers for auditability.",
    ]
    for step in apply_steps:
        p = doc.add_paragraph(step)
        _small(p)

    source_note = doc.add_paragraph(
        f"Reproducibility: run `python {SOURCE_SCRIPT.name}` to regenerate the source IPCW outputs, and "
        f"`python {Path(__file__).name}` to regenerate this explainer document."
    )
    _small(source_note, 8)

    doc.save(OUT_DOC)


def main() -> None:
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"Missing required source document: {SOURCE_DOC}")
    if not SOURCE_SCRIPT.exists():
        print(f"Warning: source script not found ({SOURCE_SCRIPT.name}); proceeding with DOCX-only provenance.")

    rows = parse_latest_ipcw_models()
    if not rows:
        raise RuntimeError("No IPCW model blocks were parsed from the source DOCX.")

    write_ipcw_doc(rows)
    print(f"Saved: {OUT_DOC.name}")
    print(f"Code implementation artifact: {Path(__file__).name}")


if __name__ == "__main__":
    main()
