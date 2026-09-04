#!/usr/bin/env python3
"""
20260904_Comprehensive_1018.py
======================
Combined bootstrap LASSO (no IPCW) + single-completion-model IPCW pipeline for
6MWT4 prediction.  Two key improvements over the 20260831_LASSO_1047 baseline:

  1.  Model 5 (recovery-trajectory) is now restricted to patients whose
      Rehab LOS category is "21-42 days" or ">42 days".  Patients outside
      those groups are excluded from the Model 5 analysis entirely (not
      counted, not imputed).  T1T2_Change values are valid only for
      qualifying patients; they are left blank and not imputed for
      non-qualifying patients.

  2.  IPCW section uses ONE well-specified completion model (the richest
      clinically justified predictor set: demographics + stroke info +
      comorbidities + acute complications + NIHSS out + functional T1 with
      gait speed) to derive stabilized weights.  The same weights are then
      applied uniformly across retained IPCW outcome regressions.
      This is consistent with conventional IPCW practice, avoids conflating
      "tier effect on outcome" with "tier effect on weighting", and prevents
      re-deriving a new completion model per tier.

Outputs (written to repo root):
    20260904_Comprehensive_1018.docx  – combined LASSO + IPCW report
    20260904_Comprehensive_1018.xlsx  – full dataset with per-patient predictions

Author: auto-generated 2026-09-04
"""

from __future__ import annotations

import math
import warnings
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from sklearn.base import clone
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import ExtraTreesClassifier, RandomForestClassifier
from sklearn.exceptions import ConvergenceWarning
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LassoCV, LogisticRegression, Ridge
from sklearn.metrics import (
    balanced_accuracy_score,
    mean_absolute_error,
    r2_score,
)
from sklearn.model_selection import KFold, StratifiedKFold, cross_val_predict, cross_validate
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler

warnings.filterwarnings("ignore", category=ConvergenceWarning)

# ─────────────────────────────────────────────────────────────────────────────
# FILE PATHS
# ─────────────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent
INPUT_XLSX = ROOT / "20260903_Comprehensive_1327.xlsx"
OUTPUT_DOCX = ROOT / "20260904_Comprehensive_1018.docx"
OUTPUT_XLSX = ROOT / "20260904_Comprehensive_1018.xlsx"

# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL SETTINGS
# ─────────────────────────────────────────────────────────────────────────────
RANDOM_STATE = 42
N_BOOTSTRAP_LASSO = 1000
N_BOOTSTRAP_IPCW = 500
CV_FOLDS = 5
ALPHAS = np.logspace(-4, 2, 120)
STABILITY_THRESHOLD = 0.70
WINSOR_PCT = 99
WINSORIZATION_THRESHOLDS = [90, 95, 99]
NUMERATOR_MODE = "intercept"

# Rehab LOS categories that qualify for Model 5 / T1T2_Change analysis
QUALIFYING_REHAB_LOS = ["21-42 days", ">42 days"]

# ─────────────────────────────────────────────────────────────────────────────
# VARIABLE CATEGORIES
# ─────────────────────────────────────────────────────────────────────────────
DEMOGRAPHICS_PAC = ["Age", "Sex, F0 M1"]

STROKE_INFO = [
    "Dissection", "ACA", "Undetermined", "HemorrhageStroke",
    "LVS", "LVO",
    "Side_Right", "Side_Left", "Side_Bilateral",
    "Loc_CortSub", "Loc_Subcortical", "Loc_Infratentorial",
]

COMORBIDITIES_PAC = [
    "AF", "DM", "HTN", "Dyslipidemia", "CAD", "CKD",
    "RestrictiveLung", "GIUlcer", "LiverCirrhosis", "Hepatitis",
    "Parkinsonism", "Malignancy", "OldStroke", "Dementia", "Psychiatric", "Gout",
]

ACUTE_COMPLICATIONS_PAC = ["Pneumonia", "UTI", "GIB", "Cellulitis"]

FUNCTIONAL_T1_PLUS_GS_IMPUTED = [
    "MRS1", "BI1", "FOIS1", "MNA1", "EuroQoL5D1", "IADL1",
    "BBS1", "Gait_Speed_1_Imputed", "FuglUE1", "FuglSEN1",
]

FUNCTIONAL_T1_NO_GS = [
    x for x in FUNCTIONAL_T1_PLUS_GS_IMPUTED if x != "Gait_Speed_1_Imputed"
]

NIHSS_OUT = [
    "ConsOut", "AnswerOut", "OrderOut", "EOMOut", "VisualOut",
    "FacialOut", "LUOut", "RUOut", "LLOut", "RLOut",
    "Coordinateout", "SensoryOut", "LanguageOut", "ArticulateOut", "NeglectOut",
]

NIHSS_IN = [
    "ConsIn", "AnswerIn", "OrderIn", "EOMIn", "VisualIn",
    "FaceIn", "LUIn", "RUIn", "LLIn", "RLIn",
    "CoordinateIn", "SensoryIn", "LanguageIn", "ArticulateIn", "NeglectIn",
]

T1T2_IMPROVEMENT = [
    "BI_T1T2_Change", "BBS_T1T2_Change", "MRS_T1T2_Change",
    "FOIS_T1T2_Change", "MNA_T1T2_Change", "IADL_T1T2_Change",
    "FuglUE_T1T2_Change", "FuglSEN_T1T2_Change", "EuroQoL5D_T1T2_Change",
]

FUNCTIONAL_AUDIT_SPECS = [
    {
        "name": "BI",
        "t1": "BI1",
        "t2": "BI2",
        "t3": "BI3",
        "t4": "BI 4",
        "t4_label": "BI4",
        "change": "BI_T1T2_Change",
        "valid_min": 0,
        "valid_max": 100,
        "auto_clear_zero_t2": True,
        "note": "BI2=0 treated as no T2 BI assessment per task request.",
    },
    {
        "name": "BBS",
        "t1": "BBS1",
        "t2": "BBS2",
        "t3": "BBS3",
        "t4": "BBS4",
        "t4_label": "BBS4",
        "change": "BBS_T1T2_Change",
        "valid_min": 0,
        "valid_max": 56,
        "auto_clear_zero_t2": True,
        "note": "T1T2-flagged BBS rows (T2=0 with non-null change) are cleared per task request.",
    },
    {
        "name": "MRS",
        "t1": "MRS1",
        "t2": "MRS2",
        "t3": "MRS3",
        "t4": "MRS4",
        "t4_label": "MRS4",
        "change": "MRS_T1T2_Change",
        "valid_min": 0,
        "valid_max": 6,
        "auto_clear_zero_t2": True,
        "note": "T1T2-flagged MRS rows (T2=0 with non-null change) are cleared per task request.",
    },
    {
        "name": "FOIS",
        "t1": "FOIS1",
        "t2": "FOIS2",
        "t3": "FOIS3",
        "t4": "FOIS4",
        "t4_label": "FOIS4",
        "change": "FOIS_T1T2_Change",
        "valid_min": 1,
        "valid_max": 7,
        "auto_clear_zero_t2": True,
        "note": "FOIS valid range starts at 1, so FOIS2=0 was treated as no assessment.",
    },
    {
        "name": "MNA",
        "t1": "MNA1",
        "t2": "MNA2",
        "t3": "MNA3",
        "t4": "MNA4",
        "t4_label": "MNA4",
        "change": "MNA_T1T2_Change",
        "valid_min": 0,
        "valid_max": 30,
        "auto_clear_zero_t2": True,
        "note": "T1T2-flagged MNA rows (T2=0 with non-null change) are cleared per task request.",
    },
    {
        "name": "IADL",
        "t1": "IADL1",
        "t2": "IADL2",
        "t3": "IADL3",
        "t4": "IADL4",
        "t4_label": "IADL4",
        "change": "IADL_T1T2_Change",
        "valid_min": 0,
        "valid_max": 8,
        "auto_clear_zero_t2": True,
        "note": "T1T2-flagged IADL rows (T2=0 with non-null change) are cleared per task request.",
    },
    {
        "name": "FuglUE",
        "t1": "FuglUE1",
        "t2": "FuglUE2",
        "t3": "FuglUE3",
        "t4": "FuglUE4",
        "t4_label": "FuglUE4",
        "change": "FuglUE_T1T2_Change",
        "valid_min": 0,
        "valid_max": 66,
        "auto_clear_zero_t2": True,
        "note": "T1T2-flagged FuglUE rows (T2=0 with non-null change) are cleared per task request.",
    },
    {
        "name": "FuglSEN",
        "t1": "FuglSEN1",
        "t2": "FuglSEN2",
        "t3": "FuglSEN3",
        "t4": "FuglSEN4",
        "t4_label": "FuglSEN4",
        "change": "FuglSEN_T1T2_Change",
        "valid_min": 0,
        "valid_max": 44,
        "auto_clear_zero_t2": True,
        "note": "T1T2-flagged FuglSEN rows (T2=0 with non-null change) are cleared per task request.",
    },
    {
        "name": "EuroQoL5D",
        "t1": "EuroQoL5D1",
        "t2": "EuroQoL5D2",
        "t3": "EuroQoL5D3",
        "t4": "EuroQoL5D4",
        "t4_label": "EuroQoL5D4",
        "change": "EuroQoL5D_T1T2_Change",
        "valid_min": 0,
        "valid_max": 15,
        "auto_clear_zero_t2": True,
        "note": "T1T2-flagged EuroQoL5D rows (T2=0 with non-null change) are cleared per task request.",
    },
]

# Richest completion predictor set for the single IPCW completion model
# (demographics + stroke info + comorbidities + acute complications +
#  NIHSS out + functional T1 with gait speed)
IPCW_COMPLETION_FEATURES = (
    DEMOGRAPHICS_PAC
    + STROKE_INFO
    + COMORBIDITIES_PAC
    + ACUTE_COMPLICATIONS_PAC
    + NIHSS_OUT
    + FUNCTIONAL_T1_PLUS_GS_IMPUTED
)

# ─────────────────────────────────────────────────────────────────────────────
# IPCW tier definitions and LASSO model definitions
# ─────────────────────────────────────────────────────────────────────────────
IPCW_TIERS: dict[str, list[str]] = {
    "Tier 1": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED,
    "Tier 2": DEMOGRAPHICS_PAC + [f for f in FUNCTIONAL_T1_PLUS_GS_IMPUTED if f != "Gait_Speed_1_Imputed"],
    "Tier 3": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + COMORBIDITIES_PAC + STROKE_INFO + ACUTE_COMPLICATIONS_PAC,
    "Tier 4": DEMOGRAPHICS_PAC + [f for f in FUNCTIONAL_T1_PLUS_GS_IMPUTED if f != "Gait_Speed_1_Imputed"] + COMORBIDITIES_PAC + STROKE_INFO + ACUTE_COMPLICATIONS_PAC,
    "Tier 5": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + T1T2_IMPROVEMENT,
}

BASE_MODEL_SPECS: dict[str, list[str]] = {
    "Model 1: Admission Functional Core (with gait speed imputed)": (
        DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED
    ),
    "Model 2: Comprehensive Clinical + NIHSS + Functional (with gait speed imputed)": (
        DEMOGRAPHICS_PAC
        + STROKE_INFO
        + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC
        + NIHSS_OUT
        + FUNCTIONAL_T1_PLUS_GS_IMPUTED
    ),
    "Model 3: Comprehensive Clinical + NIHSS + Functional (without gait speed imputed)": (
        DEMOGRAPHICS_PAC
        + STROKE_INFO
        + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC
        + NIHSS_OUT
        + FUNCTIONAL_T1_NO_GS
    ),
    "Model 4: Bedside Mobility Core": (
        DEMOGRAPHICS_PAC + ["BBS1", "Gait_Speed_1_Imputed"]
    ),
    "Model 5: Recovery Trajectory (LOS ≥21 days)": (
        DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + T1T2_IMPROVEMENT
    ),
    "Model 6: Clinical-Only Reference": (
        DEMOGRAPHICS_PAC
        + STROKE_INFO
        + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC
        + NIHSS_OUT
    ),
}

NON_REDUNDANT_TIER_MODEL_SPECS: dict[str, list[str]] = {
    "Model 7: Functional Core (without gait speed imputed)": IPCW_TIERS["Tier 2"],
    "Model 8: Comprehensive Clinical + Functional (with gait speed imputed)": IPCW_TIERS["Tier 3"],
    "Model 9: Comprehensive Clinical + Functional (without gait speed imputed)": IPCW_TIERS["Tier 4"],
}

MODEL_SPECS: dict[str, list[str]] = {**BASE_MODEL_SPECS, **NON_REDUNDANT_TIER_MODEL_SPECS}
RETAINED_MODEL_ORDER = [
    "Model 5: Recovery Trajectory (LOS ≥21 days)",
    "Model 2: Comprehensive Clinical + NIHSS + Functional (with gait speed imputed)",
    "Model 8: Comprehensive Clinical + Functional (with gait speed imputed)",
    "Model 1: Admission Functional Core (with gait speed imputed)",
    "Model 4: Bedside Mobility Core",
]

MODEL_PRED_COLS = {
    "Model 1: Admission Functional Core (with gait speed imputed)": "LASSO_Pred_Model1",
    "Model 2: Comprehensive Clinical + NIHSS + Functional (with gait speed imputed)": "LASSO_Pred_Model2",
    "Model 3: Comprehensive Clinical + NIHSS + Functional (without gait speed imputed)": "LASSO_Pred_Model3",
    "Model 4: Bedside Mobility Core": "LASSO_Pred_Model4",
    "Model 5: Recovery Trajectory (LOS ≥21 days)": "LASSO_Pred_Model5",
    "Model 6: Clinical-Only Reference": "LASSO_Pred_Model6",
    "Model 7: Functional Core (without gait speed imputed)": "LASSO_Pred_Model7",
    "Model 8: Comprehensive Clinical + Functional (with gait speed imputed)": "LASSO_Pred_Model8",
    "Model 9: Comprehensive Clinical + Functional (without gait speed imputed)": "LASSO_Pred_Model9",
}

MODEL_EXPLAINERS = {
    "Model 1: Admission Functional Core (with gait speed imputed)": (
        "Functional-admission model focused on bedside baseline impairment and mobility markers."
    ),
    "Model 2: Comprehensive Clinical + NIHSS + Functional (with gait speed imputed)": (
        "Full clinical-performance model combining neurological burden, medical complexity, and functional status including imputed gait speed."
    ),
    "Model 3: Comprehensive Clinical + NIHSS + Functional (without gait speed imputed)": (
        "Full clinical-performance model that removes imputed gait speed to test whether the broader functional battery can compensate."
    ),
    "Model 4: Bedside Mobility Core": (
        "Compact bedside model prioritising quick deployment with only age, sex, balance, and gait speed."
    ),
    "Model 5: Recovery Trajectory (LOS ≥21 days)": (
        "Recovery-trajectory model augmenting admission function with early change scores between T1 and T2. "
        "RESTRICTED to patients with Rehab LOS of 21-42 days or >42 days, because T1T2_Change variables "
        "are only meaningful for patients who have sufficient rehabilitation exposure. Patients outside "
        "these LOS categories are excluded entirely (not counted, T1T2_Change values not imputed)."
    ),
    "Model 6: Clinical-Only Reference": (
        "Reference neurological/comorbidity model without functional T1 measures, retained for comparison."
    ),
    "Model 7: Functional Core (without gait speed imputed)": (
        "Tier 2 predictor set replicated in Step 1 LASSO to isolate the impact of removing imputed gait speed."
    ),
    "Model 8: Comprehensive Clinical + Functional (with gait speed imputed)": (
        "Tier 3 predictor set replicated in Step 1 LASSO, combining broad clinical burden with functional T1 including gait speed."
    ),
    "Model 9: Comprehensive Clinical + Functional (without gait speed imputed)": (
        "Tier 4 predictor set replicated in Step 1 LASSO as the gait-speed-removed counterpart to Tier 3."
    ),
}

SCENARIOS = ["6MWT_Best_Scenario", "6MWT_Worst_Scenario"]

# ─────────────────────────────────────────────────────────────────────────────
# SHARED UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def _filter_existing(cols: list[str], df: pd.DataFrame) -> list[str]:
    return [c for c in cols if c in df.columns]


PREDICTOR_CATEGORY_MAP: dict[str, list[str]] = {
    "Demographics": DEMOGRAPHICS_PAC,
    "Stroke info": STROKE_INFO,
    "Comorbidities": COMORBIDITIES_PAC,
    "Acute complications": ACUTE_COMPLICATIONS_PAC,
    "NIHSS out": NIHSS_OUT,
    "Functional T1": FUNCTIONAL_T1_PLUS_GS_IMPUTED,
    "T1T2 improvement": T1T2_IMPROVEMENT,
}


def _predictor_category_counts(features: list[str]) -> dict[str, int]:
    feature_set = set(features)
    return {
        name: int(len(feature_set.intersection(cols)))
        for name, cols in PREDICTOR_CATEGORY_MAP.items()
    }


def _predictor_category_summary_text(features: list[str]) -> str:
    counts = _predictor_category_counts(features)
    ordered = [f"{name}={count}" for name, count in counts.items() if count > 0]
    gs_status = (
        "with gait speed imputed"
        if "Gait_Speed_1_Imputed" in features
        else "without gait speed imputed"
    )
    return f"{gs_status}; " + (", ".join(ordered) if ordered else "no categorized predictors")


def _add_table_to_doc(doc: Document, rows: list[list]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True


def _small_para(paragraph, size: int = 9) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(size)


def _fmt(v: float, d: int = 4) -> str:
    return f"{v:.{d}f}" if math.isfinite(v) else "N/A"


# ─────────────────────────────────────────────────────────────────────────────
# LASSO PIPELINE
# ─────────────────────────────────────────────────────────────────────────────

def _build_lasso_pipeline(features: list[str]) -> Pipeline:
    numeric_transformer = Pipeline([
        ("imputer", SimpleImputer(strategy="median")),
        ("scaler", StandardScaler()),
    ])
    preprocessor = ColumnTransformer([
        ("num", numeric_transformer, features)
    ])
    lasso = LassoCV(
        alphas=ALPHAS,
        cv=CV_FOLDS,
        random_state=RANDOM_STATE,
        max_iter=50_000,
        n_jobs=-1,
    )
    return Pipeline([
        ("prep", preprocessor),
        ("model", lasso),
    ])


def _top_predictors(importance: pd.DataFrame, *, limit: int = 5, positive: bool | None = None) -> pd.DataFrame:
    frame = importance.copy()
    if positive is True:
        frame = frame[frame["bootstrap_coef_mean"] > 0]
    elif positive is False:
        frame = frame[frame["bootstrap_coef_mean"] < 0]
    return frame.sort_values(
        ["selection_frequency", "abs_bootstrap_coef_mean"],
        ascending=[False, False],
    ).head(limit)


def _format_predictor_summary(frame: pd.DataFrame) -> str:
    if frame.empty:
        return "none"
    parts = []
    for _, row in frame.iterrows():
        parts.append(
            f"{row['predictor']} ({row['bootstrap_coef_mean']:+.1f}; freq {row['selection_frequency']:.0%})"
        )
    return ", ".join(parts)


def apply_manual_patient_corrections(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    corrected = df.copy()
    correction_rules = [
        {"ID": 317, "Column": "MNA2", "Corrected_Value": 16.0},
        {"ID": 528, "Column": "FuglUE2", "Corrected_Value": 62.0},
    ]

    for col in ["FuglSEN1", "FuglSEN2", "FuglSEN3", "FuglSEN4"]:
        correction_rules.extend([
            {"ID": 329, "Column": col, "Condition": lambda s: s > 44, "Corrected_Value": 44.0},
            {"ID": 340, "Column": col, "Condition": lambda s: s > 44, "Corrected_Value": 44.0},
            {"ID": 356, "Column": col, "Condition": lambda s: s > 44, "Corrected_Value": 44.0},
        ])

    applied_rows: list[dict] = []
    for rule in correction_rules:
        id_mask = pd.to_numeric(corrected.get("ID"), errors="coerce").eq(rule["ID"])
        if not id_mask.any() or rule["Column"] not in corrected.columns:
            continue

        series = pd.to_numeric(corrected[rule["Column"]], errors="coerce")
        if "Condition" in rule:
            target_mask = id_mask & rule["Condition"](series)
        else:
            target_mask = id_mask
        if not target_mask.any():
            continue

        for idx in corrected.index[target_mask]:
            original = corrected.at[idx, rule["Column"]]
            corrected.at[idx, rule["Column"]] = rule["Corrected_Value"]
            applied_rows.append({
                "ID": int(rule["ID"]),
                "Column": rule["Column"],
                "Original_Value": original,
                "Corrected_Value": rule["Corrected_Value"],
                "Reason": "Manual correction from task request",
            })

    return corrected, pd.DataFrame(applied_rows)


def extract_fuglsen_outliers(df: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict] = []
    for col in ["FuglSEN1", "FuglSEN2", "FuglSEN3", "FuglSEN4"]:
        if col not in df.columns:
            continue
        values = pd.to_numeric(df[col], errors="coerce")
        mask = values > 44
        if mask.any():
            flagged = pd.DataFrame({
                "ID": df.get("ID"),
                "Visit_Column": col,
                "Observed_Value": values,
                "Allowed_Max": 44,
                "Rehab_LOS_Category": df.get("Rehab_LOS_Category"),
            }).loc[mask]
            rows.extend(flagged.to_dict("records"))
    return pd.DataFrame(rows)


def audit_and_clean_t1t2_data(df: pd.DataFrame) -> tuple[pd.DataFrame, dict]:
    """
    Audit T1/T2 change variables and blank implausible change scores.

    Applied cleaning rules are intentionally conservative:
      • BI_T1T2_Change is blanked for Rehab_LOS_Category <=21 days.
      • BI_T1T2_Change is also blanked for Rehab_LOS_Category 21-42 days when BI2 == 0.
      • FOIS_T1T2_Change is blanked whenever FOIS2 == 0 because FOIS valid scores start at 1.
      • Any T1T2 change attached to an out-of-range T2 value or itself out of range is blanked.
    Other zero-at-T2 patterns are surfaced in audit tables for manual review.
    """
    cleaned = df.copy()
    summary_rows: list[dict] = []
    flagged_rows: list[dict] = []
    illogical_rows: list[dict] = []
    raw_outlier_rows: list[dict] = []

    for spec in FUNCTIONAL_AUDIT_SPECS:
        cols = [spec["t1"], spec["t2"], spec["t3"], spec["t4"], spec["change"]]
        if any(col not in cleaned.columns for col in cols):
            continue

        t1 = pd.to_numeric(cleaned[spec["t1"]], errors="coerce")
        t2 = pd.to_numeric(cleaned[spec["t2"]], errors="coerce")
        t3 = pd.to_numeric(cleaned[spec["t3"]], errors="coerce")
        t4 = pd.to_numeric(cleaned[spec["t4"]], errors="coerce")
        change = pd.to_numeric(cleaned[spec["change"]], errors="coerce")

        valid_change = change.notna()
        negative_before = int((valid_change & change.lt(0)).sum())
        zero_t2 = t2.fillna(0).eq(0)
        later_nonzero = t3.fillna(0).ne(0) | t4.fillna(0).ne(0)
        rehab_los = (
            cleaned["Rehab_LOS_Category"].astype(str).str.strip()
            if "Rehab_LOS_Category" in cleaned.columns
            else pd.Series("", index=cleaned.index)
        )

        value_span = spec["valid_max"] - spec["valid_min"]
        t2_out_of_range = t2.notna() & ((t2 < spec["valid_min"]) | (t2 > spec["valid_max"]))
        change_out_of_range = valid_change & (
            (change < -value_span) | (change > value_span)
        )
        clear_zero_t2 = zero_t2 & valid_change if spec["auto_clear_zero_t2"] else pd.Series(False, index=cleaned.index)
        bi_short_stay_clear = pd.Series(False, index=cleaned.index)
        bi_mid_stay_zero_t2_clear = pd.Series(False, index=cleaned.index)
        if spec["name"] == "BI":
            bi_short_stay_clear = rehab_los.eq("<=21 days") & valid_change
            bi_mid_stay_zero_t2_clear = rehab_los.eq("21-42 days") & zero_t2 & valid_change
            clear_zero_t2 = clear_zero_t2 | bi_mid_stay_zero_t2_clear

        for visit_col, visit_values in [
            (spec["t1"], t1),
            (spec["t2"], t2),
            (spec["t3"], t3),
            (spec["t4"], t4),
        ]:
            below_min = visit_values < spec["valid_min"]
            if visit_col != spec["t1"]:
                below_min &= ~visit_values.eq(0)
            raw_outlier_mask = visit_values.notna() & (
                below_min | (visit_values > spec["valid_max"])
            )
            if visit_col == spec["t2"]:
                raw_outlier_mask &= ~clear_zero_t2
            if raw_outlier_mask.any():
                flagged = pd.DataFrame({
                    "ID": cleaned.get("ID"),
                    "Assessment": spec["name"],
                    "Visit_Column": visit_col,
                    "Observed_Value": visit_values,
                    "Allowed_Range": f"[{spec['valid_min']}, {spec['valid_max']}]",
                    "Affects_T1T2_Change_Cleanup": np.where(
                        visit_col == spec["t2"],
                        "Yes",
                        "No",
                    ),
                    "Rehab_LOS_Category": cleaned.get("Rehab_LOS_Category"),
                    "PAC_Program_Completion": cleaned.get("PAC_Program_Completion"),
                }).loc[raw_outlier_mask]
                raw_outlier_rows.extend(flagged.to_dict("records"))

        clear_out_of_range = valid_change & (t2_out_of_range | change_out_of_range) & ~clear_zero_t2
        clear_mask = bi_short_stay_clear | clear_zero_t2 | clear_out_of_range

        original_change = change.copy()
        cleaned.loc[clear_mask, spec["change"]] = np.nan
        cleaned_change = pd.to_numeric(cleaned[spec["change"]], errors="coerce")
        negative_after = int(cleaned_change.lt(0).sum())

        summary_rows.append({
            "Assessment": spec["name"],
            "T1_Column": spec["t1"],
            "T2_Column": spec["t2"],
            "Change_Column": spec["change"],
            "NonNull_Change_Rows": int(valid_change.sum()),
            "Negative_Before_Cleanup": negative_before,
            "T2_Zero_Rows_With_Change": int((zero_t2 & valid_change).sum()),
            "Rows_Cleared_Zero_T2_Rule": int(clear_zero_t2.sum()),
            "Rows_Cleared_BI_LOS_Rule": int(bi_short_stay_clear.sum()),
            "Rows_Cleared_Out_Of_Range": int(clear_out_of_range.sum()),
            "Rows_Cleared_Total": int(clear_mask.sum()),
            "Negative_After_Cleanup": negative_after,
            "Rows_With_T2Zero_And_Later_NonZero": int((zero_t2 & valid_change & later_nonzero).sum()),
            "Note": spec["note"],
        })

        flagged_mask = zero_t2 & valid_change
        if flagged_mask.any():
            flagged = pd.DataFrame({
                "ID": cleaned.get("ID"),
                "Assessment": spec["name"],
                "T1": t1,
                "T2": t2,
                "T3": t3,
                "T4": t4,
                "Original_Change": original_change,
                "Cleaned_Change": cleaned_change,
                "Later_NonZero_T3_or_T4": later_nonzero.map({True: "Yes", False: "No"}),
                "Action": np.where(clear_zero_t2, "Cleared", "Flagged for review"),
                "Reason": np.where(
                    clear_zero_t2,
                    spec["note"],
                    "T2=0 with non-null change; review whether 0 is a true score or missing assessment.",
                ),
                "Rehab_LOS_Category": cleaned.get("Rehab_LOS_Category"),
                "PAC_Program_Completion": cleaned.get("PAC_Program_Completion"),
            }).loc[flagged_mask]
            flagged_rows.extend(flagged.to_dict("records"))

        if clear_out_of_range.any():
            flagged = pd.DataFrame({
                "ID": cleaned.get("ID"),
                "Assessment": spec["name"],
                "Issue": np.where(
                    t2_out_of_range,
                    f"{spec['t2']} outside valid range [{spec['valid_min']}, {spec['valid_max']}]",
                    f"{spec['change']} outside allowable ±{value_span}",
                ),
                "T1": t1,
                "T2": t2,
                "T3": t3,
                "T4": t4,
                "Original_Change": original_change,
                "Cleaned_Change": cleaned_change,
                "Rehab_LOS_Category": cleaned.get("Rehab_LOS_Category"),
                "PAC_Program_Completion": cleaned.get("PAC_Program_Completion"),
            }).loc[clear_out_of_range]
            illogical_rows.extend(flagged.to_dict("records"))

    bi_cases_df = pd.DataFrame(flagged_rows)
    if not bi_cases_df.empty:
        bi_cases_df = bi_cases_df[bi_cases_df["Assessment"].eq("BI")][[
            "ID", "T1", "T2", "T3", "T4", "Original_Change", "Cleaned_Change",
            "Rehab_LOS_Category", "PAC_Program_Completion", "Reason",
        ]].rename(columns={
            "T1": "BI1",
            "T2": "BI2",
            "T3": "BI3",
            "T4": "BI4",
            "Original_Change": "Original_BI_T1T2_Change",
            "Cleaned_Change": "Cleaned_BI_T1T2_Change",
        })

    audit = {
        "summary_df": pd.DataFrame(summary_rows),
        "flagged_cases_df": pd.DataFrame(flagged_rows),
        "illogical_values_df": pd.DataFrame(illogical_rows),
        "raw_outliers_df": pd.DataFrame(raw_outlier_rows),
        "bi_cases_df": bi_cases_df if not isinstance(bi_cases_df, list) else pd.DataFrame(),
    }
    return cleaned, audit


def fit_bootstrap_lasso(
    df: pd.DataFrame,
    features: list[str],
    model_name: str,
    model_seed: int = RANDOM_STATE,
    model5_restrict: bool = False,
) -> dict:
    """
    Fit a bootstrap LASSO model and return a results dictionary.

    Parameters
    ----------
    df : pd.DataFrame
        Full dataset.
    features : list[str]
        Candidate predictor columns (already filtered to those in df).
    model_name : str
        Human-readable model label.
    model_seed : int
        Random seed (offset per model for reproducibility).
    model5_restrict : bool
        If True, restrict analysis to patients with Rehab_LOS_Category in
        QUALIFYING_REHAB_LOS ('21-42 days' or '>42 days'); used for
        models that include T1T2_Change predictors.
    """
    if model5_restrict and "Rehab_LOS_Category" in df.columns:
        model_df = df[
            df["6MWT4"].notna()
            & df["Rehab_LOS_Category"].isin(QUALIFYING_REHAB_LOS)
        ].copy()
    else:
        model_df = df[df["6MWT4"].notna()].copy()

    X = model_df[features]
    y = model_df["6MWT4"].to_numpy(dtype=float)

    n_patients = len(model_df)
    n_input = len(features)

    base_pipe = _build_lasso_pipeline(features)
    base_pipe.fit(X, y)
    base_lasso: LassoCV = base_pipe.named_steps["model"]
    base_coef = base_lasso.coef_.copy()
    best_alpha = float(base_lasso.alpha_)

    y_train_pred = base_pipe.predict(X)
    train_r2 = float(r2_score(y, y_train_pred))
    train_mae = float(mean_absolute_error(y, y_train_pred))

    outer_cv = KFold(n_splits=CV_FOLDS, shuffle=True, random_state=RANDOM_STATE)
    y_oof = cross_val_predict(
        _build_lasso_pipeline(features), X, y,
        cv=outer_cv, n_jobs=-1,
    )
    cv_r2 = float(r2_score(y, y_oof))
    cv_mae = float(mean_absolute_error(y, y_oof))

    rng = np.random.default_rng(model_seed)
    coef_boot = np.zeros((N_BOOTSTRAP_LASSO, n_input), dtype=float)
    alpha_boot = np.zeros(N_BOOTSTRAP_LASSO, dtype=float)

    for b in range(N_BOOTSTRAP_LASSO):
        idx = rng.integers(0, n_patients, size=n_patients)
        Xb = X.iloc[idx]
        yb = y[idx]

        bp = _build_lasso_pipeline(features)
        bp.fit(Xb, yb)
        lb: LassoCV = bp.named_steps["model"]

        coef_boot[b, :] = lb.coef_
        alpha_boot[b] = lb.alpha_

    _COEF_TOL = 1e-9
    sel_freq = (np.abs(coef_boot) > _COEF_TOL).mean(axis=0)
    coef_mean = coef_boot.mean(axis=0)
    coef_sd = coef_boot.std(axis=0, ddof=1)

    importance = pd.DataFrame({
        "predictor": features,
        "base_coef": base_coef,
        "bootstrap_coef_mean": coef_mean,
        "bootstrap_coef_sd": coef_sd,
        "selection_frequency": sel_freq,
        "abs_bootstrap_coef_mean": np.abs(coef_mean),
    }).sort_values(
        ["selection_frequency", "abs_bootstrap_coef_mean"],
        ascending=[False, False],
    ).reset_index(drop=True)

    n_final_nonzero = int((np.abs(base_coef) > 1e-9).sum())
    n_stable = int((sel_freq >= STABILITY_THRESHOLD).sum())
    category_counts = _predictor_category_counts(features)

    return {
        "model_name": model_name,
        "features": features,
        "n_patients": n_patients,
        "n_input_variables": n_input,
        "n_final_nonzero_base": n_final_nonzero,
        "n_stable_selected": n_stable,
        "best_alpha_full_fit": best_alpha,
        "alpha_boot_median": float(np.median(alpha_boot)),
        "alpha_boot_q25": float(np.quantile(alpha_boot, 0.25)),
        "alpha_boot_q75": float(np.quantile(alpha_boot, 0.75)),
        "cv_r2": cv_r2,
        "cv_mae": cv_mae,
        "train_r2": train_r2,
        "train_mae": train_mae,
        "base_pipe": base_pipe,
        "importance": importance,
        "category_counts": category_counts,
        "gait_speed_imputed_status": (
            "with gait speed imputed"
            if "Gait_Speed_1_Imputed" in features
            else "without gait speed imputed"
        ),
        "predictor_category_summary": _predictor_category_summary_text(features),
    }


# ─────────────────────────────────────────────────────────────────────────────
# IPCW PIPELINE – SINGLE COMPLETION MODEL
# ─────────────────────────────────────────────────────────────────────────────

def compute_ipcw_weights_single_model(
    df: pd.DataFrame,
    completion_features: list[str],
    winsor_pct: float = WINSOR_PCT,
    numerator_mode: str = NUMERATOR_MODE,
) -> dict:
    """
    Compute stabilized IPCW weights using ONE completion model applied
    uniformly to all outcome-tier regressions.

    The denominator model is fitted on the richest clinically justified
    predictor set (IPCW_COMPLETION_FEATURES).  The same weights are later
    passed to each tier's Ridge regression, consistent with conventional
    IPCW practice where the missingness/completion model is specified once.

    Parameters
    ----------
    df : full patient dataset
    completion_features : list of columns used in the denominator model
        (pre-filtered to those present in df)
    winsor_pct : percentile threshold for weight winsorization
    numerator_mode : 'intercept' (marginal completion rate, default)

    Returns
    -------
    dict with weights, diagnostics, sensitivity tables
    """
    valid = _filter_existing(completion_features, df)
    if not valid:
        raise ValueError("No valid completion features found.")

    completion_status = df["PAC_Program_Completion"].astype("string")
    eligible_mask = completion_status.notna()
    completed = completion_status.eq("Completed PAC program").astype(int)

    X = df.loc[eligible_mask, valid]
    y = completed.loc[eligible_mask].to_numpy()

    imputer = SimpleImputer(strategy="median")
    X_imp = pd.DataFrame(imputer.fit_transform(X), columns=valid, index=X.index)

    # Denominator model: P(completion | richest feature set)
    denom_pipe = Pipeline([
        ("sc", StandardScaler()),
        ("lr", LogisticRegression(max_iter=5000, solver="lbfgs", C=1.0)),
    ])
    denom_pipe.fit(X_imp, y)
    p_denom_raw = denom_pipe.predict_proba(X_imp)[:, 1]

    near_zero = float(np.mean(p_denom_raw < 0.05))
    near_one = float(np.mean(p_denom_raw > 0.95))
    p_denom_diag = {
        "mean": float(np.mean(p_denom_raw)),
        "sd": float(np.std(p_denom_raw)),
        "min": float(np.min(p_denom_raw)),
        "p10": float(np.percentile(p_denom_raw, 10)),
        "p25": float(np.percentile(p_denom_raw, 25)),
        "p50": float(np.median(p_denom_raw)),
        "p75": float(np.percentile(p_denom_raw, 75)),
        "p90": float(np.percentile(p_denom_raw, 90)),
        "max": float(np.max(p_denom_raw)),
        "pct_near_zero": near_zero,
        "pct_near_one": near_one,
        "positivity_flag": near_zero > 0.05 or near_one > 0.05,
    }
    p_denom = np.clip(p_denom_raw, 1e-4, 1 - 1e-4)

    # Numerator (intercept-only: marginal completion rate)
    p_numer = np.full(len(y), y.mean(), dtype="float64")
    numer_description = "Intercept-only (marginal completion rate)"

    raw_weights = np.where(y == 1, p_numer / p_denom, np.nan)
    completer_weights = raw_weights[~np.isnan(raw_weights)]

    sensitivity: dict[int, dict] = {}
    all_threshold_pcts = sorted(set(WINSORIZATION_THRESHOLDS) | {winsor_pct})
    for pct in all_threshold_pcts:
        lo, hi = np.nanpercentile(completer_weights, [100 - pct, pct])
        w_clipped = np.clip(completer_weights, lo, hi)
        pct_truncated = float(np.mean((completer_weights < lo) | (completer_weights > hi)))
        sensitivity[pct] = {
            "lower": float(lo),
            "upper": float(hi),
            "mean": float(np.mean(w_clipped)),
            "sd": float(np.std(w_clipped)),
            "min": float(np.min(w_clipped)),
            "max": float(np.max(w_clipped)),
            "pct_truncated": pct_truncated,
        }

    lo_main = sensitivity[winsor_pct]["lower"]
    hi_main = sensitivity[winsor_pct]["upper"]
    winsorized = np.where(
        ~np.isnan(raw_weights),
        np.clip(raw_weights, lo_main, hi_main),
        np.nan,
    )

    weights_series = pd.Series(np.nan, index=df.index, dtype="float64")
    weights_series.loc[X_imp.index] = winsorized

    raw_weights_series = pd.Series(np.nan, index=df.index, dtype="float64")
    raw_weights_series.loc[X_imp.index] = raw_weights

    wt = winsorized[~np.isnan(winsorized)]
    weight_diag = {
        "mean": float(np.mean(wt)),
        "sd": float(np.std(wt)),
        "min": float(np.min(wt)),
        "max": float(np.max(wt)),
        "pct_truncated": sensitivity[winsor_pct]["pct_truncated"],
        "mean_near_one": abs(float(np.mean(wt)) - 1.0) < 0.2,
    }

    return {
        "weights": weights_series,
        "raw_weights_series": raw_weights_series,
        "raw_completer_weights": completer_weights,
        "imputer": imputer,
        "features": valid,
        "n_features_completion_model": len(valid),
        "denominator_model": (
            f"LogisticRegression(solver='lbfgs', C=1.0) on richest feature set "
            f"({len(valid)} variables: demographics + stroke info + comorbidities + "
            "acute complications + NIHSS out + functional T1)"
        ),
        "numerator_model": numer_description,
        "winsorization": f"{100 - winsor_pct}th/{winsor_pct}th percentile",
        "p_denom_diagnostics": p_denom_diag,
        "weight_diagnostics": weight_diag,
        "sensitivity": sensitivity,
        "y_eligible": y,
        "X_imp": X_imp,
    }


# ─────────────────────────────────────────────────────────────────────────────
# BINARY CLASSIFIER (walk / no-walk for non-completers)
# ─────────────────────────────────────────────────────────────────────────────

def _binary_candidates(features: list[str]) -> dict[str, Pipeline]:
    numeric_pipe = Pipeline([("imp", SimpleImputer(strategy="median")), ("sc", StandardScaler())])
    imp_only = Pipeline([("imp", SimpleImputer(strategy="median"))])

    logistic = Pipeline([
        ("prep", ColumnTransformer([("num", numeric_pipe, features)])),
        ("model", LogisticRegression(max_iter=5000, solver="liblinear", class_weight="balanced", random_state=RANDOM_STATE)),
    ])
    rf = Pipeline([
        ("prep", ColumnTransformer([("num", imp_only, features)])),
        ("model", RandomForestClassifier(n_estimators=300, class_weight="balanced", random_state=RANDOM_STATE, n_jobs=-1)),
    ])
    et = Pipeline([
        ("prep", ColumnTransformer([("num", imp_only, features)])),
        ("model", ExtraTreesClassifier(n_estimators=300, class_weight="balanced", random_state=RANDOM_STATE, n_jobs=-1)),
    ])
    return {
        "LogisticRegression(class_weight='balanced', solver='liblinear')": logistic,
        "RandomForestClassifier(n_estimators=300, class_weight='balanced')": rf,
        "ExtraTreesClassifier(n_estimators=300, class_weight='balanced')": et,
    }


def select_binary_classifier(
    df: pd.DataFrame,
    features: list[str],
    scenario: str,
    forced_model_name: str | None = None,
) -> dict:
    valid = _filter_existing(features, df)
    model_df = df[df[scenario].notna()].copy()
    X = model_df[valid]
    y = model_df[scenario].astype(int)
    cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)
    cv_splits = list(cv.split(X, y))
    candidates = _binary_candidates(valid)

    rows = []
    for name, pipe in candidates.items():
        scores = cross_validate(clone(pipe), X, y, cv=cv_splits, n_jobs=-1,
                                scoring=["balanced_accuracy", "accuracy", "f1"])
        rows.append({
            "model": name,
            "balanced_accuracy": float(np.mean(scores["test_balanced_accuracy"])),
            "accuracy": float(np.mean(scores["test_accuracy"])),
            "f1": float(np.mean(scores["test_f1"])),
        })

    if forced_model_name is not None:
        if forced_model_name not in candidates:
            raise ValueError(f"Forced binary model not found: {forced_model_name}")
        best = next(r for r in rows if r["model"] == forced_model_name)
    else:
        best = max(rows, key=lambda r: (r["balanced_accuracy"], r["accuracy"], r["f1"]))
    best_pipe = clone(candidates[best["model"]])
    best_pipe.fit(X, y)
    return {"best_model_name": best["model"], "best_pipeline": best_pipe, "features": valid,
            "leaderboard": rows, "bal_acc": best["balanced_accuracy"]}


# ─────────────────────────────────────────────────────────────────────────────
# WEIGHTED RIDGE REGRESSION
# ─────────────────────────────────────────────────────────────────────────────

def fit_weighted_ridge(
    df: pd.DataFrame,
    features: list[str],
    ipcw: dict,
    binary: dict,
) -> dict:
    """
    Fit weighted Ridge regression using pre-computed IPCW weights.

    The weights passed in are the SAME across all retained IPCW models (single completion
    model).  Only the Ridge regression features vary per tier.
    """
    valid = _filter_existing(features, df)
    completion_status = df["PAC_Program_Completion"].astype("string")
    completer_mask = completion_status.eq("Completed PAC program") & df["6MWT4"].notna()
    df_comp = df.loc[completer_mask].copy()
    weights = ipcw["weights"].loc[completer_mask].fillna(1.0).to_numpy()

    X_raw = df_comp[valid]
    y = df_comp["6MWT4"].to_numpy()
    cv_splits = list(KFold(n_splits=CV_FOLDS, shuffle=True, random_state=RANDOM_STATE).split(X_raw))
    oof = np.zeros(len(df_comp), dtype="float64")

    for tr, te in cv_splits:
        imp = SimpleImputer(strategy="median")
        X_tr = imp.fit_transform(X_raw.iloc[tr])
        X_te = imp.transform(X_raw.iloc[te])
        m = Ridge(alpha=1.0)
        m.fit(X_tr, y[tr], sample_weight=weights[tr])
        oof[te] = np.maximum(0, m.predict(X_te))

    r2 = float(r2_score(y, oof, sample_weight=weights))
    mae = float(np.average(np.abs(y - oof), weights=weights))

    # Sensitivity at different winsor thresholds
    sensitivity_r2: dict[int, float] = {}
    sensitivity_mae: dict[int, float] = {}
    raw_w = ipcw["raw_weights_series"].loc[completer_mask].fillna(1.0).to_numpy()
    for pct in WINSORIZATION_THRESHOLDS:
        lo, hi = np.percentile(raw_w[np.isfinite(raw_w)], [100 - pct, pct])
        w_alt = np.clip(raw_w, lo, hi)
        oof_alt = np.zeros(len(df_comp), dtype="float64")
        for tr, te in cv_splits:
            imp = SimpleImputer(strategy="median")
            X_tr = imp.fit_transform(X_raw.iloc[tr])
            X_te = imp.transform(X_raw.iloc[te])
            m = Ridge(alpha=1.0)
            m.fit(X_tr, y[tr], sample_weight=w_alt[tr])
            oof_alt[te] = np.maximum(0, m.predict(X_te))
        sensitivity_r2[pct] = float(r2_score(y, oof_alt, sample_weight=w_alt))
        sensitivity_mae[pct] = float(np.average(np.abs(y - oof_alt), weights=w_alt))

    imp_final = SimpleImputer(strategy="median")
    X_all = pd.DataFrame(imp_final.fit_transform(X_raw), columns=valid, index=df_comp.index)
    final_model = Ridge(alpha=1.0)
    final_model.fit(X_all, y, sample_weight=weights)

    noncompleter_mask = completion_status.ne("Completed PAC program") & completion_status.notna()
    df_nc = df.loc[noncompleter_mask].copy()
    X_nc_raw = df_nc[valid]
    X_nc = pd.DataFrame(imp_final.transform(X_nc_raw), columns=valid, index=df_nc.index)
    walk_pred = binary["best_pipeline"].predict(X_nc_raw).astype(int)
    pred_6mwt4 = np.zeros(len(df_nc), dtype="float64")
    if walk_pred.any():
        pred_6mwt4[walk_pred == 1] = np.maximum(0, final_model.predict(X_nc[walk_pred == 1]))

    importance = (
        pd.DataFrame({"predictor": valid, "abs_coef": np.abs(final_model.coef_)})
        .sort_values("abs_coef", ascending=False).reset_index(drop=True)
    )

    return {
        "weighted_r2": r2,
        "weighted_mae": mae,
        "sensitivity_r2": sensitivity_r2,
        "sensitivity_mae": sensitivity_mae,
        "model_name": "Ridge(alpha=1.0) with single-completion-model IPCW weights",
        "importance": importance,
        "final_model": final_model,
        "imputer": imp_final,
        "noncomp_index": df_nc.index.tolist(),
        "noncomp_walk_pred": walk_pred,
        "noncomp_6mwt4_pred": pred_6mwt4,
        "n_noncomp_pred_walk": int(walk_pred.sum()),
        "n_noncomp_pred_no_walk": int((walk_pred == 0).sum()),
    }


# ─────────────────────────────────────────────────────────────────────────────
# BOOTSTRAP UQ FOR IPCW PIPELINE
# ─────────────────────────────────────────────────────────────────────────────

def _bootstrap_ipcw_once(
    df: pd.DataFrame,
    completion_features: list[str],
    tier_features: list[str],
    scenario: str,
    rng: np.random.Generator,
) -> tuple[float, float]:
    """
    One bootstrap resample: re-fit the single completion model, re-derive
    weights, refit Ridge for the given tier.  Stratified by completion status.
    """
    completion_status = df["PAC_Program_Completion"].astype("string")
    groups = completion_status.fillna("Missing")
    idx_resampled: list[np.ndarray] = []
    for grp_val in groups.unique():
        grp_idx = np.where(groups == grp_val)[0]
        idx_resampled.append(rng.choice(grp_idx, size=len(grp_idx), replace=True))
    boot_idx = np.concatenate(idx_resampled)
    df_boot = df.iloc[boot_idx].reset_index(drop=True)

    try:
        ipcw_boot = compute_ipcw_weights_single_model(df_boot, completion_features)
    except Exception:
        return float("nan"), float("nan")

    try:
        valid = _filter_existing(tier_features, df_boot)
        model_df = df_boot[df_boot[scenario].notna()].copy()
        X_bin = model_df[valid]
        y_bin = model_df[scenario].astype(int)
        bin_pipe = Pipeline([
            ("prep", ColumnTransformer([
                ("num", Pipeline([("imp", SimpleImputer(strategy="median")),
                                  ("sc", StandardScaler())]), valid),
            ])),
            ("model", LogisticRegression(max_iter=1000, solver="liblinear",
                                         class_weight="balanced", random_state=RANDOM_STATE)),
        ])
        bin_pipe.fit(X_bin, y_bin)
    except Exception:
        return float("nan"), float("nan")

    try:
        comp_mask = (df_boot["PAC_Program_Completion"].astype("string").eq("Completed PAC program")
                     & df_boot["6MWT4"].notna())
        df_comp = df_boot.loc[comp_mask].copy()
        if len(df_comp) < 10:
            return float("nan"), float("nan")
        weights = ipcw_boot["weights"].loc[comp_mask].fillna(1.0).to_numpy()
        X_raw = df_comp[valid]
        y = df_comp["6MWT4"].to_numpy()
        n_splits = min(5, max(2, len(df_comp) // 10))
        cv_splits = list(KFold(n_splits=n_splits, shuffle=True,
                               random_state=RANDOM_STATE).split(X_raw))
        oof = np.zeros(len(df_comp), dtype="float64")
        for tr, te in cv_splits:
            imp = SimpleImputer(strategy="median")
            X_tr = imp.fit_transform(X_raw.iloc[tr])
            X_te = imp.transform(X_raw.iloc[te])
            m = Ridge(alpha=1.0)
            m.fit(X_tr, y[tr], sample_weight=weights[tr])
            oof[te] = np.maximum(0, m.predict(X_te))
        r2 = float(r2_score(y, oof, sample_weight=weights))
        mae = float(np.average(np.abs(y - oof), weights=weights))
        return r2, mae
    except Exception:
        return float("nan"), float("nan")


def bootstrap_ci_ipcw(
    df: pd.DataFrame,
    completion_features: list[str],
    tier_features: list[str],
    scenario: str,
    n_bootstrap: int = N_BOOTSTRAP_IPCW,
    seed: int = RANDOM_STATE,
) -> dict:
    rng = np.random.default_rng(seed)
    r2_boot: list[float] = []
    mae_boot: list[float] = []

    for _ in range(n_bootstrap):
        r2_b, mae_b = _bootstrap_ipcw_once(df, completion_features, tier_features, scenario, rng)
        if math.isfinite(r2_b):
            r2_boot.append(r2_b)
        if math.isfinite(mae_b):
            mae_boot.append(mae_b)

    def _ci(vals: list[float]) -> tuple[float, float, float]:
        if not vals:
            return float("nan"), float("nan"), float("nan")
        a = np.array(vals)
        return float(np.mean(a)), float(np.percentile(a, 2.5)), float(np.percentile(a, 97.5))

    r2_mean, r2_lo, r2_hi = _ci(r2_boot)
    mae_mean, mae_lo, mae_hi = _ci(mae_boot)
    return {
        "n_valid_r2": len(r2_boot),
        "n_valid_mae": len(mae_boot),
        "r2_boot_mean": r2_mean, "r2_ci_lo": r2_lo, "r2_ci_hi": r2_hi,
        "mae_boot_mean": mae_mean, "mae_ci_lo": mae_lo, "mae_ci_hi": mae_hi,
    }


# ─────────────────────────────────────────────────────────────────────────────
# DOCX REPORT
# ─────────────────────────────────────────────────────────────────────────────

def _model_rankings(results: list[dict]) -> dict[str, int]:
    ranked = sorted(results, key=lambda x: x["cv_r2"], reverse=True)
    return {row["model_name"]: idx for idx, row in enumerate(ranked, start=1)}


def _executive_bullets(results: list[dict]) -> list[str]:
    ranked = sorted(results, key=lambda x: x["cv_r2"], reverse=True)
    best = ranked[0]
    second = ranked[1]
    compact = min(results, key=lambda x: x["n_input_variables"])
    model_lookup = {r["model_name"].split(":", 1)[0]: r for r in results}
    model5 = model_lookup.get("Model 5")
    model2 = model_lookup.get("Model 2")
    model8 = model_lookup.get("Model 8")
    model1 = model_lookup.get("Model 1")
    model4 = model_lookup.get("Model 4")
    return [
        (
            f"Best overall out-of-fold fit was achieved by {best['model_name']} "
            f"(CV R² {best['cv_r2']:.4f}, MAE {best['cv_mae']:.2f} m), narrowly ahead of "
            f"{second['model_name']}."
        ),
        (
            f"Retained Step 1 models in this update are ordered as: "
            f"Model 5, Model 2, Model 8, Model 1, Model 4."
        ),
        (
            f"Model 5 (recovery-trajectory) was restricted to {model5['n_patients']:,} patients with "
            f"Rehab LOS ≥21 days (Rehab_LOS_Category = '21-42 days' or '>42 days'). "
            f"CV R² = {model5['cv_r2']:.4f}, MAE = {model5['cv_mae']:.2f} m."
        ),
        (
            f"Among retained models, CV R² values were: "
            f"Model 2={model2['cv_r2']:.4f}, Model 8={model8['cv_r2']:.4f}, "
            f"Model 1={model1['cv_r2']:.4f}, Model 4={model4['cv_r2']:.4f}."
        ),
        (
            f"The most compact model was {compact['model_name']} with {compact['n_input_variables']} predictors "
            f"while retaining CV R² {compact['cv_r2']:.4f}."
        ),
    ]


def write_combined_report(
    lasso_results: list[dict],
    ipcw_completion_info: dict,
    ipcw_tier_results: list[dict],
    audit_info: dict,
) -> None:
    doc = Document()

    title = doc.add_paragraph()
    title.add_run(
        "Comprehensive LASSO + IPCW Report for 6MWT4 Prediction – 2026-09-04"
    ).bold = True

    intro = doc.add_paragraph(
        "This report combines two analyses:\n"
        "(1) Bootstrap LASSO regression models (no IPCW) predicting 6MWT4 as a "
        "continuous outcome — non-redundant model set aligned to clinically distinct "
        "predictor groups, with updated Model 5 restricted to "
        "patients with sufficient rehabilitation exposure (Rehab LOS ≥21 days).\n"
        "(2) IPCW-weighted Ridge regression across retained Step 1 models, using a "
        "single completion model (richest predictor set) to derive stabilized weights "
        "applied uniformly across all tiers — consistent with conventional IPCW "
        "practice."
    )
    _small_para(intro, 9)

    doc.add_paragraph()
    audit_heading = doc.add_paragraph()
    audit_heading.add_run("T1T2 Data Quality Audit and Cleanup").bold = True

    audit_summary = audit_info["summary_df"]
    bi_cases = audit_info["bi_cases_df"]
    illogical_values = audit_info["illogical_values_df"]
    raw_outliers = audit_info["raw_outliers_df"]
    manual_corrections = audit_info.get("manual_corrections_df", pd.DataFrame())
    fuglsen_outliers_pre = audit_info.get("fuglsen_outliers_pre_df", pd.DataFrame())

    bi_row = (
        audit_summary.loc[audit_summary["Assessment"].eq("BI")].iloc[0]
        if not audit_summary.empty and audit_summary["Assessment"].eq("BI").any()
        else None
    )
    if bi_row is not None:
        audit_intro = doc.add_paragraph(
            f"BI_T1T2_Change had {int(bi_row['Negative_Before_Cleanup'])} negative values before cleanup. "
            f"{int(bi_row['Rows_Cleared_BI_LOS_Rule'])} rows were blanked by the BI LOS<=21 rule, and "
            f"{int(bi_row['Rows_Cleared_Zero_T2_Rule'])} additional rows were blanked for BI2=0 in the 21-42 day LOS group. "
            "Remaining negative BI change values therefore reflect observed deterioration rather than missing-T2 artifacts."
        )
        _small_para(audit_intro, 9)

    summary_rows = [[
        "Assessment", "Change rows", "Negative before", "T2=0 rows",
        "Cleared LOS rule", "Cleared zero-T2", "Cleared out-of-range", "Negative after",
    ]]
    for _, row in audit_summary.iterrows():
        summary_rows.append([
            row["Assessment"],
            int(row["NonNull_Change_Rows"]),
            int(row["Negative_Before_Cleanup"]),
            int(row["T2_Zero_Rows_With_Change"]),
            int(row.get("Rows_Cleared_BI_LOS_Rule", 0)),
            int(row["Rows_Cleared_Zero_T2_Rule"]),
            int(row["Rows_Cleared_Out_Of_Range"]),
            int(row["Negative_After_Cleanup"]),
        ])
    if len(summary_rows) > 1:
        _add_table_to_doc(doc, summary_rows)

    if not illogical_values.empty:
        doc.add_paragraph()
        illogical_heading = doc.add_paragraph()
        illogical_heading.add_run("Illogical values cleared from T1T2 change fields").bold = True
        illogical_rows = [[
            "ID", "Assessment", "Issue", "T1", "T2", "T3", "T4", "Original change",
        ]]
        for _, row in illogical_values.iterrows():
            illogical_rows.append([
                int(row["ID"]) if pd.notna(row["ID"]) else "",
                row["Assessment"],
                row["Issue"],
                _fmt(row["T1"], 1),
                _fmt(row["T2"], 1),
                _fmt(row["T3"], 1),
                _fmt(row["T4"], 1),
                _fmt(row["Original_Change"], 1),
            ])
        _add_table_to_doc(doc, illogical_rows)

    if not raw_outliers.empty:
        doc.add_paragraph()
        raw_heading = doc.add_paragraph()
        raw_heading.add_run("Other functional raw-score outliers found during audit").bold = True
        raw_rows = [[
            "ID", "Assessment", "Visit column", "Observed value", "Allowed range", "Affects T1T2 cleanup",
        ]]
        for _, row in raw_outliers.iterrows():
            raw_rows.append([
                int(row["ID"]) if pd.notna(row["ID"]) else "",
                row["Assessment"],
                row["Visit_Column"],
                _fmt(row["Observed_Value"], 1),
                row["Allowed_Range"],
                row["Affects_T1T2_Change_Cleanup"],
            ])
        _add_table_to_doc(doc, raw_rows)

    if not bi_cases.empty:
        doc.add_paragraph()
        bi_heading = doc.add_paragraph()
        bi_heading.add_run("BI2 = 0 cases (no BI T2 assessment; change blanked)").bold = True
        bi_rows = [[
            "ID", "BI1", "BI2", "BI3", "BI4", "Original BI change",
            "Rehab LOS", "PAC completion",
        ]]
        for _, row in bi_cases.iterrows():
            bi_rows.append([
                int(row["ID"]) if pd.notna(row["ID"]) else "",
                _fmt(row["BI1"], 1),
                _fmt(row["BI2"], 1),
                _fmt(row["BI3"], 1),
                _fmt(row["BI4"], 1),
                _fmt(row["Original_BI_T1T2_Change"], 1),
                row["Rehab_LOS_Category"],
                row["PAC_Program_Completion"],
            ])
        _add_table_to_doc(doc, bi_rows)

    if not fuglsen_outliers_pre.empty:
        doc.add_paragraph()
        fh = doc.add_paragraph()
        fh.add_run("FuglSEN T1-T4 out-of-range values identified (>44)").bold = True
        fr = [["ID", "Visit column", "Observed value", "Allowed max", "Rehab LOS"]]
        for _, row in fuglsen_outliers_pre.iterrows():
            fr.append([
                int(row["ID"]) if pd.notna(row["ID"]) else "",
                row["Visit_Column"],
                _fmt(row["Observed_Value"], 1),
                _fmt(row["Allowed_Max"], 1),
                row["Rehab_LOS_Category"],
            ])
        _add_table_to_doc(doc, fr)

    # ── PART 1: LASSO MODELS ─────────────────────────────────────────────────
    doc.add_paragraph()
    p1h = doc.add_paragraph()
    p1h.add_run("PART 1: Bootstrap LASSO Models (No IPCW)").bold = True

    intro2 = doc.add_paragraph(
        f"{len(lasso_results)} bootstrap LASSO models predicting 6MWT4. "
        f"Missing predictor values are handled by median imputation; predictors are "
        f"z-score standardised. Regularisation alpha is chosen by {CV_FOLDS}-fold LassoCV. "
        f"True bootstrap ({N_BOOTSTRAP_LASSO:,} resamples) assesses coefficient stability. "
        "Performance is reported both on the full training set (apparent) and via "
        f"{CV_FOLDS}-fold out-of-fold (OOF) cross-validation.\n\n"
        "T1T2-model update (2026-09-04): T1T2_Change variables are only meaningful for "
        "patients with sufficient rehabilitation exposure. Any LASSO model using T1T2 predictors is therefore "
        "restricted to patients whose Rehab_LOS_Category is '21-42 days' or '>42 days'. "
        "Patients outside these groups are excluded from the Model 5 analysis entirely — "
        "they are not counted, and their T1T2_Change values are not imputed. "
        "In addition, BI_T1T2_Change is blanked for LOS<=21 days and for BI2=0 within 21-42 days LOS, "
        "FOIS_T1T2_Change is blanked whenever FOIS2=0, "
        "and impossible change values linked to out-of-range T2 data are removed before modeling. "
        "This ensures that change scores reflect a genuine rehabilitation-period change "
        "rather than a mix of short-stay noise and missing values."
    )
    _small_para(intro2, 9)

    doc.add_paragraph()
    comp_heading = doc.add_paragraph()
    comp_heading.add_run("Executive comparison across all Step 1 LASSO models").bold = True
    comparison_rows = [["Rank", "Model", "N", "Predictors", "GS status", "CV R²", "CV MAE (m)"]]
    for rank, r in enumerate(sorted(lasso_results, key=lambda x: x["cv_r2"], reverse=True), start=1):
        comparison_rows.append([
            rank,
            r["model_name"].split(":", 1)[0],
            r["n_patients"],
            r["n_input_variables"],
            r["gait_speed_imputed_status"],
            f"{r['cv_r2']:.4f}",
            f"{r['cv_mae']:.2f}",
        ])
    _add_table_to_doc(doc, comparison_rows)

    for bullet in _executive_bullets(lasso_results):
        para = doc.add_paragraph(f"• {bullet}")
        _small_para(para, 9)

    ranks = _model_rankings(lasso_results)

    for r in lasso_results:
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run(r["model_name"]).bold = True

        feat_preview = ", ".join(r["features"][:15])
        if len(r["features"]) > 15:
            feat_preview += f" … (+{len(r['features']) - 15} more)"
        pred_para = doc.add_paragraph(f"Predictors ({len(r['features'])}): {feat_preview}")
        _small_para(pred_para, 8)
        cat_para = doc.add_paragraph(
            f"Predictor categories: {r['predictor_category_summary']}"
        )
        _small_para(cat_para, 8)

        _add_table_to_doc(doc, [
            ["Metric", "Value"],
            ["Patients (non-missing 6MWT4, qualifying LOS for Model 5)", r["n_patients"]],
            ["Input variable count", r["n_input_variables"]],
            ["Predictor category summary", r["predictor_category_summary"]],
            ["Final selected (non-zero in full-fit)", r["n_final_nonzero_base"]],
            [
                f"Stably selected (selection freq ≥ {STABILITY_THRESHOLD:.0%})",
                r["n_stable_selected"],
            ],
            ["Best alpha (full-data fit)", f"{r['best_alpha_full_fit']:.6g}"],
            [
                "Bootstrap alpha: median [IQR]",
                (
                    f"{r['alpha_boot_median']:.6g}"
                    f" [{r['alpha_boot_q25']:.6g}, {r['alpha_boot_q75']:.6g}]"
                ),
            ],
            ["CV R² (OOF)", f"{r['cv_r2']:.4f}"],
            ["CV MAE (OOF)", f"{r['cv_mae']:.2f} m"],
            ["Train R² (apparent)", f"{r['train_r2']:.4f}"],
            ["Train MAE (apparent)", f"{r['train_mae']:.2f} m"],
        ])

        top_positive = _top_predictors(r["importance"], positive=True)
        top_negative = _top_predictors(r["importance"], positive=False)
        explainer = doc.add_paragraph(
            "Explainer: "
            f"{MODEL_EXPLAINERS[r['model_name']]} "
            f"It ranked #{ranks[r['model_name']]} of {len(lasso_results)} by CV R². "
            f"The apparent-to-cross-validated R² gap was {r['train_r2'] - r['cv_r2']:.4f}, "
            f"with {r['n_stable_selected']} predictors meeting the {STABILITY_THRESHOLD:.0%} stability threshold."
        )
        _small_para(explainer, 9)

        pos_para = doc.add_paragraph(
            "Strongest positive contributors: "
            f"{_format_predictor_summary(top_positive)}."
        )
        _small_para(pos_para, 8)
        neg_para = doc.add_paragraph(
            "Strongest negative contributors: "
            f"{_format_predictor_summary(top_negative)}."
        )
        _small_para(neg_para, 8)

        doc.add_paragraph()
        imp_heading = doc.add_paragraph()
        imp_heading.add_run(
            "Predictor importance – top 20 (sorted by bootstrap selection frequency)"
        ).bold = True

        imp = r["importance"].head(20)
        rows = [[
            "Predictor", "Direction", "Full-fit Coef",
            "Boot Mean Coef", "Boot SD", "Selection Freq", "Stable",
        ]]
        for _, row in imp.iterrows():
            rows.append([
                row["predictor"],
                "Positive" if row["bootstrap_coef_mean"] >= 0 else "Negative",
                f"{row['base_coef']:.5f}",
                f"{row['bootstrap_coef_mean']:.5f}",
                f"{row['bootstrap_coef_sd']:.5f}",
                f"{row['selection_frequency']:.3f}",
                "Yes" if row["selection_frequency"] >= STABILITY_THRESHOLD else "No",
            ])
        _add_table_to_doc(doc, rows)

    # ── PART 2: IPCW WITH SINGLE COMPLETION MODEL ────────────────────────────
    doc.add_paragraph()
    p2h = doc.add_paragraph()
    p2h.add_run("PART 2: IPCW-Weighted Ridge Regression (Single Completion Model)").bold = True

    ipcw_intro = doc.add_paragraph(
        "This section estimates 6MWT4 for completers and non-completers using "
        "Inverse Probability of Completion Weighting (IPCW).  "
        "A key methodological improvement over the prior session is the use of a "
        "SINGLE completion model, fitted once on the richest clinically justified "
        "predictor set, whose stabilized weights are then applied uniformly across "
        "all retained-model Ridge regressions.\n\n"
        "Rationale for single completion model (per conventional IPCW practice):\n"
        "• Re-deriving a separate completion model per tier conflates 'tier effect on "
        "outcome' with 'tier effect on weighting', making it difficult to isolate "
        "tier-specific predictive value from tier-specific selection bias.\n"
        "• A single well-specified completion model (using all available clinical "
        "information) provides the most stable and efficient weight estimates.\n"
        "• Applying the same weights uniformly across outcome models is consistent "
        "with the standard causal inference framing, where the weighting model "
        "addresses the selection mechanism once, independently of the outcome model.\n\n"
        f"Completion model predictor set ({ipcw_completion_info['n_features_completion_model']} features): "
        f"{ipcw_completion_info['denominator_model']}.\n"
        f"Numerator: {ipcw_completion_info['numerator_model']}.\n"
        f"Winsorization: {ipcw_completion_info['winsorization']}.\n"
        f"Bootstrap: B={N_BOOTSTRAP_IPCW}, stratified by completion status."
    )
    _small_para(ipcw_intro, 9)

    # Completion model diagnostics
    doc.add_paragraph()
    pd_diag = ipcw_completion_info["p_denom_diagnostics"]
    wd = ipcw_completion_info["weight_diagnostics"]

    dh = doc.add_paragraph()
    dh.add_run("Completion Model Positivity Diagnostics").bold = True

    flag_text = (
        "⚠ POSITIVITY FLAG: >5% of p_denom values near 0 or 1."
        if pd_diag["positivity_flag"]
        else "No positivity flag (< 5% near 0 or 1)."
    )
    _add_table_to_doc(doc, [
        ["p_denom statistic", "Value"],
        ["Mean", _fmt(pd_diag["mean"], 4)],
        ["SD", _fmt(pd_diag["sd"], 4)],
        ["Min", _fmt(pd_diag["min"], 4)],
        ["P10", _fmt(pd_diag["p10"], 4)],
        ["Median", _fmt(pd_diag["p50"], 4)],
        ["P90", _fmt(pd_diag["p90"], 4)],
        ["Max", _fmt(pd_diag["max"], 4)],
        ["% near 0 (<0.05)", _fmt(pd_diag["pct_near_zero"] * 100, 1) + "%"],
        ["% near 1 (>0.95)", _fmt(pd_diag["pct_near_one"] * 100, 1) + "%"],
        ["Positivity assessment", flag_text],
    ])

    doc.add_paragraph()
    wdh = doc.add_paragraph()
    wdh.add_run("Stabilized Weight Distribution").bold = True
    mean_note = (
        " (mean near 1 — reasonable numerator/denominator pairing)"
        if wd["mean_near_one"]
        else " ⚠ mean not near 1 — review specification"
    )
    _add_table_to_doc(doc, [
        ["Weight statistic", "Value"],
        ["Mean", _fmt(wd["mean"], 4) + mean_note],
        ["SD", _fmt(wd["sd"], 4)],
        ["Min", _fmt(wd["min"], 4)],
        ["Max", _fmt(wd["max"], 4)],
        ["% truncated at winsor threshold", _fmt(wd["pct_truncated"] * 100, 1) + "%"],
    ])

    # Per tier/scenario results
    for res in ipcw_tier_results:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(f"{res['model']} | {res['scenario_label']} Scenario").bold = True

        reg = res["regression"]
        ci = res["bootstrap_ci"]

        _add_table_to_doc(doc, [
            ["Field", "Value"],
            ["Model", res["model"]],
            ["Scenario", res["scenario_label"]],
            ["Completion model", "(shared single model — see above)"],
            ["Selected binary model", res["binary"]["best_model_name"]],
            ["Binary OOF balanced accuracy", _fmt(res["binary"]["bal_acc"], 3)],
            ["Weighted regression model", reg["model_name"]],
            ["Weighted OOF R²", _fmt(reg["weighted_r2"], 4)],
            ["Weighted OOF MAE", _fmt(reg["weighted_mae"], 1)],
            ["Bootstrap R² (mean / 95% CI)", f"{_fmt(ci['r2_boot_mean'], 4)} [{_fmt(ci['r2_ci_lo'], 4)}, {_fmt(ci['r2_ci_hi'], 4)}]"],
            ["Bootstrap MAE (mean / 95% CI)", f"{_fmt(ci['mae_boot_mean'], 1)} [{_fmt(ci['mae_ci_lo'], 1)}, {_fmt(ci['mae_ci_hi'], 1)}]"],
            ["Bootstrap valid resamples (R²/MAE)", f"{ci['n_valid_r2']} / {ci['n_valid_mae']}"],
            ["Non-completers predicted to walk", reg["n_noncomp_pred_walk"]],
            ["Non-completers predicted not to walk", reg["n_noncomp_pred_no_walk"]],
            ["Model feature count", len(res["features"])],
        ])

        doc.add_paragraph()
        sh = doc.add_paragraph()
        sh.add_run("Winsorization-Threshold Sensitivity (tier Ridge OOF R² and MAE)").bold = True
        sens_rows = [["Threshold", "Winsor lower", "Winsor upper", "Wt mean", "% truncated", "OOF R²", "OOF MAE"]]
        for pct in WINSORIZATION_THRESHOLDS:
            s = ipcw_completion_info["sensitivity"][pct]
            sens_rows.append([
                f"{100-pct}th/{pct}th",
                _fmt(s["lower"], 3), _fmt(s["upper"], 3),
                _fmt(s["mean"], 4),
                _fmt(s["pct_truncated"] * 100, 1) + "%",
                _fmt(reg["sensitivity_r2"].get(pct, float("nan")), 4),
                _fmt(reg["sensitivity_mae"].get(pct, float("nan")), 1),
            ])
        _add_table_to_doc(doc, sens_rows)

        r2_vals = [reg["sensitivity_r2"].get(p, float("nan")) for p in WINSORIZATION_THRESHOLDS]
        finite_r2 = [v for v in r2_vals if math.isfinite(v)]
        if finite_r2:
            r2_spread = max(finite_r2) - min(finite_r2)
            stability_note = (
                "Stable (|ΔR²| < 0.05 across thresholds)."
                if r2_spread < 0.05
                else "⚠ Unstable (|ΔR²| ≥ 0.05 across thresholds)."
            )
        else:
            stability_note = "Cannot assess."
        p_stab = doc.add_paragraph(f"Sensitivity: {stability_note}")
        _small_para(p_stab, 9)

    # ── Methodology explainer ─────────────────────────────────────────────────
    doc.add_paragraph()
    mh = doc.add_paragraph()
    mh.add_run("Methodology Notes").bold = True

    meth_items = [
        ("Model 5 — Rehab LOS Restriction",
         "T1T2_Change variables (BI, BBS, MRS, FOIS, MNA, IADL, FuglUE, FuglSEN, EuroQoL5D) "
         "capture functional change between T1 (post-acute assessment) and T2 (discharge). "
         "These change scores are only clinically interpretable for patients who had sufficient "
         "rehabilitation exposure.  Patients with Rehab LOS < 21 days (categories '0-7 days', "
         "'8-14 days', '15-20 days') have minimal rehabilitation time and their T1T2_Change "
         "values conflate short-stay selection effects with genuine rehabilitation-induced change. "
         "Model 5 is therefore restricted to Rehab_LOS_Category ∈ {'21-42 days', '>42 days'}. "
         "Excluded patients are not counted in the patient N, and their T1T2_Change values are "
         "left blank and not imputed. In this update, BI_T1T2_Change is blanked for LOS<=21 days and also blanked for BI2=0 within 21-42 days LOS, "
         "FOIS_T1T2_Change is blanked whenever FOIS2=0, and out-of-range T2-linked change values are removed."),
        ("Single IPCW Completion Model",
         "In the prior session (20260831_IPCW_2331), a separate completion model was fitted for "
         "each tier using only that tier's features.  This creates a methodological inconsistency: "
         "the weighting model changes across compared models, so differences in weighted performance partly "
         "reflect differences in the quality of the completion model rather than differences in "
         "the outcome model's predictor set.  The corrected approach (this session) fits ONE "
         "completion model using the richest available predictor set "
         "(demographics + stroke info + comorbidities + acute complications + NIHSS out + "
         "functional T1 with imputed gait speed) and applies the resulting weights uniformly "
         "across the retained-model regressions.  This is consistent with the standard causal "
         "inference / IPCW framework."),
        ("Intercept-only numerator",
         "The numerator in a stabilized IPCW weight controls how far the weights deviate from 1. "
         "An intercept-only numerator (marginal completion rate) is the most common and "
         "best-validated choice: simple, robust to numerator mis-specification, and typically "
         "gives weights whose mean is near 1 when the denominator model is well-specified."),
        ("Full-pipeline bootstrap",
         "The two-stage nature of the IPCW pipeline (completion weighting + outcome model) "
         "means OOF metrics that treat weights as fixed understate uncertainty. "
         "The bootstrap (B=500, stratified by completion status) re-estimates all stages — "
         "completion model, weight construction, binary classifier, Ridge fitting — within each "
         "resample, propagating full pipeline uncertainty into the reported 95% CIs."),
    ]
    for title_str, body_str in meth_items:
        h = doc.add_paragraph()
        h.add_run(title_str).bold = True
        p = doc.add_paragraph(body_str)
        _small_para(p, 9)

    repro = doc.add_paragraph(
        f"Reproducibility: run `python {Path(__file__).name}` to regenerate this report and the Excel output."
    )
    _small_para(repro, 8)

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX.name}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    if not INPUT_XLSX.exists():
        raise FileNotFoundError(f"Missing required input dataset: {INPUT_XLSX.name}")
    print(f"Loading: {INPUT_XLSX.name}")
    df = pd.read_excel(INPUT_XLSX)
    print(f"Dataset: {df.shape[0]:,} rows × {df.shape[1]} columns")

    fuglsen_outliers_pre = extract_fuglsen_outliers(df)
    df, manual_corrections_df = apply_manual_patient_corrections(df)
    if not manual_corrections_df.empty:
        print("\nManual data corrections applied:")
        for _, row in manual_corrections_df.iterrows():
            print(
                f"  ID {int(row['ID'])} {row['Column']}: "
                f"{row['Original_Value']} -> {row['Corrected_Value']}"
            )

    df, audit_info = audit_and_clean_t1t2_data(df)
    audit_info["manual_corrections_df"] = manual_corrections_df
    audit_info["fuglsen_outliers_pre_df"] = fuglsen_outliers_pre
    audit_summary_df = audit_info["summary_df"]
    if not audit_summary_df.empty:
        print("\nT1T2 audit cleanup:")
        for _, row in audit_summary_df.iterrows():
            print(
                f"  {row['Assessment']}: cleared {int(row['Rows_Cleared_Total'])} change values "
                f"(LOS-rule={int(row.get('Rows_Cleared_BI_LOS_Rule', 0))}, "
                f"zero-T2={int(row['Rows_Cleared_Zero_T2_Rule'])}, "
                f"out-of-range={int(row['Rows_Cleared_Out_Of_Range'])})"
            )

    retained_model_specs = {
        model_name: MODEL_SPECS[model_name]
        for model_name in RETAINED_MODEL_ORDER
        if model_name in MODEL_SPECS
    }
    all_candidate_cols = list({c for cols in retained_model_specs.values() for c in cols})
    for c in all_candidate_cols + ["6MWT4"] + NIHSS_IN:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    # ── PART 1: Bootstrap LASSO ───────────────────────────────────────────────
    print("\n" + "=" * 78)
    print("PART 1: Bootstrap LASSO Models")
    lasso_results: list[dict] = []
    for i, model_name in enumerate(RETAINED_MODEL_ORDER):
        candidate_features = retained_model_specs.get(model_name, [])
        valid_features = _filter_existing(candidate_features, df)

        print("\n" + "-" * 78)
        print(f"  {model_name}")
        print(f"  Candidate: {len(candidate_features)} → In data: {len(valid_features)}")

        if not valid_features:
            print("  !! No valid features – skipping.")
            continue

        is_t1t2_model = any(col in candidate_features for col in T1T2_IMPROVEMENT)
        if is_t1t2_model and "Rehab_LOS_Category" in df.columns:
            n_qualifying = int(
                df["Rehab_LOS_Category"].isin(QUALIFYING_REHAB_LOS).sum()
            )
            print(f"  T1T2 model: restricted to Rehab LOS ∈ {QUALIFYING_REHAB_LOS} "
                  f"({n_qualifying:,} patients qualify)")

        res = fit_bootstrap_lasso(
            df, valid_features, model_name,
            model_seed=RANDOM_STATE + i,
            model5_restrict=is_t1t2_model,
        )
        lasso_results.append(res)

        print(
            f"  Patients: {res['n_patients']:,} | Input vars: {res['n_input_variables']} | "
            f"Non-zero: {res['n_final_nonzero_base']} | Stable: {res['n_stable_selected']}"
        )
        print(
            f"  CV  R²: {res['cv_r2']:.4f}  MAE: {res['cv_mae']:.2f} m  |  "
            f"Train R²: {res['train_r2']:.4f}  MAE: {res['train_mae']:.2f} m"
        )

    # ── PART 2: IPCW – single completion model ────────────────────────────────
    print("\n" + "=" * 78)
    print("PART 2: IPCW – computing single completion model")

    completion_features = _filter_existing(IPCW_COMPLETION_FEATURES, df)
    print(f"  Completion model features: {len(completion_features)}")

    ipcw_completion = compute_ipcw_weights_single_model(df, completion_features)
    wd = ipcw_completion["weight_diagnostics"]
    pd_diag = ipcw_completion["p_denom_diagnostics"]
    print(f"  Numerator: {ipcw_completion['numerator_model']}")
    print(f"  Weights: mean={wd['mean']:.4f}  SD={wd['sd']:.4f}  "
          f"min={wd['min']:.4f}  max={wd['max']:.4f}  %trunc={wd['pct_truncated']*100:.1f}%")
    if pd_diag["positivity_flag"]:
        print(f"  ⚠ Positivity flag: {pd_diag['pct_near_zero']*100:.1f}% near 0, "
              f"{pd_diag['pct_near_one']*100:.1f}% near 1")

    ipcw_tier_results: list[dict] = []
    output_df = df.copy()

    for model_name in RETAINED_MODEL_ORDER:
        model_features = retained_model_specs.get(model_name, [])
        valid_model_features = _filter_existing(model_features, df)
        if not valid_model_features:
            continue

        model_label = model_name.split(":", 1)[0]
        print(f"\n  {model_label} ({len(valid_model_features)} features)")

        for scenario in SCENARIOS:
            scenario_label = scenario.replace("6MWT_", "").replace("_Scenario", "")
            print(f"    Scenario: {scenario_label}", end="  ", flush=True)

            binary_result = select_binary_classifier(
                df,
                valid_model_features,
                scenario,
            )
            regression_result = fit_weighted_ridge(df, valid_model_features, ipcw_completion, binary_result)

            print(
                f"BinAcc={binary_result['bal_acc']:.3f}  "
                f"R²={regression_result['weighted_r2']:.4f}  "
                f"MAE={regression_result['weighted_mae']:.1f}",
                end="  bootstrap...", flush=True,
            )
            ci = bootstrap_ci_ipcw(
                df, completion_features, valid_model_features, scenario,
                n_bootstrap=N_BOOTSTRAP_IPCW,
            )
            print(f"R² CI=[{ci['r2_ci_lo']:.4f}, {ci['r2_ci_hi']:.4f}]")

            model_token = model_label.replace(" ", "")
            walk_col = f"IPCW_Walk_{model_token}_{scenario_label}"
            pred_col = f"IPCW_6MWT4_{model_token}_{scenario_label}"
            output_df[walk_col] = pd.Series(pd.NA, index=output_df.index, dtype="Int64")
            output_df[pred_col] = np.nan
            nc_idx = regression_result["noncomp_index"]
            output_df.loc[nc_idx, walk_col] = regression_result["noncomp_walk_pred"]
            output_df.loc[nc_idx, pred_col] = regression_result["noncomp_6mwt4_pred"]

            ipcw_tier_results.append({
                "model": model_label,
                "scenario": scenario,
                "scenario_label": scenario_label,
                "features": valid_model_features,
                "binary": binary_result,
                "regression": regression_result,
                "bootstrap_ci": ci,
            })

    if "Rehab_LOS_Category" in output_df.columns:
        non_qualifying_los = output_df["Rehab_LOS_Category"].notna() & ~output_df["Rehab_LOS_Category"].isin(QUALIFYING_REHAB_LOS)
        t1t2_cols_present = [c for c in T1T2_IMPROVEMENT if c in output_df.columns]
        if t1t2_cols_present:
            output_df.loc[non_qualifying_los, t1t2_cols_present] = np.nan

    # Add LASSO predictions to output_df
    for r in lasso_results:
        col_name = MODEL_PRED_COLS[r["model_name"]]
        pipe: Pipeline = r["base_pipe"]
        features: list[str] = r["features"]
        if any(col in features for col in T1T2_IMPROVEMENT):
            # Only predict for qualifying patients; leave others as NaN
            qualifying_mask = (
                df["Rehab_LOS_Category"].isin(QUALIFYING_REHAB_LOS)
                if "Rehab_LOS_Category" in df.columns
                else pd.Series(True, index=df.index)
            )
            preds = pd.Series(np.nan, index=df.index)
            q_idx = df.index[qualifying_mask]
            preds.loc[q_idx] = pipe.predict(df.loc[q_idx, features])
            output_df[col_name] = preds
        else:
            output_df[col_name] = pipe.predict(df[features])

    # Summary sheet for IPCW
    summary_rows = []
    for res in ipcw_tier_results:
        ci_s = res["bootstrap_ci"]
        reg = res["regression"]
        summary_rows.append({
            "Model": res["model"],
            "Scenario": res["scenario_label"],
            "Completion_model": f"Single model ({len(completion_features)} features)",
            "Numerator_mode": NUMERATOR_MODE,
            "Binary_model": res["binary"]["best_model_name"],
            "Binary_OOF_Bal_Acc": round(res["binary"]["bal_acc"], 3),
            "Weighted_OOF_R2": round(reg["weighted_r2"], 4),
            "Weighted_OOF_MAE": round(reg["weighted_mae"], 1),
            "Boot_R2_mean": round(ci_s["r2_boot_mean"], 4),
            "Boot_R2_CI_lo": round(ci_s["r2_ci_lo"], 4),
            "Boot_R2_CI_hi": round(ci_s["r2_ci_hi"], 4),
            "Boot_MAE_mean": round(ci_s["mae_boot_mean"], 1),
            "Boot_MAE_CI_lo": round(ci_s["mae_ci_lo"], 1),
            "Boot_MAE_CI_hi": round(ci_s["mae_ci_hi"], 1),
            "Boot_valid_resamples": ci_s["n_valid_r2"],
            "N_noncomp_pred_walk": reg["n_noncomp_pred_walk"],
            "N_noncomp_pred_no_walk": reg["n_noncomp_pred_no_walk"],
        })
    summary_df = pd.DataFrame(summary_rows)
    if not summary_df.empty:
        summary_df["Performance_Rank_within_Scenario"] = (
            summary_df.groupby("Scenario")["Weighted_OOF_R2"]
            .rank(method="dense", ascending=False)
            .astype(int)
        )
        summary_df = summary_df.sort_values(
            ["Scenario", "Performance_Rank_within_Scenario", "Weighted_OOF_MAE"],
            ascending=[True, True, True],
        ).reset_index(drop=True)

    # LASSO summary sheet
    lasso_summary_rows = []
    for r in lasso_results:
        cat = r["category_counts"]
        lasso_summary_rows.append({
            "Model": r["model_name"],
            "N_patients": r["n_patients"],
            "Input_vars": r["n_input_variables"],
            "Gait_speed_status": r["gait_speed_imputed_status"],
            "Predictor_categories_summary": r["predictor_category_summary"],
            "N_Demographics": cat["Demographics"],
            "N_Stroke_info": cat["Stroke info"],
            "N_Comorbidities": cat["Comorbidities"],
            "N_Acute_complications": cat["Acute complications"],
            "N_NIHSS_out": cat["NIHSS out"],
            "N_Functional_T1": cat["Functional T1"],
            "N_T1T2_improvement": cat["T1T2 improvement"],
            "Final_nonzero": r["n_final_nonzero_base"],
            "Stable_selected": r["n_stable_selected"],
            "Best_alpha": round(r["best_alpha_full_fit"], 6),
            "CV_R2": round(r["cv_r2"], 4),
            "CV_MAE": round(r["cv_mae"], 2),
            "Train_R2": round(r["train_r2"], 4),
            "Train_MAE": round(r["train_mae"], 2),
        })
    lasso_summary_df = pd.DataFrame(lasso_summary_rows)
    if not lasso_summary_df.empty:
        lasso_summary_df["Performance_Rank"] = (
            lasso_summary_df["CV_R2"].rank(method="dense", ascending=False).astype(int)
        )
        lasso_summary_df = lasso_summary_df.sort_values(
            ["Performance_Rank", "CV_MAE"], ascending=[True, True]
        ).reset_index(drop=True)

    with pd.ExcelWriter(OUTPUT_XLSX, engine="openpyxl") as writer:
        output_df.to_excel(writer, sheet_name="Predictions", index=False)
        lasso_summary_df.to_excel(writer, sheet_name="LASSO_Summary", index=False)
        summary_df.to_excel(writer, sheet_name="IPCW_Summary", index=False)
        audit_info["summary_df"].to_excel(writer, sheet_name="T1T2_Audit_Summary", index=False)
        audit_info["bi_cases_df"].to_excel(writer, sheet_name="BI2_No_Assessment", index=False)
        audit_info["flagged_cases_df"].to_excel(writer, sheet_name="T1T2_Flagged_Cases", index=False)
        audit_info["illogical_values_df"].to_excel(writer, sheet_name="Illogical_T1T2", index=False)
        audit_info["raw_outliers_df"].to_excel(writer, sheet_name="Functional_Outliers", index=False)
        audit_info["fuglsen_outliers_pre_df"].to_excel(writer, sheet_name="FuglSEN_Outliers", index=False)
        audit_info["manual_corrections_df"].to_excel(writer, sheet_name="Manual_Corrections", index=False)
    print(f"\nSaved: {OUTPUT_XLSX.name}")

    write_combined_report(lasso_results, ipcw_completion, ipcw_tier_results, audit_info)

    print("\n=== LASSO SUMMARY ===")
    for rank, r in enumerate(sorted(lasso_results, key=lambda x: x["cv_r2"], reverse=True), start=1):
        print(f"  #{rank:<2} {r['model_name'][:56]:<56}  N={r['n_patients']:>4}  CV R²={r['cv_r2']:.4f}  MAE={r['cv_mae']:.2f}")

    print("\n=== IPCW SUMMARY ===")
    grouped = {}
    for r in ipcw_tier_results:
        grouped.setdefault(r["scenario_label"], []).append(r)
    for scenario_label, rows in grouped.items():
        print(f"  Scenario={scenario_label}")
        ranked_rows = sorted(rows, key=lambda x: x["regression"]["weighted_r2"], reverse=True)
        for rank, r in enumerate(ranked_rows, start=1):
            ci = r["bootstrap_ci"]
            print(
                f"    #{rank:<2} {r['model']:<8}  "
                f"BinAcc={r['binary']['bal_acc']:.3f}  "
                f"Wt R²={r['regression']['weighted_r2']:.4f}  "
                f"MAE={r['regression']['weighted_mae']:.1f}  "
                f"Boot CI=[{ci['r2_ci_lo']:.4f},{ci['r2_ci_hi']:.4f}]"
            )


if __name__ == "__main__":
    main()
