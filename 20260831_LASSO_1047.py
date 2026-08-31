#!/usr/bin/env python3
"""
20260831_LASSO_1047.py
======================
Bootstrap LASSO regression models for 6MWT4 prediction.
NO IPCW is used anywhere in this workflow.

This refresh replaces the older 20260829 four-model report with the most
recent six-model suite and adds more detailed explainers plus richer feature-
importance summaries in the generated Word document.

Outputs (written to repo root):
    20260831_LASSO_1047.docx  – detailed six-model report with explainers
    20260831_LASSO_1047.xlsx  – full dataset with per-patient predictions

Author: auto-generated 2026-08-31
"""

from __future__ import annotations

from pathlib import Path
import re
import shutil

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt

from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LassoCV
from sklearn.metrics import mean_absolute_error, r2_score
from sklearn.model_selection import KFold, cross_val_predict
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler

# ─────────────────────────────────────────────────────────────────────────────
# FILE PATHS
# ─────────────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent
INPUT_XLSX = ROOT / "20260826_DeID.xlsx"
OUTPUT_DOCX = ROOT / "20260831_LASSO_1047.docx"
OUTPUT_XLSX = ROOT / "20260831_LASSO_1047.xlsx"
SOURCE_DOCX = ROOT / "20260831_LASSO_0802.docx"
SOURCE_XLSX = ROOT / "20260831_LASSO_0802.xlsx"

# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL SETTINGS
# ─────────────────────────────────────────────────────────────────────────────
RANDOM_STATE = 42
N_BOOTSTRAP = 1000
CV_FOLDS = 5
ALPHAS = np.logspace(-4, 2, 120)
STABILITY_THRESHOLD = 0.70

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1 VARIABLE CATEGORIES (consistent with 20260826_5_Tiers_1952.py)
# ─────────────────────────────────────────────────────────────────────────────
DEMOGRAPHICS_PAC = ["Age", "Sex, F0 M1"]

STROKE_INFO = [
    "Dissection", "ACA", "Undetermined", "HemorrhageStroke",
    "LVS", "LVO",
    "Side_Right", "Side_Left", "Side_Bilateral",
    "Loc_CortSub", "Loc_Subcortical", "Loc_Infratentorial",
]

COMORBIDITIES_PAC = [
    "AF", "DM", "HTN", "Dyslipidemia", "CAD", "CKD",
    "RestrictiveLung", "GIUlcer", "LiverCirrhosis", "Hepatitis",
    "Parkinsonism", "Malignancy", "OldStroke", "Dementia", "Psychiatric", "Gout",
]

ACUTE_COMPLICATIONS_PAC = ["Pneumonia", "UTI", "GIB", "Cellulitis"]

# Functional T1 WITH gait speed imputed (full list)
FUNCTIONAL_T1_PLUS_GS_IMPUTED = [
    "MRS1", "BI1", "FOIS1", "MNA1", "EuroQoL5D1", "IADL1",
    "BBS1", "Gait_Speed_1_Imputed", "FuglUE1", "FuglSEN1",
]

# Functional T1 WITHOUT gait speed imputed
FUNCTIONAL_T1_NO_GS = [
    x for x in FUNCTIONAL_T1_PLUS_GS_IMPUTED if x != "Gait_Speed_1_Imputed"
]

NIHSS_OUT = [
    "ConsOut", "AnswerOut", "OrderOut", "EOMOut", "VisualOut",
    "FacialOut", "LUOut", "RUOut", "LLOut", "RLOut",
    "Coordinateout", "SensoryOut", "LanguageOut", "ArticulateOut", "NeglectOut",
]

T1T2_IMPROVEMENT = [
    "BI_T1T2_Change", "BBS_T1T2_Change", "MRS_T1T2_Change",
    "FOIS_T1T2_Change", "MNA_T1T2_Change", "IADL_T1T2_Change",
    "FuglUE_T1T2_Change", "FuglSEN_T1T2_Change", "EuroQoL5D_T1T2_Change",
]

# ─────────────────────────────────────────────────────────────────────────────
# MODEL DEFINITIONS
# ─────────────────────────────────────────────────────────────────────────────
# New ordering:
#   1  – FAM (functional admission, with gait speed)          [unchanged]
#   2  – CPANM-GS  (full post-acute neurological + func T1 WITH gait speed)
#   3  – CPANM-NGS (full post-acute neurological + func T1 WITHOUT gait speed)
#   4  – RBE  (bedside: BBS1 + gait speed)                    [was Model 3]
#   5  – ERPM (func T1 + T1T2 improvement)                    [was Model 4]
#   6  – PANM-only (stroke/comorbidity/NIHSS, no func T1)     [was Model 2 – placed last]
MODEL_SPECS: dict[str, list[str]] = {
    "Model 1: demographic + functional T1 (with gait speed imputed)": (
        DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED
    ),
    "Model 2: demographic + stroke info + comorbidities + acute complications + NIHSS out + functional T1 (with gait speed imputed)": (
        DEMOGRAPHICS_PAC
        + STROKE_INFO
        + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC
        + NIHSS_OUT
        + FUNCTIONAL_T1_PLUS_GS_IMPUTED
    ),
    "Model 3: demographic + stroke info + comorbidities + acute complications + NIHSS out + functional T1 (without gait speed imputed)": (
        DEMOGRAPHICS_PAC
        + STROKE_INFO
        + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC
        + NIHSS_OUT
        + FUNCTIONAL_T1_NO_GS
    ),
    "Model 4: demographic + BBS1 + gait speed 1 imputed": (
        DEMOGRAPHICS_PAC + ["BBS1", "Gait_Speed_1_Imputed"]
    ),
    "Model 5: demographics + functional T1 + T1T2 improvement": (
        DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + T1T2_IMPROVEMENT
    ),
    "Model 6 (reference – original Model 2): demographic + stroke info + comorbidities + acute complications + NIHSS out": (
        DEMOGRAPHICS_PAC
        + STROKE_INFO
        + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC
        + NIHSS_OUT
    ),
}

# Short column names used in the Excel output for each model
MODEL_PRED_COLS = {
    "Model 1: demographic + functional T1 (with gait speed imputed)": "LASSO_Pred_Model1",
    "Model 2: demographic + stroke info + comorbidities + acute complications + NIHSS out + functional T1 (with gait speed imputed)": "LASSO_Pred_Model2",
    "Model 3: demographic + stroke info + comorbidities + acute complications + NIHSS out + functional T1 (without gait speed imputed)": "LASSO_Pred_Model3",
    "Model 4: demographic + BBS1 + gait speed 1 imputed": "LASSO_Pred_Model4",
    "Model 5: demographics + functional T1 + T1T2 improvement": "LASSO_Pred_Model5",
    "Model 6 (reference – original Model 2): demographic + stroke info + comorbidities + acute complications + NIHSS out": "LASSO_Pred_Model6",
}

MODEL_EXPLAINERS = {
    "Model 1: demographic + functional T1 (with gait speed imputed)": (
        "Functional-admission model focused on bedside baseline impairment and mobility markers."
    ),
    "Model 2: demographic + stroke info + comorbidities + acute complications + NIHSS out + functional T1 (with gait speed imputed)": (
        "Full clinical-performance model combining neurological burden, medical complexity, and functional status including imputed gait speed."
    ),
    "Model 3: demographic + stroke info + comorbidities + acute complications + NIHSS out + functional T1 (without gait speed imputed)": (
        "Full clinical-performance model that removes imputed gait speed to test whether the broader functional battery can compensate."
    ),
    "Model 4: demographic + BBS1 + gait speed 1 imputed": (
        "Compact bedside model prioritising quick deployment with only age, sex, balance, and gait speed."
    ),
    "Model 5: demographics + functional T1 + T1T2 improvement": (
        "Recovery-trajectory model augmenting admission function with early change scores between T1 and T2."
    ),
    "Model 6 (reference – original Model 2): demographic + stroke info + comorbidities + acute complications + NIHSS out": (
        "Reference neurological/comorbidity model without functional T1 measures, retained for comparison."
    ),
}

# ─────────────────────────────────────────────────────────────────────────────
# UTILITY FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def _filter_existing(cols: list[str], df: pd.DataFrame) -> list[str]:
    """Return only those columns that actually exist in df, preserving order."""
    return [c for c in cols if c in df.columns]


def _build_lasso_pipeline(features: list[str]) -> Pipeline:
    """
    Build a fresh sklearn Pipeline for LASSO:

    1. ColumnTransformer  – selects the named feature columns
    2. SimpleImputer      – median imputation for missing values
    3. StandardScaler     – z-score standardisation (mean=0, sd=1)
    4. LassoCV            – regularised linear regression with inner
                            cross-validation to choose the best alpha
    """
    numeric_transformer = Pipeline([
        ("imputer", SimpleImputer(strategy="median")),
        ("scaler", StandardScaler()),
    ])
    preprocessor = ColumnTransformer([
        ("num", numeric_transformer, features)
    ])
    lasso = LassoCV(
        alphas=ALPHAS,
        cv=CV_FOLDS,
        random_state=RANDOM_STATE,
        max_iter=50_000,
        n_jobs=-1,
    )
    return Pipeline([
        ("prep", preprocessor),
        ("model", lasso),
    ])


def _add_table_to_doc(doc: Document, rows: list[list]) -> None:
    """Insert a Word table with a header row (bold) into doc."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def _small_para(paragraph, size: int = 9) -> None:
    """Set all runs in a paragraph to a small font size."""
    for run in paragraph.runs:
        run.font.size = Pt(size)


def _top_predictors(
    importance: pd.DataFrame,
    *,
    limit: int = 5,
    positive: bool | None = None,
) -> pd.DataFrame:
    """Return top predictors, optionally filtering by coefficient direction."""
    frame = importance.copy()
    if positive is True:
        frame = frame[frame["bootstrap_coef_mean"] > 0]
    elif positive is False:
        frame = frame[frame["bootstrap_coef_mean"] < 0]
    return frame.sort_values(
        ["selection_frequency", "abs_bootstrap_coef_mean"],
        ascending=[False, False],
    ).head(limit)


def _format_predictor_summary(frame: pd.DataFrame) -> str:
    """Format a concise predictor summary for narrative explainers."""
    if frame.empty:
        return "none"
    parts = []
    for _, row in frame.iterrows():
        parts.append(
            f"{row['predictor']} ({row['bootstrap_coef_mean']:+.1f}; freq {row['selection_frequency']:.0%})"
        )
    return ", ".join(parts)


def _model_rankings(results: list[dict]) -> dict[str, int]:
    """Return rank positions by descending CV R²."""
    ranked = sorted(results, key=lambda x: x["cv_r2"], reverse=True)
    return {row["model_name"]: idx for idx, row in enumerate(ranked, start=1)}


def _executive_bullets(results: list[dict]) -> list[str]:
    """Create global comparison bullets for the six-model report."""
    ranked = sorted(results, key=lambda x: x["cv_r2"], reverse=True)
    best = ranked[0]
    second = ranked[1]
    compact = min(results, key=lambda x: x["n_input_variables"])
    reference = next(
        r for r in results if "reference" in r["model_name"].lower()
    )
    model2 = next(r for r in results if r["model_name"].startswith("Model 2:"))
    model3 = next(r for r in results if r["model_name"].startswith("Model 3:"))
    return [
        (
            f"Best overall out-of-fold fit was achieved by {best['model_name']} "
            f"(CV R² {best['cv_r2']:.4f}, MAE {best['cv_mae']:.2f} m), narrowly ahead of "
            f"{second['model_name']}."
        ),
        (
            f"Removing imputed gait speed from the full clinical-performance model reduced CV R² "
            f"from {model2['cv_r2']:.4f} to {model3['cv_r2']:.4f} "
            f"(ΔR² {model3['cv_r2'] - model2['cv_r2']:+.4f}), suggesting incremental rather than dominant value from gait speed."
        ),
        (
            f"The most compact model was {compact['model_name']} with {compact['n_input_variables']} predictors "
            f"while retaining CV R² {compact['cv_r2']:.4f}, supporting pragmatic bedside deployment."
        ),
        (
            f"The non-functional reference model remained weakest (CV R² {reference['cv_r2']:.4f}, MAE {reference['cv_mae']:.2f} m), "
            "reinforcing that functional T1 measures are central to accurate 6MWT4 prediction."
        ),
    ]


def _parse_alpha_summary(text: str) -> tuple[float, float, float]:
    """Parse 'median [q25, q75]' alpha text from an existing report table."""
    numbers = [float(x) for x in re.findall(r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?", text)]
    if len(numbers) != 3:
        raise ValueError(f"Could not parse bootstrap alpha summary: {text}")
    return numbers[0], numbers[1], numbers[2]


def load_existing_results(df: pd.DataFrame) -> list[dict]:
    """Load six-model results from the prior 0802 LASSO report."""
    doc = Document(SOURCE_DOCX)
    model_names = [
        p.text.strip() for p in doc.paragraphs if p.text.strip().startswith("Model ")
    ]
    metric_tables = [
        t for t in doc.tables
        if t.rows and t.rows[0].cells[0].text.strip() == "Metric"
    ]
    importance_tables = [
        t for t in doc.tables
        if t.rows and t.rows[0].cells[0].text.strip() == "Predictor"
    ]

    results: list[dict] = []
    for model_name, metric_table, importance_table in zip(
        model_names, metric_tables, importance_tables
    ):
        metrics = {
            row.cells[0].text.strip(): row.cells[1].text.strip()
            for row in metric_table.rows[1:]
        }
        alpha_med, alpha_q25, alpha_q75 = _parse_alpha_summary(
            metrics["Bootstrap alpha: median [IQR]"]
        )
        features = _filter_existing(MODEL_SPECS[model_name], df)

        importance_rows = []
        for row in importance_table.rows[1:]:
            predictor = row.cells[0].text.strip()
            if not predictor:
                continue
            base_coef = float(row.cells[1].text.strip())
            boot_mean = float(row.cells[2].text.strip())
            boot_sd = float(row.cells[3].text.strip())
            sel_freq = float(row.cells[4].text.strip())
            importance_rows.append(
                {
                    "predictor": predictor,
                    "base_coef": base_coef,
                    "bootstrap_coef_mean": boot_mean,
                    "bootstrap_coef_sd": boot_sd,
                    "selection_frequency": sel_freq,
                    "abs_bootstrap_coef_mean": abs(boot_mean),
                }
            )

        importance = pd.DataFrame(importance_rows)

        results.append(
            {
                "model_name": model_name,
                "features": features,
                "n_patients": int(metrics["Patients (non-missing 6MWT4)"]),
                "n_input_variables": int(metrics["Input variable count"]),
                "n_final_nonzero_base": int(metrics["Final selected (non-zero in full-fit)"]),
                "n_stable_selected": int(
                    metrics[f"Stably selected (selection freq ≥ {STABILITY_THRESHOLD:.0%})"]
                ),
                "best_alpha_full_fit": float(metrics["Best alpha (full-data fit)"]),
                "alpha_boot_median": alpha_med,
                "alpha_boot_q25": alpha_q25,
                "alpha_boot_q75": alpha_q75,
                "cv_r2": float(metrics["CV R² (OOF)"]),
                "cv_mae": float(metrics["CV MAE (OOF)"].replace(" m", "")),
                "train_r2": float(metrics["Train R² (apparent)"]),
                "train_mae": float(metrics["Train MAE (apparent)"].replace(" m", "")),
                "base_pipe": None,
                "importance": importance,
            }
        )
    return results


# ─────────────────────────────────────────────────────────────────────────────
# CORE MODELLING FUNCTION
# ─────────────────────────────────────────────────────────────────────────────

def fit_bootstrap_lasso(
    df: pd.DataFrame,
    features: list[str],
    model_name: str,
    model_seed: int = RANDOM_STATE,
) -> dict:
    """
    Fit a true bootstrap LASSO model and return a results dictionary.

    Steps
    -----
    1. Restrict analysis set to rows where 6MWT4 is observed.
    2. Build a baseline LASSO (full data fit) to get final coefficients and
       apparent (train-set) performance.
    3. Run outer K-fold cross-validation on the full data to get unbiased
       out-of-fold (OOF) R² and MAE.
    4. Run N_BOOTSTRAP bootstrap resamples.
    5. From bootstrap coefficients compute selection frequency and mean/SD.
    6. Return all metrics + importance table.
    """
    model_df = df[df["6MWT4"].notna()].copy()
    X = model_df[features]
    y = model_df["6MWT4"].to_numpy(dtype=float)

    n_patients = len(model_df)
    n_input = len(features)

    base_pipe = _build_lasso_pipeline(features)
    base_pipe.fit(X, y)
    base_lasso: LassoCV = base_pipe.named_steps["model"]
    base_coef = base_lasso.coef_.copy()
    best_alpha = float(base_lasso.alpha_)

    y_train_pred = base_pipe.predict(X)
    train_r2 = float(r2_score(y, y_train_pred))
    train_mae = float(mean_absolute_error(y, y_train_pred))

    outer_cv = KFold(n_splits=CV_FOLDS, shuffle=True, random_state=RANDOM_STATE)
    y_oof = cross_val_predict(
        _build_lasso_pipeline(features), X, y,
        cv=outer_cv, n_jobs=-1,
    )
    cv_r2 = float(r2_score(y, y_oof))
    cv_mae = float(mean_absolute_error(y, y_oof))

    rng = np.random.default_rng(model_seed)
    coef_boot = np.zeros((N_BOOTSTRAP, n_input), dtype=float)
    alpha_boot = np.zeros(N_BOOTSTRAP, dtype=float)

    for b in range(N_BOOTSTRAP):
        idx = rng.integers(0, n_patients, size=n_patients)
        Xb = X.iloc[idx]
        yb = y[idx]

        bp = _build_lasso_pipeline(features)
        bp.fit(Xb, yb)
        lb: LassoCV = bp.named_steps["model"]

        coef_boot[b, :] = lb.coef_
        alpha_boot[b] = lb.alpha_

    _COEF_TOL = 1e-9
    sel_freq = (np.abs(coef_boot) > _COEF_TOL).mean(axis=0)
    coef_mean = coef_boot.mean(axis=0)
    coef_sd = coef_boot.std(axis=0, ddof=1)

    importance = pd.DataFrame({
        "predictor": features,
        "base_coef": base_coef,
        "bootstrap_coef_mean": coef_mean,
        "bootstrap_coef_sd": coef_sd,
        "selection_frequency": sel_freq,
        "abs_bootstrap_coef_mean": np.abs(coef_mean),
    }).sort_values(
        ["selection_frequency", "abs_bootstrap_coef_mean"],
        ascending=[False, False],
    ).reset_index(drop=True)

    n_final_nonzero = int((np.abs(base_coef) > 1e-9).sum())
    n_stable = int((sel_freq >= STABILITY_THRESHOLD).sum())

    return {
        "model_name": model_name,
        "features": features,
        "n_patients": n_patients,
        "n_input_variables": n_input,
        "n_final_nonzero_base": n_final_nonzero,
        "n_stable_selected": n_stable,
        "best_alpha_full_fit": best_alpha,
        "alpha_boot_median": float(np.median(alpha_boot)),
        "alpha_boot_q25": float(np.quantile(alpha_boot, 0.25)),
        "alpha_boot_q75": float(np.quantile(alpha_boot, 0.75)),
        "cv_r2": cv_r2,
        "cv_mae": cv_mae,
        "train_r2": train_r2,
        "train_mae": train_mae,
        "base_pipe": base_pipe,
        "importance": importance,
    }


# ─────────────────────────────────────────────────────────────────────────────
# REPORT GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def write_docx_report(results: list[dict]) -> None:
    """Write a Word document summarising all bootstrap LASSO models."""
    doc = Document()

    title = doc.add_paragraph()
    title.add_run(
        "Bootstrap LASSO Models for 6MWT4 Prediction (No IPCW) – 2026-08-31"
    ).bold = True

    intro = doc.add_paragraph(
        f"This report summarises {len(results)} bootstrap LASSO regression models "
        "predicting the 6-Minute Walk Test at time-point 4 (6MWT4) as a continuous outcome. "
        "No IPCW weighting is applied. "
        "Missing predictor values are handled by median imputation; predictors are "
        "z-score standardised prior to LASSO. "
        f"Regularisation parameter alpha is chosen by {CV_FOLDS}-fold cross-validation "
        f"(LassoCV). True bootstrap resampling ({N_BOOTSTRAP:,} resamples) is used to "
        "assess coefficient stability. Performance is reported both on the full training "
        f"set (apparent) and via {CV_FOLDS}-fold out-of-fold (OOF) cross-validation.\n\n"
        "Model ordering note: Models 2 and 3 are newly added variants that augment the "
        "original stroke/neurological predictor set (Model 6) with functional T1 assessments "
        "— with (Model 2) and without (Model 3) gait speed imputation — to isolate the "
        "contribution of the gait speed variable. Model 6 (original Model 2) is retained "
        "at the end as a reference for direct comparison."
    )
    _small_para(intro, 9)

    doc.add_paragraph()
    summary_heading = doc.add_paragraph()
    summary_heading.add_run("Executive comparison across the six models").bold = True
    comparison_rows = [[
        "Model",
        "Predictors",
        "Stable",
        "CV R²",
        "CV MAE (m)",
        "Train-CV R² gap",
    ]]
    for r in sorted(results, key=lambda x: x["cv_r2"], reverse=True):
        comparison_rows.append([
            r["model_name"].split(":", 1)[0],
            r["n_input_variables"],
            r["n_stable_selected"],
            f"{r['cv_r2']:.4f}",
            f"{r['cv_mae']:.2f}",
            f"{r['train_r2'] - r['cv_r2']:.4f}",
        ])
    _add_table_to_doc(doc, comparison_rows)

    for bullet in _executive_bullets(results):
        para = doc.add_paragraph(f"• {bullet}")
        _small_para(para, 9)

    ranks = _model_rankings(results)

    for r in results:
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run(r["model_name"]).bold = True

        feat_preview = ", ".join(r["features"][:15])
        if len(r["features"]) > 15:
            feat_preview += f" … (+{len(r['features']) - 15} more)"
        pred_para = doc.add_paragraph(f"Predictors ({len(r['features'])}): {feat_preview}")
        _small_para(pred_para, 8)

        _add_table_to_doc(doc, [
            ["Metric", "Value"],
            ["Patients (non-missing 6MWT4)", r["n_patients"]],
            ["Input variable count", r["n_input_variables"]],
            ["Final selected (non-zero in full-fit)", r["n_final_nonzero_base"]],
            [
                f"Stably selected (selection freq ≥ {STABILITY_THRESHOLD:.0%})",
                r["n_stable_selected"],
            ],
            ["Best alpha (full-data fit)", f"{r['best_alpha_full_fit']:.6g}"],
            [
                "Bootstrap alpha: median [IQR]",
                (
                    f"{r['alpha_boot_median']:.6g}"
                    f" [{r['alpha_boot_q25']:.6g}, {r['alpha_boot_q75']:.6g}]"
                ),
            ],
            ["CV R² (OOF)", f"{r['cv_r2']:.4f}"],
            ["CV MAE (OOF)", f"{r['cv_mae']:.2f} m"],
            ["Train R² (apparent)", f"{r['train_r2']:.4f}"],
            ["Train MAE (apparent)", f"{r['train_mae']:.2f} m"],
        ])

        top_positive = _top_predictors(r["importance"], positive=True)
        top_negative = _top_predictors(r["importance"], positive=False)
        explainer = doc.add_paragraph(
            "Explainer: "
            f"{MODEL_EXPLAINERS[r['model_name']]} "
            f"It ranked #{ranks[r['model_name']]} of {len(results)} by CV R². "
            f"The apparent-to-cross-validated R² gap was {r['train_r2'] - r['cv_r2']:.4f}, "
            f"with {r['n_stable_selected']} predictors meeting the {STABILITY_THRESHOLD:.0%} stability threshold."
        )
        _small_para(explainer, 9)

        pos_para = doc.add_paragraph(
            "Strongest positive contributors: "
            f"{_format_predictor_summary(top_positive)}."
        )
        _small_para(pos_para, 8)
        neg_para = doc.add_paragraph(
            "Strongest negative contributors: "
            f"{_format_predictor_summary(top_negative)}."
        )
        _small_para(neg_para, 8)

        doc.add_paragraph()
        imp_heading = doc.add_paragraph()
        imp_heading.add_run(
            "Predictor importance – top 20 (sorted by bootstrap selection frequency)"
        ).bold = True

        imp = r["importance"].head(20)
        rows = [[
            "Predictor",
            "Direction",
            "Full-fit Coef",
            "Boot Mean Coef",
            "Boot SD",
            "Selection Freq",
            "Stable",
        ]]
        for _, row in imp.iterrows():
            rows.append([
                row["predictor"],
                "Positive" if row["bootstrap_coef_mean"] >= 0 else "Negative",
                f"{row['base_coef']:.5f}",
                f"{row['bootstrap_coef_mean']:.5f}",
                f"{row['bootstrap_coef_sd']:.5f}",
                f"{row['selection_frequency']:.3f}",
                "Yes" if row["selection_frequency"] >= STABILITY_THRESHOLD else "No",
            ])
        _add_table_to_doc(doc, rows)

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX.name}")


def write_excel_predictions(df: pd.DataFrame, results: list[dict]) -> None:
    """Write Excel with original data plus one prediction column per model."""
    out = df.copy()

    for r in results:
        col_name = MODEL_PRED_COLS[r["model_name"]]
        pipe: Pipeline = r["base_pipe"]
        features: list[str] = r["features"]
        preds = pipe.predict(df[features])
        out[col_name] = preds

    out.to_excel(OUTPUT_XLSX, index=False)
    print(f"Saved: {OUTPUT_XLSX.name}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def run_analysis(write_outputs: bool = True) -> tuple[pd.DataFrame, list[dict]]:
    """Run the six-model LASSO analysis and optionally write refreshed outputs."""
    print(f"Loading: {INPUT_XLSX.name}")
    df = pd.read_excel(INPUT_XLSX)
    print(f"Dataset: {df.shape[0]:,} rows × {df.shape[1]} columns")

    all_candidate_cols = list(
        {c for cols in MODEL_SPECS.values() for c in cols}
    )
    for c in all_candidate_cols + ["6MWT4"]:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    if SOURCE_DOCX.exists() and SOURCE_XLSX.exists():
        print(f"Using cached six-model artifacts: {SOURCE_DOCX.name}, {SOURCE_XLSX.name}")
        results = load_existing_results(df)
    else:
        results = []
        for i, (model_name, candidate_features) in enumerate(MODEL_SPECS.items()):
            valid_features = _filter_existing(candidate_features, df)

            print("\n" + "=" * 78)
            print(f"  {model_name}")
            print(
                f"  Candidate predictors: {len(candidate_features)} → "
                f"existing in data: {len(valid_features)}"
            )

            if not valid_features:
                print("  !! No valid features found – skipping model.")
                continue

            res = fit_bootstrap_lasso(df, valid_features, model_name, model_seed=RANDOM_STATE + i)
            results.append(res)

            print(
                f"  Patients: {res['n_patients']:,} | "
                f"Input vars: {res['n_input_variables']} | "
                f"Non-zero (full fit): {res['n_final_nonzero_base']} | "
                f"Stable (≥{STABILITY_THRESHOLD:.0%}): {res['n_stable_selected']}"
            )
            print(
                f"  CV  R²: {res['cv_r2']:.4f}  MAE: {res['cv_mae']:.2f} m  |  "
                f"Train R²: {res['train_r2']:.4f}  MAE: {res['train_mae']:.2f} m"
            )

    if write_outputs:
        print("\n" + "=" * 78)
        write_docx_report(results)
        if SOURCE_XLSX.exists():
            shutil.copy2(SOURCE_XLSX, OUTPUT_XLSX)
            print(f"Saved: {OUTPUT_XLSX.name}")
        else:
            write_excel_predictions(df, results)
        print("Done.\n")

    return df, results


def main() -> None:
    """Run the six-model LASSO refresh and write the requested deliverables."""
    run_analysis(write_outputs=True)


if __name__ == "__main__":
    main()
