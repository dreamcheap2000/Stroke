#!/usr/bin/env python3
"""
IPCW_Code_202608082227.py
Unadjusted vs IPCW-adjusted LassoCV predictor coefficients for 6MWT4.
Uses pre-computed IPCW weights from Tip_Over_Analysis_artifacts_out_202608081955.csv.
Outputs: IPCW_Table_202608082227.docx
"""

import os
import numpy as np
import pandas as pd
from sklearn.linear_model import LassoCV
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 1. Load data ──────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
artifacts_df = pd.read_csv(os.path.join(SCRIPT_DIR,
                            'Tip_Over_Analysis_artifacts_out_202608081955.csv'))

# ── 2. Feature sets (same as the reference code) ──────────────────────────────
demographics = ['cov_Age', 'cov_Sex, F0 M1']
stroke_info  = ['cov_Dissection', 'cov_ACA', 'cov_Undetermined',
                'cov_HemorrhageStroke', 'cov_Side_Right', 'cov_Side_Left',
                'cov_Side_Bilateral', 'cov_Loc_CortSub', 'cov_Loc_Subcortical',
                'cov_Loc_Infratentorial', 'cov_LVS', 'cov_LVO']
comorbidities = ['cov_AF', 'cov_DM', 'cov_HTN', 'cov_Dyslipidemia', 'cov_CAD',
                 'cov_CKD', 'cov_RestrictiveLung', 'cov_GIUlcer',
                 'cov_LiverCirrhosis', 'cov_Hepatitis', 'cov_Parkinsonism',
                 'cov_Malignancy', 'cov_OldStroke', 'cov_Dementia',
                 'cov_Psychiatric', 'cov_Gout']
nihss_out_cols = ['cov_ConsOut', 'cov_AnswerOut', 'cov_OrderOut', 'cov_EOMOut',
                  'cov_VisualOut', 'cov_FacialOut', 'cov_LUOut', 'cov_RUOut',
                  'cov_LLOut', 'cov_RLOut', 'cov_Coordinateout', 'cov_SensoryOut',
                  'cov_LanguageOut', 'cov_ArticulateOut', 'cov_NeglectOut']
acute_complications = ['cov_Pneumonia', 'cov_UTI', 'cov_GIB', 'cov_Cellulitis']
acute_treatment     = ['cov_tPA', 'cov_IA', 'cov_tPAIA']
functional_t1_plus_gs = ['cov_MRS1', 'cov_BI1', 'cov_FOIS1', 'cov_MNA1',
                          'cov_EuroQoL5D1', 'cov_IADL1', 'cov_BBS1',
                          'cov_Gait_Speed_1', 'cov_FuglUE1', 'cov_FuglSEN1',
                          'cov_CCAT1']

features_raw = (demographics + stroke_info + comorbidities +
                functional_t1_plus_gs + nihss_out_cols +
                acute_complications + acute_treatment)
features = [f for f in features_raw if f in artifacts_df.columns]
missing_feats = set(features_raw) - set(features)
if missing_feats:
    print(f'Features not found (excluded): {missing_feats}')

target = 'outcome_6mwt4'

# ── 3. Completer subset ───────────────────────────────────────────────────────
model_df = artifacts_df[artifacts_df['completer'] == 1].copy()
model_df = model_df[features + [target, 'ipcw']].dropna()
print(f'Sample size (completers with full data): N = {len(model_df)}')

X_raw = model_df[features].values
y     = model_df[target].values
w     = model_df['ipcw'].values

scaler = StandardScaler()
X      = scaler.fit_transform(X_raw)

# ── 4. Fit unadjusted and IPCW-adjusted LassoCV ───────────────────────────────
RANDOM_STATE = 42
CV_FOLDS     = 5
MAX_ITER     = 20000

lasso_unadj = LassoCV(cv=CV_FOLDS, random_state=RANDOM_STATE, max_iter=MAX_ITER)
lasso_unadj.fit(X, y)
print(f'Unadjusted  α={lasso_unadj.alpha_:.4f}')

lasso_ipcw = LassoCV(cv=CV_FOLDS, random_state=RANDOM_STATE, max_iter=MAX_ITER)
lasso_ipcw.fit(X, y, sample_weight=w)
print(f'IPCW-adj    α={lasso_ipcw.alpha_:.4f}')

coef_unadj = lasso_unadj.coef_
coef_ipcw  = lasso_ipcw.coef_

# ── 5. Bootstrap CIs and p-values (B = 1 000 resamples) ──────────────────────
B  = 1000
n  = len(y)
rng = np.random.default_rng(RANDOM_STATE)

boot_unadj = np.zeros((B, len(features)))
boot_ipcw  = np.zeros((B, len(features)))

for b in range(B):
    idx = rng.integers(0, n, size=n)
    Xb, yb, wb = X[idx], y[idx], w[idx]

    m_u = LassoCV(cv=CV_FOLDS, random_state=RANDOM_STATE, max_iter=MAX_ITER)
    m_u.fit(Xb, yb)
    boot_unadj[b] = m_u.coef_

    m_i = LassoCV(cv=CV_FOLDS, random_state=RANDOM_STATE, max_iter=MAX_ITER)
    m_i.fit(Xb, yb, sample_weight=wb)
    boot_ipcw[b] = m_i.coef_

    if (b + 1) % 100 == 0:
        print(f'  Bootstrap {b + 1}/{B}')

def ci95(boot_coefs_col):
    lo = np.percentile(boot_coefs_col, 2.5)
    hi = np.percentile(boot_coefs_col, 97.5)
    return lo, hi

def pvalue(point_est, boot_coefs_col):
    """Two-sided bootstrap p-value via sign-flip proportion."""
    # Proportion of bootstrap estimates on the opposite side of zero from the point estimate
    if point_est >= 0:
        p = np.mean(boot_coefs_col <= 0)
    else:
        p = np.mean(boot_coefs_col >= 0)
    return float(np.clip(2 * p, 0, 1))

# ── 6. Assemble results table ─────────────────────────────────────────────────
rows = []
for i, feat in enumerate(features):
    c_u = coef_unadj[i]
    c_i = coef_ipcw[i]
    # Only include if non-zero in either model
    if c_u == 0 and c_i == 0:
        continue
    lo_u, hi_u = ci95(boot_unadj[:, i])
    lo_i, hi_i = ci95(boot_ipcw[:, i])
    p_u = pvalue(c_u, boot_unadj[:, i])
    p_i = pvalue(c_i, boot_ipcw[:, i])
    rows.append({
        'Predictor'         : feat,
        'beta_unadj'        : c_u,
        'ci_lo_unadj'       : lo_u,
        'ci_hi_unadj'       : hi_u,
        'p_unadj'           : p_u,
        'beta_ipcw'         : c_i,
        'ci_lo_ipcw'        : lo_i,
        'ci_hi_ipcw'        : hi_i,
        'p_ipcw'            : p_i,
    })

results_df = pd.DataFrame(rows).sort_values('beta_ipcw', ascending=False).reset_index(drop=True)
print(f'\nNon-zero predictors: {len(results_df)}')

# ── 7. Write DOCX table ───────────────────────────────────────────────────────
def fmt_coef(v):
    return f'{v:.3f}' if v != 0 else '0.000'

def fmt_ci(lo, hi):
    return f'[{lo:.3f}, {hi:.3f}]'

def fmt_p(p):
    if p < 0.001:
        return '<0.001'
    return f'{p:.3f}'

def shade_row(row, hex_color):
    """Apply background shading to a table row."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    trPr.append(shd)

def set_cell_font(cell, bold=False, size_pt=9, color_rgb=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = bold
            run.font.size = Pt(size_pt)
            if color_rgb:
                run.font.color.rgb = RGBColor(*color_rgb)

doc = Document()

# Title
title_para = doc.add_paragraph()
title_run  = title_para.add_run(
    'Table. Unadjusted and IPCW-Adjusted LassoCV Predictor Coefficients for 6MWT at Discharge (T4)')
title_run.font.bold = True
title_run.font.size = Pt(10)

# Merged header layout:
# Row 0 (span): Predictor | (blank) | Unadjusted (span 3) | IPCW-Adjusted (span 3)
# Row 1: Predictor | β | 95% CI | p-value | β | 95% CI | p-value
NUM_COLS = 7
table = doc.add_table(rows=2, cols=NUM_COLS)
table.style = 'Table Grid'

# ---- header row 0 ----
hdr0 = table.rows[0]
shade_row(hdr0, '2E4057')  # dark blue

def write_header_cell(cell, text, bold=True):
    cell.text = text
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.bold = bold
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

write_header_cell(hdr0.cells[0], 'Predictor')
write_header_cell(hdr0.cells[1], 'Unadjusted')
# merge cols 1-3
hdr0.cells[1].merge(hdr0.cells[2])
hdr0.cells[2].merge(hdr0.cells[3])
write_header_cell(hdr0.cells[4], 'IPCW-Adjusted')
hdr0.cells[4].merge(hdr0.cells[5])
hdr0.cells[5].merge(hdr0.cells[6])

# ---- header row 1 ----
hdr1 = table.rows[1]
shade_row(hdr1, '4A90D9')  # medium blue
sub_headers = ['Predictor', 'β', '95% CI', 'p-value', 'β', '95% CI', 'p-value']
for j, txt in enumerate(sub_headers):
    write_header_cell(hdr1.cells[j], txt)

# ---- data rows ----
for ri, row_data in results_df.iterrows():
    row = table.add_row()
    fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
    shade_row(row, fill)

    vals = [
        row_data['Predictor'].replace('cov_', ''),
        fmt_coef(row_data['beta_unadj']),
        fmt_ci(row_data['ci_lo_unadj'], row_data['ci_hi_unadj']),
        fmt_p(row_data['p_unadj']),
        fmt_coef(row_data['beta_ipcw']),
        fmt_ci(row_data['ci_lo_ipcw'], row_data['ci_hi_ipcw']),
        fmt_p(row_data['p_ipcw']),
    ]
    for j, v in enumerate(vals):
        cell = row.cells[j]
        cell.text = v
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(9)
                # Highlight significant IPCW p-value
                if j == 6 and row_data['p_ipcw'] < 0.05:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

# ---- Legend ----
doc.add_paragraph()
legend = doc.add_paragraph()
legend.add_run('Legend. ').font.bold = True
legend.add_run(
    'β = standardized LassoCV regression coefficient (StandardScaler applied to all predictors). '
    '95% CI = bootstrap 95% confidence interval (1 000 resamples, percentile method). '
    'p-value = two-sided bootstrap p-value. '
    'Unadjusted = LassoCV fitted on completers without sample weights. '
    'IPCW-Adjusted = LassoCV fitted on completers using stabilized inverse-probability-of-completion weights '
    '(pre-computed in Tip_Over_Analysis_artifacts_out_202608081955.csv). '
    'Only predictors with a non-zero coefficient in at least one model are shown. '
    'Red bold p-values indicate statistical significance (p < 0.05) in the IPCW-adjusted model.'
)
for run in legend.runs:
    if not run.bold:
        run.font.size = Pt(8)
    else:
        run.font.size = Pt(8)

# ── 8. Save ───────────────────────────────────────────────────────────────────
out_docx = os.path.join(SCRIPT_DIR, 'IPCW_Table_202608082227.docx')
doc.save(out_docx)
print(f'\nTable saved to: {out_docx}')
print('\nPreview of results:')
print(results_df[['Predictor', 'beta_unadj', 'p_unadj', 'beta_ipcw', 'p_ipcw']].to_string(index=False))
