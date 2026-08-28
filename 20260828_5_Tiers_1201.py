#!/usr/bin/env python3
"""
Tiered IPCW-based 6MWT4 extrapolation workflow.

This script treats the final patient-level output as a composite workflow rather
than as one combined model:

1. Build stabilized inverse-probability-of-completion weights (IPCW) for each
   feature tier using a completion model.
2. For each tier and each walking scenario (Best / Worst), select the best
   binary walking classifier by cross-validated balanced accuracy.
3. Fit a weighted Ridge regression on observed 6MWT4 among completers only.
4. Apply the scenario-specific classifier to non-completers to decide whether
   they are predicted to walk, then use the weighted Ridge model to generate a
   scenario-specific 6MWT4 prediction for those patients.
5. Save a patient-level Excel file plus a DOCX summary documenting the exact
   binary and regression models used for each tier/scenario combination.
"""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from sklearn.base import clone
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import ExtraTreesClassifier, RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.inspection import permutation_importance
from sklearn.linear_model import LogisticRegression, Ridge
from sklearn.metrics import (
    accuracy_score,
    balanced_accuracy_score,
    f1_score,
    mean_absolute_error,
    precision_score,
    r2_score,
    recall_score,
)
from sklearn.model_selection import KFold, StratifiedKFold, cross_val_predict, cross_validate
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler

# ─────────────────────────────────────────────────────────────────────────────
# File locations
# ─────────────────────────────────────────────────────────────────────────────

ROOT = Path(__file__).resolve().parent
INPUT_XLSX = ROOT / "20260826_DeID.xlsx"
OUTPUT_XLSX = ROOT / "20260828_DeID_IPCW_Predicted_6MWT.xlsx"
OUTPUT_REPORT = ROOT / "20260828_5_Tiers_1201.docx"

RANDOM_STATE = 42

# ─────────────────────────────────────────────────────────────────────────────
# Predictor groups carried forward from the prior analysis scripts
# ─────────────────────────────────────────────────────────────────────────────

DEMOGRAPHICS_PAC = ["Age", "Sex, F0 M1"]
STROKE_INFO = [
    "Dissection",
    "ACA",
    "Undetermined",
    "HemorrhageStroke",
    "LVS",
    "LVO",
    "Side_Right",
    "Side_Left",
    "Side_Bilateral",
    "Loc_CortSub",
    "Loc_Subcortical",
    "Loc_Infratentorial",
]
COMORBIDITIES_PAC = [
    "AF",
    "DM",
    "HTN",
    "Dyslipidemia",
    "CAD",
    "CKD",
    "RestrictiveLung",
    "GIUlcer",
    "LiverCirrhosis",
    "Hepatitis",
    "Parkinsonism",
    "Malignancy",
    "OldStroke",
    "Dementia",
    "Psychiatric",
    "Gout",
]
ACUTE_COMPLICATIONS_PAC = ["Pneumonia", "UTI", "GIB", "Cellulitis"]
FUNCTIONAL_T1_PLUS_GS_IMPUTED = [
    "MRS1",
    "BI1",
    "FOIS1",
    "MNA1",
    "EuroQoL5D1",
    "IADL1",
    "BBS1",
    "Gait_Speed_1_Imputed",
    "FuglUE1",
    "FuglSEN1",
]
NIHSS_OUT = [
    "ConsOut",
    "AnswerOut",
    "OrderOut",
    "EOMOut",
    "VisualOut",
    "FacialOut",
    "LUOut",
    "RUOut",
    "LLOut",
    "RLOut",
    "Coordinateout",
    "SensoryOut",
    "LanguageOut",
    "ArticulateOut",
    "NeglectOut",
]
NIHSS_IN = [
    "ConsIn",
    "AnswerIn",
    "OrderIn",
    "EOMIn",
    "VisualIn",
    "FaceIn",
    "LUIn",
    "RUIn",
    "LLIn",
    "RLIn",
    "CoordinateIn",
    "SensoryIn",
    "LanguageIn",
    "ArticulateIn",
    "NeglectIn",
]
T1T2_IMPROVEMENT = [
    "BI_T1T2_Change",
    "BBS_T1T2_Change",
    "MRS_T1T2_Change",
    "FOIS_T1T2_Change",
    "MNA_T1T2_Change",
    "IADL_T1T2_Change",
    "FuglUE_T1T2_Change",
    "FuglSEN_T1T2_Change",
    "EuroQoL5D_T1T2_Change",
]


def _remove(lst: list[str], item: str) -> list[str]:
    return [x for x in lst if x != item]


TIERS: dict[str, list[str]] = {
    "Tier 1": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED,
    "Tier 2": DEMOGRAPHICS_PAC + _remove(FUNCTIONAL_T1_PLUS_GS_IMPUTED, "Gait_Speed_1_Imputed"),
    "Tier 3": (
        DEMOGRAPHICS_PAC
        + FUNCTIONAL_T1_PLUS_GS_IMPUTED
        + COMORBIDITIES_PAC
        + STROKE_INFO
        + ACUTE_COMPLICATIONS_PAC
    ),
    "Tier 4": (
        DEMOGRAPHICS_PAC
        + _remove(FUNCTIONAL_T1_PLUS_GS_IMPUTED, "Gait_Speed_1_Imputed")
        + COMORBIDITIES_PAC
        + STROKE_INFO
        + ACUTE_COMPLICATIONS_PAC
    ),
    "Tier 5": DEMOGRAPHICS_PAC + FUNCTIONAL_T1_PLUS_GS_IMPUTED + T1T2_IMPROVEMENT,
}

SCENARIOS = ["6MWT_Best_Scenario", "6MWT_Worst_Scenario"]


# ─────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────────────────────────────────────

def _filter_existing(cols: list[str], df: pd.DataFrame) -> list[str]:
    return [c for c in cols if c in df.columns]


def _make_binary_cv() -> StratifiedKFold:
    return StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)


def _make_regression_cv() -> KFold:
    return KFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)


def _tier_token(tier_name: str) -> str:
    return tier_name.replace(" ", "_")


def _scenario_token(scenario_name: str) -> str:
    return scenario_name.replace("6MWT_", "").replace("_Scenario", "")


def _set_small_font(paragraph, size: int = 9) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(size)


def _add_table(doc: Document, rows: list[list[object]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
            for run in table.cell(i, j).paragraphs[0].runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True


def _weighted_mae(y_true: np.ndarray, y_pred: np.ndarray, weights: np.ndarray) -> float:
    return float(np.average(np.abs(y_true - y_pred), weights=weights))


# ─────────────────────────────────────────────────────────────────────────────
# Binary walking models
# ─────────────────────────────────────────────────────────────────────────────

def _binary_candidates(features: list[str]) -> dict[str, Pipeline]:
    numeric_pipe = Pipeline(
        [
            ("imp", SimpleImputer(strategy="median")),
            ("sc", StandardScaler()),
        ]
    )
    imputer_only_pipe = Pipeline([("imp", SimpleImputer(strategy="median"))])

    logistic = Pipeline(
        [
            ("prep", ColumnTransformer([("num", numeric_pipe, features)])),
            (
                "model",
                LogisticRegression(
                    max_iter=5000,
                    solver="liblinear",
                    class_weight="balanced",
                    random_state=RANDOM_STATE,
                ),
            ),
        ]
    )
    random_forest = Pipeline(
        [
            ("prep", ColumnTransformer([("num", imputer_only_pipe, features)])),
            (
                "model",
                RandomForestClassifier(
                    n_estimators=400,
                    class_weight="balanced",
                    random_state=RANDOM_STATE,
                    n_jobs=-1,
                ),
            ),
        ]
    )
    extra_trees = Pipeline(
        [
            ("prep", ColumnTransformer([("num", imputer_only_pipe, features)])),
            (
                "model",
                ExtraTreesClassifier(
                    n_estimators=400,
                    class_weight="balanced",
                    random_state=RANDOM_STATE,
                    n_jobs=-1,
                ),
            ),
        ]
    )
    return {
        "LogisticRegression(class_weight='balanced', solver='liblinear')": logistic,
        "RandomForestClassifier(n_estimators=400, class_weight='balanced')": random_forest,
        "ExtraTreesClassifier(n_estimators=400, class_weight='balanced')": extra_trees,
    }


def evaluate_binary(df: pd.DataFrame, features: list[str], scenario: str) -> dict:
    """
    Select the best scenario-specific walking classifier.

    The training target is already scenario-defined in the input workbook, so
    this step chooses the best binary model for the current tier/scenario pair.
    """
    valid = _filter_existing(features, df)
    model_df = df[df[scenario].notna()].copy()
    X = model_df[valid]
    y = model_df[scenario].astype(int)

    cv = _make_binary_cv()
    cv_splits = list(cv.split(X, y))
    candidates = _binary_candidates(valid)

    leaderboard_rows = []
    for model_name, pipeline in candidates.items():
        scores = cross_validate(
            clone(pipeline),
            X,
            y,
            cv=cv_splits,
            n_jobs=-1,
            scoring=["accuracy", "balanced_accuracy", "precision", "recall", "f1"],
        )
        leaderboard_rows.append(
            {
                "model": model_name,
                "accuracy": float(np.mean(scores["test_accuracy"])),
                "balanced_accuracy": float(np.mean(scores["test_balanced_accuracy"])),
                "precision": float(np.mean(scores["test_precision"])),
                "recall": float(np.mean(scores["test_recall"])),
                "f1": float(np.mean(scores["test_f1"])),
            }
        )

    leaderboard = (
        pd.DataFrame(leaderboard_rows)
        .sort_values(["balanced_accuracy", "accuracy", "f1"], ascending=False)
        .reset_index(drop=True)
    )
    best_model_name = str(leaderboard.iloc[0]["model"])
    best_pipeline = clone(candidates[best_model_name])

    oof_predictions = cross_val_predict(clone(best_pipeline), X, y, cv=cv_splits, n_jobs=-1)
    metrics = {
        "accuracy": float(accuracy_score(y, oof_predictions)),
        "balanced_accuracy": float(balanced_accuracy_score(y, oof_predictions)),
        "precision": float(precision_score(y, oof_predictions, zero_division=0)),
        "recall": float(recall_score(y, oof_predictions, zero_division=0)),
        "f1": float(f1_score(y, oof_predictions, zero_division=0)),
    }

    best_pipeline.fit(X, y)

    # Permutation importance is summarized only to keep the final DOCX readable.
    fold_importances = []
    for fold_idx, (train_idx, test_idx) in enumerate(cv_splits):
        fold_pipeline = clone(candidates[best_model_name])
        fold_pipeline.fit(X.iloc[train_idx], y.iloc[train_idx])
        fold_importance = permutation_importance(
            fold_pipeline,
            X.iloc[test_idx],
            y.iloc[test_idx],
            scoring="balanced_accuracy",
            n_repeats=10,
            random_state=RANDOM_STATE + fold_idx,
            n_jobs=-1,
        )
        fold_importances.append(fold_importance.importances_mean)

    importance = (
        pd.DataFrame(
            {
                "predictor": valid,
                "importance_mean": np.vstack(fold_importances).mean(axis=0),
            }
        )
        .sort_values("importance_mean", ascending=False)
        .reset_index(drop=True)
    )

    return {
        "scenario": scenario,
        "features": valid,
        "n_total": int(len(model_df)),
        "n_positive": int(y.sum()),
        "n_negative": int((y == 0).sum()),
        "leaderboard": leaderboard,
        "best_model_name": best_model_name,
        "best_pipeline": best_pipeline,
        "metrics": metrics,
        "importance": importance,
    }


# ─────────────────────────────────────────────────────────────────────────────
# IPCW completion weighting
# ─────────────────────────────────────────────────────────────────────────────

def compute_ipcw_weights(df: pd.DataFrame, features: list[str]) -> dict:
    """
    Estimate stabilized IPCW weights for completers.

    Denominator model:
        P(completed PAC | tier features)
    Numerator model:
        P(completed PAC | Age) when Age exists, otherwise the first tier feature
    """
    valid = _filter_existing(features, df)
    completion_status = df["PAC_Program_Completion"].astype("string")
    eligible_mask = completion_status.notna()
    completed = completion_status.eq("Completed PAC program").astype(int)

    X = df.loc[eligible_mask, valid]
    y = completed.loc[eligible_mask].to_numpy()

    imputer = SimpleImputer(strategy="median")
    X_imp = pd.DataFrame(imputer.fit_transform(X), columns=valid, index=X.index)

    denom_pipe = Pipeline(
        [
            ("sc", StandardScaler()),
            ("lr", LogisticRegression(max_iter=5000, solver="lbfgs", C=1.0)),
        ]
    )
    denom_pipe.fit(X_imp, y)
    p_denom = np.clip(denom_pipe.predict_proba(X_imp)[:, 1], 1e-4, 1 - 1e-4)

    stabilizer_feature = "Age" if "Age" in X_imp.columns else valid[0]
    X_numer = X_imp[[stabilizer_feature]]
    numer_pipe = Pipeline(
        [
            ("sc", StandardScaler()),
            ("lr", LogisticRegression(max_iter=1000, solver="lbfgs", C=1e6)),
        ]
    )
    numer_pipe.fit(X_numer, y)
    p_numer = np.clip(numer_pipe.predict_proba(X_numer)[:, 1], 1e-4, 1 - 1e-4)

    raw_weights = np.where(y == 1, p_numer / p_denom, np.nan)
    completer_weights = raw_weights[~np.isnan(raw_weights)]
    lower, upper = np.nanpercentile(completer_weights, [1, 99])
    winsorized = np.where(~np.isnan(raw_weights), np.clip(raw_weights, lower, upper), np.nan)

    weights = pd.Series(np.nan, index=df.index, dtype="float64")
    weights.loc[X_imp.index] = winsorized

    return {
        "weights": weights,
        "features": valid,
        "denominator_model": "LogisticRegression(solver='lbfgs', C=1.0) on tier features",
        "numerator_model": (
            f"LogisticRegression(solver='lbfgs', C=1e6) on stabilizer feature: {stabilizer_feature}"
        ),
        "winsorization": "1st/99th percentile",
        "mean_weight": float(np.nanmean(winsorized)),
        "min_weight": float(np.nanmin(completer_weights)),
        "max_weight": float(np.nanmax(completer_weights)),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Weighted 6MWT4 regression plus patient-level extrapolation
# ─────────────────────────────────────────────────────────────────────────────

def predict_6mwt4_ipcw(
    df: pd.DataFrame,
    features: list[str],
    ipcw_result: dict,
    binary_result: dict,
) -> dict:
    """
    Fit the weighted completer-only 6MWT4 model and generate patient-level
    predictions for non-completers for the current tier/scenario pair.
    """
    valid = _filter_existing(features, df)
    completion_status = df["PAC_Program_Completion"].astype("string")

    # Weighted regression is trained only where 6MWT4 is observed among completers.
    completer_mask = completion_status.eq("Completed PAC program") & df["6MWT4"].notna()
    df_completers = df.loc[completer_mask].copy()
    weights = ipcw_result["weights"].loc[completer_mask].fillna(1.0).to_numpy()

    X_raw = df_completers[valid]
    y = df_completers["6MWT4"].to_numpy()
    regression_cv = list(_make_regression_cv().split(X_raw))
    oof_predictions = np.zeros(len(df_completers), dtype="float64")

    # Manual CV keeps IPCW sample weights aligned inside each training fold.
    for train_idx, test_idx in regression_cv:
        fold_imputer = SimpleImputer(strategy="median")
        X_train = fold_imputer.fit_transform(X_raw.iloc[train_idx])
        X_test = fold_imputer.transform(X_raw.iloc[test_idx])

        fold_model = Ridge(alpha=1.0)
        fold_model.fit(X_train, y[train_idx], sample_weight=weights[train_idx])
        oof_predictions[test_idx] = np.maximum(0, fold_model.predict(X_test))

    weighted_r2 = float(r2_score(y, oof_predictions, sample_weight=weights))
    weighted_mae = _weighted_mae(y, oof_predictions, weights)

    imputer = SimpleImputer(strategy="median")
    X_all = pd.DataFrame(imputer.fit_transform(X_raw), columns=valid, index=df_completers.index)
    final_model = Ridge(alpha=1.0)
    final_model.fit(X_all, y, sample_weight=weights)

    importance = (
        pd.DataFrame(
            {
                "predictor": valid,
                "coefficient": final_model.coef_,
                "abs_coefficient": np.abs(final_model.coef_),
            }
        )
        .sort_values("abs_coefficient", ascending=False)
        .reset_index(drop=True)
    )

    # Non-completers receive a scenario-specific walk prediction and then a
    # 6MWT4 prediction. Non-walkers keep a 0-valued predicted 6MWT4 so every
    # scenario/tier combination is populated for non-completers.
    noncompleter_mask = completion_status.ne("Completed PAC program") & completion_status.notna()
    df_noncompleters = df.loc[noncompleter_mask].copy()
    X_noncompleters = pd.DataFrame(
        imputer.transform(df_noncompleters[valid]),
        columns=valid,
        index=df_noncompleters.index,
    )

    walk_predictions = binary_result["best_pipeline"].predict(X_noncompleters).astype(int)
    predicted_6mwt4 = np.zeros(len(df_noncompleters), dtype="float64")
    walk_rows = walk_predictions == 1
    if walk_rows.any():
        predicted_6mwt4[walk_rows] = np.maximum(0, final_model.predict(X_noncompleters.loc[walk_rows]))

    return {
        "features": valid,
        "n_completers": int(completer_mask.sum()),
        "n_noncompleters": int(noncompleter_mask.sum()),
        "n_noncomp_pred_walk": int(walk_predictions.sum()),
        "n_noncomp_pred_no_walk": int((walk_predictions == 0).sum()),
        "weighted_r2": weighted_r2,
        "weighted_mae": weighted_mae,
        "model_name": "Ridge(alpha=1.0) with IPCW sample weights",
        "final_model": final_model,
        "imputer": imputer,
        "importance": importance,
        "noncomp_index": df_noncompleters.index.tolist(),
        "noncomp_walk_pred": walk_predictions,
        "noncomp_6mwt4_pred": predicted_6mwt4,
    }


# ─────────────────────────────────────────────────────────────────────────────
# DOCX summary
# ─────────────────────────────────────────────────────────────────────────────

def write_report(all_results: list[dict]) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("20260828 Tiered IPCW 6MWT4 Summary").bold = True

    intro = doc.add_paragraph(
        "This report documents the final composite prediction workflow: "
        "IPCW completion-weighting, a scenario-specific binary walking "
        "classifier, and a weighted Ridge regression for observed 6MWT4 among completers."
    )
    _set_small_font(intro)

    for result in all_results:
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run(f"{result['tier']} | {result['scenario_label']} Scenario").bold = True

        doc.add_paragraph(
            f"Final patient-level predictions were generated by: "
            f"{result['ipcw_completion']['denominator_model']} + "
            f"{result['binary']['best_model_name']} + "
            f"{result['regression']['model_name']}."
        )

        _add_table(
            doc,
            [
                ["Field", "Value"],
                ["Tier", result["tier"]],
                ["Scenario", result["scenario_label"]],
                ["IPCW denominator model", result["ipcw_completion"]["denominator_model"]],
                ["IPCW numerator model", result["ipcw_completion"]["numerator_model"]],
                ["IPCW winsorization", result["ipcw_completion"]["winsorization"]],
                ["Selected binary model", result["binary"]["best_model_name"]],
                ["Binary OOF balanced accuracy", f'{result["binary"]["metrics"]["balanced_accuracy"]:.3f}'],
                ["Binary OOF F1", f'{result["binary"]["metrics"]["f1"]:.3f}'],
                ["Weighted regression model", result["regression"]["model_name"]],
                ["Weighted OOF R²", f'{result["regression"]["weighted_r2"]:.4f}'],
                ["Weighted OOF MAE", f'{result["regression"]["weighted_mae"]:.1f}'],
                ["Non-completers predicted to walk", result["regression"]["n_noncomp_pred_walk"]],
                ["Non-completers predicted not to walk", result["regression"]["n_noncomp_pred_no_walk"]],
                ["Feature count", len(result["features"])],
            ],
        )

        feature_text = ", ".join(result["features"])
        feature_paragraph = doc.add_paragraph(f"Feature list: {feature_text}")
        _set_small_font(feature_paragraph, size=8)

        top_predictors = result["regression"]["importance"].head(10)
        _add_table(
            doc,
            [["Top weighted Ridge predictors", "|Coefficient|"]]
            + [
                [row["predictor"], f'{row["abs_coefficient"]:.4f}']
                for _, row in top_predictors.iterrows()
            ],
        )

        leaderboard = result["binary"]["leaderboard"]
        _add_table(
            doc,
            [["Binary candidates", "Accuracy", "Balanced Acc", "Precision", "Recall", "F1"]]
            + [
                [
                    row["model"],
                    f'{row["accuracy"]:.3f}',
                    f'{row["balanced_accuracy"]:.3f}',
                    f'{row["precision"]:.3f}',
                    f'{row["recall"]:.3f}',
                    f'{row["f1"]:.3f}',
                ]
                for _, row in leaderboard.iterrows()
            ],
        )

    doc.save(OUTPUT_REPORT)


# ─────────────────────────────────────────────────────────────────────────────
# Main execution
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    df = pd.read_excel(INPUT_XLSX)
    print(f"Loaded: {df.shape[0]} rows × {df.shape[1]} columns")

    # Coerce all analysis features into numeric form before modeling.
    all_candidate_features = (
        DEMOGRAPHICS_PAC
        + STROKE_INFO
        + COMORBIDITIES_PAC
        + ACUTE_COMPLICATIONS_PAC
        + FUNCTIONAL_T1_PLUS_GS_IMPUTED
        + NIHSS_OUT
        + NIHSS_IN
        + T1T2_IMPROVEMENT
    )
    for column in all_candidate_features:
        if column in df.columns:
            df[column] = pd.to_numeric(df[column], errors="coerce")
    df["6MWT4"] = pd.to_numeric(df.get("6MWT4", np.nan), errors="coerce")

    output_df = df.copy()
    all_results: list[dict] = []

    for tier_name, tier_features in TIERS.items():
        valid_features = _filter_existing(tier_features, df)
        if not valid_features:
            continue

        print(f"\n{'=' * 72}")
        print(f"Processing {tier_name} with {len(valid_features)} predictors")
        ipcw_completion = compute_ipcw_weights(df, valid_features)
        print(f"  IPCW mean weight: {ipcw_completion['mean_weight']:.4f}")

        for scenario in SCENARIOS:
            scenario_label = _scenario_token(scenario)
            print(f"  Scenario: {scenario_label}")

            binary_result = evaluate_binary(df, valid_features, scenario)
            print(
                f"    Binary model: {binary_result['best_model_name']} | "
                f"OOF balanced accuracy={binary_result['metrics']['balanced_accuracy']:.3f}"
            )

            regression_result = predict_6mwt4_ipcw(df, valid_features, ipcw_completion, binary_result)
            print(
                f"    Weighted Ridge: OOF R²={regression_result['weighted_r2']:.4f} | "
                f"OOF MAE={regression_result['weighted_mae']:.1f} | "
                f"predicted walkers={regression_result['n_noncomp_pred_walk']}"
            )

            walk_column = f"IPCW_Walk_{_tier_token(tier_name)}_{scenario_label}"
            pred_column = f"IPCW_6MWT4_{_tier_token(tier_name)}_{scenario_label}"
            output_df[walk_column] = pd.Series(pd.NA, index=output_df.index, dtype="Int64")
            output_df[pred_column] = np.nan

            noncomp_index = regression_result["noncomp_index"]
            output_df.loc[noncomp_index, walk_column] = regression_result["noncomp_walk_pred"]
            output_df.loc[noncomp_index, pred_column] = regression_result["noncomp_6mwt4_pred"]

            all_results.append(
                {
                    "tier": tier_name,
                    "scenario": scenario,
                    "scenario_label": scenario_label,
                    "features": valid_features,
                    "ipcw_completion": ipcw_completion,
                    "binary": binary_result,
                    "regression": regression_result,
                }
            )

    output_df.to_excel(OUTPUT_XLSX, index=False)
    print(f"\nSaved dataset: {OUTPUT_XLSX.name}")

    write_report(all_results)
    print(f"Saved report: {OUTPUT_REPORT.name}")

    print("\n=== SUMMARY ===")
    print(f"{'Tier':<10} {'Scenario':<10} {'BalAcc':>8} {'F1':>8} {'Wt R²':>10} {'Wt MAE':>10} {'Pred Walk':>10}")
    print("-" * 80)
    for result in all_results:
        print(
            f"{result['tier']:<10} {result['scenario_label']:<10} "
            f"{result['binary']['metrics']['balanced_accuracy']:>8.3f} "
            f"{result['binary']['metrics']['f1']:>8.3f} "
            f"{result['regression']['weighted_r2']:>10.4f} "
            f"{result['regression']['weighted_mae']:>10.1f} "
            f"{result['regression']['n_noncomp_pred_walk']:>10}"
        )


if __name__ == "__main__":
    main()
